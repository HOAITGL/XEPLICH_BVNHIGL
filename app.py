from models.schedule_lock import ScheduleLock

from functools import wraps

def session_login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect('/login')
        return f(*args, **kwargs)
    return decorated_function

import sys
import os
sys.path.insert(0, os.path.abspath(os.path.dirname(__file__)))
from flask import Flask, render_template, request, redirect, url_for, flash, session
from flask_login import login_required
from datetime import datetime, timedelta
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate, upgrade
from models import ScheduleSignature
from models.ScheduleSignature import ScheduleSignature
from extensions import db  # Sử dụng đối tượng db đã khởi tạo trong extensions.py
from openpyxl import Workbook
from io import BytesIO
from logging.handlers import RotatingFileHandler
from flask_migrate import Migrate
from models.permission import Permission
from models.unit_config import UnitConfig
from models.user import User
from models.hazard_config import HazardConfig
from datetime import date
import calendar
from models import User, Shift, HazardConfig, ChamCong
from sqlalchemy import text as sql_text
from datetime import date, datetime, timedelta
import calendar, unicodedata
from flask import jsonify
from flask import jsonify, request, session
from sqlalchemy import func
import unicodedata


from logging.handlers import RotatingFileHandler
import logging, os

def setup_logging(app):
    if not os.path.exists('logs'):
        os.mkdir('logs')
    log_handler = RotatingFileHandler('logs/activity.log', maxBytes=1000000, backupCount=5)
    log_handler.setLevel(logging.INFO)
    log_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    log_handler.setFormatter(log_formatter)
    
    # Tránh thêm nhiều handler nếu đã có
    if not any(isinstance(h, RotatingFileHandler) for h in app.logger.handlers):
        app.logger.addHandler(log_handler)
    
    app.logger.setLevel(logging.INFO)

app = Flask(__name__)

# Lấy DATABASE_URL từ môi trường (Render), nếu không có -> fallback SQLite
db_url = os.getenv('DATABASE_URL')
if not db_url:
    db_url = 'sqlite:///database.db'  # fallback local
elif db_url.startswith("postgres://"):
    # Render thường trả postgres:// nhưng SQLAlchemy cần postgresql://
    db_url = db_url.replace("postgres://", "postgresql://", 1)

app.config['SQLALCHEMY_DATABASE_URI'] = db_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = 'lichtruc2025'

db.init_app(app)
migrate = Migrate(app, db)

setup_logging(app)
app.logger.info('[BOOT] Logging system started')

from sqlalchemy import text

@app.before_first_request
def add_missing_columns():
    with app.app_context():
        # Danh sách cột cần kiểm tra
        required_columns = {
            'shift': [
                ('order', 'INTEGER DEFAULT 0')
            ],
            'user': [
                ('contract_type', 'TEXT'),
                ('signature_file', 'TEXT'),   # 👈 thêm dòng này
                ('start_year', 'INTEGER')     # 👈 thêm dòng này
            ]
        }

        for table, columns in required_columns.items():
            for col_name, col_type in columns:
                # Kiểm tra cột tồn tại
                if 'postgres' in db.engine.url.drivername:
                    # Postgres: kiểm tra qua information_schema
                    check_col = db.session.execute(sql_text(f"""
                        SELECT column_name FROM information_schema.columns
                        WHERE table_name='{table}' AND column_name='{col_name}';
                    """)).fetchall()
                    column_exists = len(check_col) > 0
                else:
                    # SQLite: kiểm tra qua PRAGMA
                    existing_cols = db.session.execute(sql_text(f"PRAGMA table_info({table});")).fetchall()
                    existing_col_names = [col[1] for col in existing_cols]
                    column_exists = col_name in existing_col_names

                # Nếu chưa có cột thì thêm
                if not column_exists:
                    db.session.execute(sql_text(f'ALTER TABLE {table} ADD COLUMN "{col_name}" {col_type};'))
                    db.session.commit()
                    print(f"✅ Đã thêm cột '{col_name}' vào bảng {table}.")

                    # Nếu là cột order trong shift → cập nhật giá trị mặc định
                    if table == 'shift' and col_name == 'order':
                        shifts = Shift.query.order_by(Shift.id).all()
                        for i, s in enumerate(shifts):
                            s.order = i
                        db.session.commit()
                        print("✅ Đã cập nhật giá trị mặc định cho cột 'order'.")

                    # Nếu là cột contract_type trong user → set mặc định 'biên chế'
                    if table == 'user' and col_name == 'contract_type':
                        db.session.execute(sql_text("UPDATE \"user\" SET contract_type = 'biên chế' WHERE contract_type IS NULL;"))
                        db.session.commit()
                        print("✅ Đã set mặc định contract_type = 'biên chế' cho tất cả user cũ.")
                else:
                    print(f"ℹ️ Cột '{col_name}' đã tồn tại trong bảng {table}, bỏ qua.")

        # 🔍 Debug: in danh sách cột của bảng user
        insp = db.inspect(db.engine)
        cols = [c["name"] for c in insp.get_columns("user")]
        print("📌 Các cột bảng user:", cols)

# ✅ Tạo bảng nếu thiếu (dùng cho Render khi không gọi __main__)
with app.app_context():
    from sqlalchemy import inspect
    inspector = inspect(db.engine)
    existing_tables = inspector.get_table_names()
    required_tables = {'user', 'permission'}

    if not required_tables.issubset(set(existing_tables)):
        db.create_all()
        print("✅ Đã tạo bảng user/permission trên Render.")
    else:
        print("✅ Các bảng chính đã tồn tại.")


migrate = Migrate(app, db)

with app.app_context():
    try:
        upgrade()
        print("✅ Đã tự động chạy flask db upgrade.")
    except Exception as e:
        print(f"❌ Lỗi khi upgrade database: {e}")


# ✅ Định nghĩa admin_required
from functools import wraps
from flask import redirect, url_for, flash
from flask_login import current_user

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated:
            flash("⚠️ Bạn cần đăng nhập.", "danger")
            return redirect(url_for('login'))
        if current_user.role != 'admin':
            flash("❌ Chức năng này chỉ dành cho quản trị viên.", "danger")
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function
    allowed_routes = ['login']
    if 'user_id' not in session and request.endpoint not in allowed_routes:
        return redirect('/login')

from flask_login import LoginManager

# Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

# Import models
from models.user import User
from models.shift import Shift
from models.schedule import Schedule
from models.ca_config import CaConfiguration
from models.clinic_room import ClinicRoom
from models.ca import Ca

# Import logic
from scheduler.logic import generate_schedule

# app.py
@app.context_processor
def inject_permissions():
    if 'user_id' not in session:
        return dict(user=None, allowed_modules=[])

    role = (session.get('role') or '').lower()
    user = {
        'name': session.get('name', 'Người dùng'),
        'role': session.get('role'),
        'department': session.get('department'),
    }

    FULL_ADMIN = [
        'trang_chu','xem_lich_truc','yeu_cau_cv_ngoai_gio','don_nghi_phep',
        'xep_lich_truc','tong_hop_khth','cham_cong','bang_cham_cong_2',
        'bang_cong_gop','bang_tinh_tien_truc','cau_hinh_ca_truc','thiet_lap_phong_kham',
        'nhan_su_theo_khoa','cau_hinh_tien_truc','cau_hinh_doc_hai','thiet_lap_khoa_hscc',
        'phan_quyen','danh_sach_cong_viec','unit_config','xem_log','doi_mat_khau'
    ]

    BASELINE = {
        'admin1': [
            'trang_chu','xem_lich_truc','yeu_cau_cv_ngoai_gio','don_nghi_phep',
            'xep_lich_truc','tong_hop_khth','cham_cong','bang_cham_cong_2',
            'bang_cong_gop','bang_tinh_tien_truc','cau_hinh_ca_truc','thiet_lap_phong_kham',
            'nhan_su_theo_khoa','cau_hinh_tien_truc','danh_sach_cong_viec','doi_mat_khau'
        ],
        'manager': [
            'trang_chu','xem_lich_truc','don_nghi_phep','xep_lich_truc','tong_hop_khth',
            'cham_cong','bang_cham_cong_2','bang_cong_gop','bang_tinh_tien_truc','doi_mat_khau'
        ],
        'user': ['trang_chu','xem_lich_truc','don_nghi_phep','doi_mat_khau']
    }

    # 1) Admin: luôn full
    if role == 'admin':
        return dict(user=user, allowed_modules=FULL_ADMIN)

    # 2) Vai trò khác: ưu tiên DB, nếu không có thì baseline
    rows = Permission.query.filter_by(user_id=session['user_id']).all()
    if rows:
        allowed = sorted({r.module_name for r in rows if getattr(r, 'can_access', True)})
    else:
        allowed = BASELINE.get(role, BASELINE['user'])

    return dict(user=user, allowed_modules=allowed)


@app.route('/module-config', methods=['GET', 'POST'])
def edit_module_config():
    if session.get('role') != 'admin':
        return "Bạn không có quyền truy cập chức năng này."

    config_path = os.path.join(os.path.dirname(__file__), 'modules_config.json')

    if request.method == 'POST':
        try:
            data = request.form.get('config_json')
            parsed = json.loads(data)  # kiểm tra JSON hợp lệ
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(parsed, f, indent=2, ensure_ascii=False)
            flash("✅ Đã cập nhật cấu hình phân hệ.", "success")
        except Exception as e:
            flash(f"❌ Lỗi khi cập nhật: {str(e)}", "danger")
        return redirect('/module-config')

    with open(config_path, encoding='utf-8') as f:
        current_config = f.read()

    return render_template('edit_module_config.html', config=current_config)

from flask_login import login_user, logout_user
from flask import redirect, url_for, flash
from flask import request, redirect, render_template, flash
from models.user import User

@app.route('/')
@login_required
def index():
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username).first()

        app.logger.info(f"[LOGIN] User '{username}' đăng nhập từ IP {request.remote_addr}")

        # Không tiết lộ tài khoản hay mật khẩu đúng/sai cụ thể
        if not user or user.password != password:
            flash("Tên đăng nhập hoặc mật khẩu không đúng.", "danger")
            return redirect(url_for('login'))

        # Đăng nhập thành công
        login_user(user)  # ✅ Bắt buộc nếu dùng Flask-Login
        session['user_id'] = user.id
        session['role'] = user.role
        session['department'] = user.department

        app.logger.info(f"[LOGIN] Tài khoản '{username}' đã đăng nhập từ IP {request.remote_addr}")
        flash('Đăng nhập thành công!', 'success')
        return redirect(url_for('index'))

    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash("Bạn đã đăng xuất.", "info")
    return redirect(url_for('login'))

# ================== Helpers Nghỉ phép / Phép năm ==================
from datetime import date, datetime, timedelta, time
import unicodedata  # (nếu cần cho chỗ khác)
from models import db
from models.shift import Shift
from models.user import User

# === Ca "Nghỉ phép" mã P ===
def get_or_create_leave_shift():
    """
    Trả về ca 'Nghỉ phép' (code='P'). Nếu chưa có thì tạo mới.
    """
    leave = Shift.query.filter_by(code="P").first()
    if not leave:
        leave = Shift(
            name="Nghỉ phép",
            code="P",
            start_time=time(7, 0),
            end_time=time(7, 0),
            duration=24.0,
        )
        db.session.add(leave)
        db.session.flush()  # có leave.id ngay trong phiên
    return leave

# === Ngày lễ (có thể thay bằng nạp từ DB) ===
HOLIDAYS = {
    date(2025, 1, 1),
    date(2025, 4, 30),
    date(2025, 5, 1),
    date(2025, 9, 2),
}

WORKDAYS = {0, 1, 2, 3, 4}  # Thứ 2..Thứ 6

def count_workdays(start_date: date, end_date: date, holidays=None, workdays=WORKDAYS) -> int:
    """
    Đếm số ngày làm việc từ start_date..end_date (bao gồm 2 đầu),
    loại trừ T7/CN và các ngày trong 'holidays'.
    """
    if not start_date or not end_date:
        return 0
    if start_date > end_date:
        start_date, end_date = end_date, start_date

    hs = set(holidays or [])
    d, total = start_date, 0
    while d <= end_date:
        if d.weekday() in workdays and d not in hs:
            total += 1
        d += timedelta(days=1)
    return total

# === Tổng phép theo thâm niên ===
def _guess_join_year(user, fallback_year: int) -> int:
    """
    Suy đoán năm vào làm từ các field thường gặp trên user.
    Nếu không có, trả về fallback_year để không làm hỏng phép tính.
    """
    for attr in ("start_year", "join_year", "year_joined", "start_work_year"):
        y = getattr(user, attr, None)
        if isinstance(y, int) and 1900 <= y <= 2100:
            return y
    for attr in ("start_date", "join_date", "hire_date"):
        d = getattr(user, attr, None)
        if d:
            try:
                return d.year
            except Exception:
                pass
    return fallback_year

def calc_total_leave_days(user: User, year: int, base: int = 12) -> int:
    """
    Quy tắc: 12 ngày cơ bản + mỗi 5 năm thâm niên +1 ngày.
    Ví dụ vào làm 2018, năm 2025: thâm niên 7 năm -> 12 + 1 = 13.
    """
    join_year = _guess_join_year(user, year)
    years_worked = max(year - join_year, 0)
    return base + (years_worked // 5)

def leave_balance_by_requests(user_id: int, year: int):
    """
    Tính (total, used, remaining) THEO ĐƠN NGHỈ:
      - total: theo thâm niên (calc_total_leave_days)
      - used : cộng 'Số ngày' (days_off) của từng đơn; nếu thiếu days_off hoặc đơn bắc qua năm -> tính ngày làm việc T2–T6 trong phần giao với năm.
      - remaining = max(total - used, 0)
    Tương thích mọi CSDL (không dùng strftime/extract trong SQL).
    """
    from models.leave_request import LeaveRequest

    user = User.query.get(user_id)
    if not user:
        return 0, 0, 0

    total = calc_total_leave_days(user, year)
    y_start = date(year, 1, 1)
    y_end = date(year + 1, 1, 1) - timedelta(days=1)  # inclusive

    def _business_days_between(d1: date, d2: date, workdays=(0, 1, 2, 3, 4)) -> int:
        cur, n = d1, 0
        while cur <= d2:
            if cur.weekday() in workdays:
                n += 1
            cur += timedelta(days=1)
        return n

    used = 0
    reqs = (
        LeaveRequest.query
        .filter(
            LeaveRequest.user_id == user_id,
            LeaveRequest.start_date <= y_end,
            LeaveRequest.end_date >= y_start,
        )
        .all()
    )

    for lv in reqs:
        if not lv.start_date or not lv.end_date:
            continue
        s = max(lv.start_date, y_start)
        e = min(lv.end_date, y_end)
        if s > e:
            continue

        # Ưu tiên dùng số ngày đã nhập (nếu đơn nằm trọn trong năm)
        if getattr(lv, "days_off", None) is not None and (lv.start_date >= y_start and lv.end_date <= y_end):
            try:
                used += int(lv.days_off)
            except Exception:
                used += _business_days_between(s, e)
        else:
            used += _business_days_between(s, e)

    remaining = max(total - used, 0)
    return total, used, remaining

# === Shim giữ tương thích tên cũ (nếu nơi khác còn gọi) ===
def leave_balance_by_schedule(user_id: int, year: int):
    """
    Giữ tương thích ngược: hiện dùng chung logic theo ĐƠN NGHỈ.
    """
    return leave_balance_by_requests(user_id, year)
# ===========================================================


def leave_balance_by_requests(user_id, year):
    from models.leave_request import LeaveRequest
    user = User.query.get(user_id)
    if not user: return 0,0,0
    total = calc_total_leave_days(user, year)
    y_start, y_end = date(year,1,1), date(year+1,1,1) - timedelta(days=1)

    def business_days_between(d1, d2, workdays=(0,1,2,3,4)):
        cur, n = d1, 0
        while cur <= d2:
            if cur.weekday() in workdays: n += 1
            cur += timedelta(days=1)
        return n

    used = 0
    reqs = (LeaveRequest.query
            .filter(LeaveRequest.user_id==user_id,
                    LeaveRequest.start_date <= y_end,
                    LeaveRequest.end_date >= y_start)
            .all())
    for lv in reqs:
        if not lv.start_date or not lv.end_date: continue
        s, e = max(lv.start_date, y_start), min(lv.end_date, y_end)
        if s > e: continue
        if getattr(lv, "days_off", None) is not None and (lv.start_date >= y_start and lv.end_date <= y_end):
            try: used += int(lv.days_off)
            except: used += business_days_between(s, e)
        else:
            used += business_days_between(s, e)

    return total, used, max(total - used, 0)


# ====== đặt ở đầu file app.py (nếu chưa có) ======
from datetime import date
from flask import render_template, request, session, flash
from sqlalchemy import desc
# leave_balance_by_schedule phải được định nghĩa sẵn (helper đã gửi trước đó)
# ==================================================

@app.route('/leaves')
def leaves_list():
    from models.leave_request import LeaveRequest
    from models.user import User
    from datetime import date
    from sqlalchemy import desc

    role    = (session.get('role') or '').strip()
    my_dept = (session.get('department') or '').strip()
    my_id   = session.get('user_id')

    # --- Danh sách khoa cho dropdown ---
    if role in ('admin', 'admin1'):
        dept_rows = (db.session.query(User.department)
                     .filter(User.department.isnot(None))
                     .distinct()
                     .order_by(User.department)
                     .all())
        # Admin có thêm "Tất cả"
        departments = ['Tất cả'] + [r[0] for r in dept_rows]
    else:
        departments = [my_dept] if my_dept else []

    # --- Khoa được chọn ---
    if role in ('admin', 'admin1'):
        selected_department = request.args.get('department')
        if not selected_department:
            selected_department = 'Tất cả'  # mặc định xem toàn bệnh viện
    else:
        selected_department = my_dept

    # --- Lấy danh sách đơn theo quyền/khoa ---
    q = LeaveRequest.query.join(User, LeaveRequest.user_id == User.id)

    if role in ('admin', 'admin1'):
        # Chỉ lọc khi chọn 1 khoa cụ thể
        if selected_department and selected_department != 'Tất cả':
            q = q.filter(User.department == selected_department)
    else:
        if not my_dept:
            flash("Tài khoản chưa có thông tin khoa.", "warning")
            return render_template(
                'leaves.html',
                leaves=[], departments=departments,
                selected_department=None,
                current_department=my_dept,
                current_role=role,
                user=None,
                leave_info=None,
                balances={}
            )
        q = q.filter(User.department == my_dept)

    leaves = q.order_by(desc(LeaveRequest.start_date), desc(LeaveRequest.id)).all()

    # --- Tính cân bằng phép theo user trong danh sách ---
    year = date.today().year
    balances = {}
    user_ids = {lv.user_id for lv in leaves}
    for uid in user_ids:
        t, u, r = leave_balance_by_requests(uid, year)
        balances[uid] = {"total": t, "used": u, "remaining": r}

    # --- Banner tóm tắt cho người đang đăng nhập ---
    if my_id:
        t0, u0, r0 = leave_balance_by_requests(my_id, year)
        leave_info = {"year": year, "total": t0, "used": u0, "remaining": r0}
        current_user = User.query.get(my_id)
    else:
        leave_info = None
        current_user = None

    return render_template(
        'leaves.html',
        leaves=leaves,
        departments=departments,
        selected_department=selected_department,
        current_department=my_dept,
        current_role=role,
        user=current_user,
        leave_info=leave_info,
        balances=balances
    )

@app.route('/admin/fix-weekend-leaves')
@login_required
def fix_weekend_leaves():
    # chỉ cho admin
    if session.get('role') not in ['admin', 'admin1']:
        flash("Bạn không có quyền.", "danger")
        return redirect('/')

    leave_shift = get_or_create_leave_shift()

    # Lấy toàn bộ bản ghi 'Nghỉ phép'
    rows = Schedule.query.filter(Schedule.shift_id == leave_shift.id).all()
    removed = 0
    for r in rows:
        if r.work_date.weekday() >= 5:
            db.session.delete(r)
            removed += 1
    db.session.commit()
    flash(f"✅ Đã xoá {removed} bản ghi 'Nghỉ phép' rơi vào Thứ 7/CN.", "success")
    return redirect('/schedule')


@app.route('/leaves/add', methods=['GET', 'POST'])
@login_required
def add_leave():
    from models.leave_request import LeaveRequest
    from models.user import User
    from models.schedule import Schedule
    from utils.unit_config import get_unit_config

    user_role = session.get('role')
    user_dept = session.get('department')
    current_user_id = session.get('user_id')

    # Danh sách khoa
    if user_role in ['admin', 'admin1']:
        departments = [d[0] for d in db.session.query(User.department)
                       .filter(User.department.isnot(None))
                       .distinct().all()]
    else:
        departments = [user_dept]

    # Khoa được chọn
    if request.method == 'POST':
        selected_department = request.form.get('department') if user_role in ['admin', 'admin1'] else user_dept
    else:
        selected_department = request.args.get('department') if user_role in ['admin', 'admin1'] else user_dept

    # Danh sách user theo khoa
    if user_role in ['admin', 'admin1']:
        users = (User.query
                 .filter(User.department == selected_department)
                 .order_by(User.name).all() if selected_department else [])
    else:
        # Non-admin chỉ được tạo đơn cho chính mình
        users = [db.session.get(User, current_user_id)]

    # Xử lý tạo đơn
    if request.method == 'POST' and 'user_id' in request.form:
        user_id_str = request.form.get('user_id', '').strip()
        if not user_id_str.isdigit():
            flash("❌ Vui lòng chọn nhân viên hợp lệ.", "danger")
            return redirect('/leaves/add')

        user_id = int(user_id_str)
        try:
            start_date = datetime.strptime(request.form['start_date'], '%Y-%m-%d').date()
            end_date   = datetime.strptime(request.form['end_date'],   '%Y-%m-%d').date()
        except Exception:
            flash("❌ Ngày bắt đầu/kết thúc không hợp lệ.", "danger")
            return redirect('/leaves/add')

        if start_date > end_date:
            flash("❌ Khoảng thời gian nghỉ không hợp lệ.", "danger")
            return redirect('/leaves/add')

        reason   = request.form.get('reason')
        location = request.form.get('location')

        # Thông tin bổ sung (nếu có)
        birth_day   = request.form.get('birth_day')
        birth_month = request.form.get('birth_month')
        birth_year  = request.form.get('birth_year')
        birth_date  = None
        if birth_day and birth_month and birth_year:
            birth_date_str = f"{birth_year}-{str(birth_month).zfill(2)}-{str(birth_day).zfill(2)}"
            try:
                birth_date = datetime.strptime(birth_date_str, '%Y-%m-%d').date()
            except Exception:
                birth_date = None

        # Cập nhật năm bắt đầu công tác (nếu người dùng nhập)
        if request.form.get('start_work_year'):
            try:
                start_work_year = int(request.form.get('start_work_year'))
                user_obj = db.session.get(User, user_id)
                if user_obj:
                    user_obj.start_year = start_work_year
            except Exception:
                pass

        # 🔒 Ràng buộc quyền: Non-admin chỉ cho chính mình & đúng khoa
        target_user = db.session.get(User, user_id)
        if not target_user:
            flash("❌ Nhân viên không tồn tại.", "danger")
            return redirect('/leaves/add')
        if user_role not in ['admin', 'admin1']:
            if user_id != current_user_id:
                flash("❌ Bạn chỉ được tạo đơn nghỉ cho chính mình.", "danger")
                return redirect('/leaves/add')
            if target_user.department != user_dept:
                flash("❌ Bạn không thể tạo đơn cho nhân viên khác khoa.", "danger")
                return redirect('/leaves/add')

        # ====== GIỚI HẠN THEO SỐ NGÀY PHÉP CÒN LẠI ======
        req_days = count_workdays(start_date, end_date, HOLIDAYS)  # số ngày xin (T2–T6, trừ lễ)
        year     = start_date.year
        total, used, remaining = leave_balance_by_schedule(user_id, year)

        # ==== CHÈN ĐOẠN NÀY VÀO ĐÂY ====
        year = start_date.year
        total, used, remaining = leave_balance_by_schedule(user_id, year)

        if remaining <= 0:
            flash(f"❌ Bạn đã dùng hết ngày phép năm {year} (đã dùng {used}/{total}).", "danger")
            return redirect(request.referrer or url_for('add_leave'))

        if req_days > remaining:
            flash(
                f"❌ Số ngày nghỉ phép còn lại của bạn là {remaining} ngày. "
                f"Bạn không thể thêm đơn nghỉ phép {req_days} ngày được "
                f"(đã dùng {used}/{total} ngày năm {year}).",
                "warning"
            )
            return redirect(request.referrer or url_for('add_leave'))

        # 1) Lưu đơn nghỉ
        leave = LeaveRequest(
            user_id=user_id,
            start_date=start_date,
            end_date=end_date,
            reason=reason,
            location=location,
            birth_date=birth_date
        )
        # (Nếu LeaveRequest có cột days_off, có thể set: leave.days_off = req_days)
        db.session.add(leave)

        # 2) Tự động chấm nghỉ -> gán ca "Nghỉ phép" (BỎ Thứ 7 & Chủ nhật)
        leave_shift = get_or_create_leave_shift()
        cur = start_date
        overwritten = 0
        created = 0
        while cur <= end_date:
            # 0..4 = Thứ 2..Thứ 6 ; 5 = Thứ 7 ; 6 = Chủ nhật
            if cur.weekday() >= 5:
                cur += timedelta(days=1)
                continue

            sched = Schedule.query.filter_by(user_id=user_id, work_date=cur).first()
            if sched:
                sched.shift_id = leave_shift.id
                overwritten += 1
            else:
                db.session.add(Schedule(user_id=user_id, shift_id=leave_shift.id, work_date=cur))
                created += 1
            cur += timedelta(days=1)

        # 2b) DỌN mọi lịch (bất kể ca nào) rơi vào T7/CN trong khoảng đơn -> để trống cuối tuần
        weekend_records = (Schedule.query
            .filter(
                Schedule.user_id == user_id,
                Schedule.work_date >= start_date,
                Schedule.work_date <= end_date
            ).all())
        removed_weekend = 0
        for s in weekend_records:
            if s.work_date.weekday() >= 5:  # 5=Thứ 7, 6=CN
                db.session.delete(s)
                removed_weekend += 1

        # 3) Commit
        try:
            db.session.commit()
            left_after = remaining - req_days
            msg = (f"✅ Đã tạo đơn nghỉ phép {req_days} ngày (T2–T6). "
                   f"Còn lại {left_after} ngày phép năm {year}. "
                   f"Cập nhật {overwritten} ngày, thêm mới {created} ngày."
                   )
            if removed_weekend:
                msg += f" Đã xoá {removed_weekend} lịch rơi vào T7/CN."
            flash(msg, "success")
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"[LEAVE_ADD_ERROR] {e}", exc_info=True)
            flash(f"❌ Lỗi khi lưu: {e}", "danger")

        return redirect('/leaves')

    # Dữ liệu hiển thị form
    unit = get_unit_config()
    return render_template(
        'add_leave.html',
        departments=departments,
        selected_department=selected_department,
        users=users,
        current_user_role=user_role,
        now=datetime.now()
    )

from models.leave_request import LeaveRequest

from datetime import datetime


@property
def full_name(self):
    return self.user.name

@property
def position(self):
    return self.user.position

@property
def birth_year(self):
    return self.birth_date.year if self.birth_date else ""

@property
def start_work_year(self):
    return self.user.start_year if hasattr(self.user, 'start_year') else ""

@property
def leave_days(self):
    return (self.end_date - self.start_date).days + 1

@app.route('/leaves/delete/<int:leave_id>', methods=['POST'])
def leaves_delete(leave_id):
    from models.leave_request import LeaveRequest

    role = session.get('role')
    my_id = session.get('user_id')

    leave = LeaveRequest.query.get_or_404(leave_id)

    # Quyền: admin/admin1 hoặc chính chủ đơn
    if role not in ('admin', 'admin1') and leave.user_id != my_id:
        return "Bạn không có quyền xoá đơn này.", 403

    db.session.delete(leave)
    db.session.commit()
    flash("🗑 Đã xoá đơn nghỉ.", "success")
    return redirect(request.referrer or url_for('leaves_list'))

from flask import request, render_template, redirect, session
from collections import defaultdict
import csv
import os
from models import User  # ✅ đúng

from models.work_request import WorkRequest

@app.route("/yeu-cau-xu-ly-cong-viec", methods=["GET", "POST"])
def yeu_cau_xu_ly_cong_viec():
    from collections import defaultdict
    from datetime import datetime

    if request.method == "POST":
        ngay_thang = request.form.get("ngay_thang")
        khoa = request.form.get("khoa")
        nguoi_yeu_cau = request.form.get("nguoi_yeu_cau")
        chu_ky = request.form.get("chu_ky")
        loi = request.form.get("loi")
        so_ho_so = request.form.get("so_ho_so", "")
        so_phieu = request.form.get("so_phieu", "")
        noi_dung = request.form.get("noi_dung", "")
        xac_nhan = request.form.get("xac_nhan")

        # Lấy số điện thoại người yêu cầu (nếu có)
        user_obj = User.query.filter_by(name=nguoi_yeu_cau).first()
        so_dien_thoai = user_obj.phone if user_obj and hasattr(user_obj, 'phone') else ""

        # Đánh dấu chữ ký
        def mark(name):
            return "✓" if xac_nhan == name else "✗"

        mark_hoa = mark("Hòa")
        mark_hiep = mark("Hiệp")
        mark_anh = mark("Ánh")
        mark_nam = mark("Nam")

        # Lưu vào bảng WorkRequest
        new_request = WorkRequest(
            ngay_thang=ngay_thang,
            khoa=khoa,
            loi=loi,
            so_ho_so=so_ho_so,
            so_phieu=so_phieu,
            noi_dung=noi_dung,
            nguoi_yeu_cau=nguoi_yeu_cau,
            so_dien_thoai=so_dien_thoai,
            chu_ky=chu_ky,
            hoa=mark_hoa,
            hiep=mark_hiep,
            anh=mark_anh,
            nam=mark_nam
        )
        db.session.add(new_request)
        db.session.commit()

        # Chuyển hướng sang danh sách để in
        return redirect("/")

    # --- GET: Load form ---
    staff_by_unit = defaultdict(list)
    users = User.query.filter(User.department != None).all()
    for user in users:
        if user.name and user.department:
            staff_by_unit[user.department].append(user.name)
    for dept in staff_by_unit:
        staff_by_unit[dept].sort()

    user_role = session.get('role')
    user_dept = session.get('department')

    if user_role in ['user', 'manager']:
        staff_by_unit_filtered = {user_dept: staff_by_unit.get(user_dept, [])}
        current_department = user_dept
    else:
        staff_by_unit_filtered = dict(staff_by_unit)
        current_department = ""

    # Lấy số điện thoại Nam (CNTT)
    nam_user = User.query.filter_by(name="Nam").first()
    nam_phone = nam_user.phone if nam_user and hasattr(nam_user, 'phone') else ""

    return render_template(
        "form.html",
        staff_by_unit=staff_by_unit_filtered,
        current_department=current_department,
        nam_phone=nam_phone
    )

@app.route('/api/user-phones')
def api_user_phones():
    users = User.query.filter(User.department == 'Phòng Kế hoạch TH - CNTT').all()
    result = {user.name: user.phone for user in users if user.phone}
    return result

@app.route("/danh-sach-yeu-cau")
def danh_sach_yeu_cau():
    # Lấy tất cả yêu cầu từ DB
    requests = WorkRequest.query.order_by(WorkRequest.id.desc()).all()

    # Tiêu đề bảng
    headers = [
        "Ngày tháng", "Khoa", "Người yêu cầu", "Lỗi",
        "Số hồ sơ", "Số phiếu", "Nội dung",
        "Hòa", "Hiệp", "Ánh", "Nam",
        "Chữ ký"
    ]

    # Dữ liệu kèm ID để xóa
    data = []
    for req in requests:
        data.append([
            req.id,                # ID dùng cho nút XÓA
            req.ngay_thang,
            req.khoa,
            req.nguoi_yeu_cau,
            req.loi,
            req.so_ho_so,
            req.so_phieu,
            req.noi_dung,
            req.hoa,
            req.hiep,
            req.anh,
            req.nam,
            req.chu_ky
        ])

    # Ngày hiện tại để hiển thị
    now = datetime.today()
    current_date = f"Gia Lai, ngày {now.day:02d} tháng {now.month:02d} năm {now.year}"

    return render_template(
        "danh_sach_yeu_cau.html",
        headers=headers,
        data=data,
        month=now.month,
        current_date=current_date
    )

@app.route("/xoa-yeu-cau/<int:req_id>", methods=["POST"])
def xoa_yeu_cau(req_id):
    from models.work_request import WorkRequest
    req = WorkRequest.query.get(req_id)
    if req:
        db.session.delete(req)
        db.session.commit()
    return redirect("/danh-sach-yeu-cau")


from flask import send_file
import pandas as pd

@app.route("/xuat-excel")
def xuat_excel():
    from models.work_request import WorkRequest
    import openpyxl
    from io import BytesIO
    from flask import send_file

    # Tạo file Excel mới
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "YeuCauCongViec"

    # Tiêu đề giống bảng hiển thị
    headers = [
        "Ngày tháng", "Khoa", "Người yêu cầu", "Lỗi",
        "Số hồ sơ", "Số phiếu", "Nội dung",
        "Hòa", "Hiệp", "Ánh", "Nam", "Chữ ký"
    ]
    ws.append(headers)

    # Lấy dữ liệu từ DB
    requests = WorkRequest.query.order_by(WorkRequest.id).all()
    for r in requests:
        ws.append([
            r.ngay_thang,
            r.khoa,
            r.nguoi_yeu_cau,
            r.loi,
            r.so_ho_so,
            r.so_phieu,
            r.noi_dung,
            r.hoa,
            r.hiep,
            r.anh,
            r.nam,
            r.chu_ky
        ])

    # Xuất file Excel
    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)

    return send_file(
        stream,
        as_attachment=True,
        download_name="yeu_cau_cong_viec.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

from flask import render_template
from datetime import datetime
from models.leave_request import LeaveRequest
import os

from datetime import timedelta
from datetime import datetime, timedelta

def working_days_inclusive(start_date, end_date):
    """Đếm số ngày làm việc (Thứ 2–Thứ 6), tính cả đầu & cuối."""
    d = start_date
    count = 0
    while d <= end_date:
        if d.weekday() < 5:  # 0..4 = Mon..Fri
            count += 1
        d += timedelta(days=1)
    return count

def end_date_from_workdays(start_date, days):
    """Lấy ngày kết thúc sau N ngày làm việc (T2–T6), coi start là ngày 1)."""
    d = start_date
    done = 0
    while True:
        if d.weekday() < 5:
            done += 1
            if done == days:
                return d
        d += timedelta(days=1)

def annual_leave_quota(user, as_of_date):
    """
    Chỉ tiêu phép năm = 12 + floor(thâm niên/5).
    Thâm niên = năm đang nghỉ - start_year (nếu có).
    """
    base = 12
    if not getattr(user, "start_year", None):
        return base
    years = as_of_date.year - int(user.start_year)
    extra = max(0, years // 5)
    return base + extra

# Map chức vụ viết tắt -> đầy đủ
POSITION_MAP = {
    "NV": "Nhân viên",
    "GĐ": "Giám đốc",
    "PGĐ": "Phó Giám đốc",
    "TP": "Trưởng phòng",
    "TK": "Trưởng Khoa",
    "PTK": "Phó Trưởng khoa",
    "PTP": "Phó Trưởng phòng",
    "ĐD": "Điều dưỡng",
    "ĐDT": "Điều dưỡng trưởng",
    "BS": "Bác sĩ",
    "KTV": "Kỹ thuật viên",
    "KTVT": "Kỹ thuật viên trưởng",
    "BV": "Bảo vệ",
    "HL": "Hộ lý",
    "LX": "Lái xe",
    "CV": "Chuyên viên",
    "CVC": "Chuyên viên chính"
}

@app.route('/leaves/print/<int:id>')
@login_required
def print_leave(id):
    from models.leave_request import LeaveRequest
    from utils.unit_config import get_unit_config

    leave = LeaveRequest.query.get_or_404(id)
    user = leave.user
    unit = get_unit_config()

    # Số ngày NGHỈ THỰC TẾ của đơn (bỏ T7 & CN)
    total_days = working_days_inclusive(leave.start_date, leave.end_date)

    # Chỉ tiêu phép năm theo thâm niên (12 + 1 ngày mỗi đủ 5 năm)
    quota_days = annual_leave_quota(user, leave.start_date)

    # Nếu nghỉ trọn quota bắt đầu từ ngày start, kết thúc là ngày:
    quota_end = end_date_from_workdays(leave.start_date, quota_days)

    # Chuẩn hoá chức vụ
    full_position = POSITION_MAP.get((user.position or '').strip(), user.position)

    return render_template(
        'leave_print.html',
        leave=leave,
        user=user,
        unit=unit,
        total_days=total_days,   # Số ngày công của đơn
        quota_days=quota_days,   # Chỉ tiêu năm theo thâm niên
        quota_end=quota_end,     # Ví dụ: nếu nghỉ đủ quota
        full_position=full_position,
        now=datetime.now()
    )

from flask import send_file
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

@app.route('/leaves/word/<int:leave_id>')
def export_leave_word(leave_id):
    from models.leave_request import LeaveRequest
    from models.user import User

    leave = LeaveRequest.query.get_or_404(leave_id)
    user = User.query.get(leave.user_id)

    document = Document()

    # Đặt font chữ mặc định cho toàn bộ văn bản
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Header: bảng 2 cột
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Pt(260)
    table.columns[1].width = Pt(260)
    cells = table.rows[0].cells

    # BÊNH VIỆN NHI TỈNH GIA LAI + Phòng ban (in hoa, in đậm, border bottom)
    p_left = cells[0].paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_left.add_run("BỆNH VIỆN NHI TỈNH GIA LAI\n")
    run.bold = True
    run.font.size = Pt(13)
    run = p_left.add_run(user.department.upper())
    run.bold = True
    run.font.size = Pt(14)

    # Border bottom cho đoạn bên trái
    p = p_left._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM + Độc lập - Tự do - Hạnh phúc
    p_right = cells[1].paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    run.bold = True
    run.font.size = Pt(14)
    run = p_right.add_run("Độc lập - Tự do - Hạnh phúc")
    run.bold = True
    run.italic = True
    run.font.size = Pt(14)
    run.font.underline = True

    document.add_paragraph()  # Dòng trống

    # Tiêu đề chính giữa
    title = document.add_paragraph("ĐƠN XIN NGHỈ PHÉP")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.runs[0]
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.text = run_title.text.upper()

    # Kính gửi (in nghiêng, đậm, thụt lề 1cm)
    p_kg = document.add_paragraph()
    p_kg.paragraph_format.left_indent = Pt(28)
    run_kg = p_kg.add_run("Kính gửi:")
    run_kg.bold = True
    run_kg.italic = True
    run_kg.font.size = Pt(14)

    # Danh sách kính gửi, thụt lề 3cm
    p_list = document.add_paragraph()
    p_list.paragraph_format.left_indent = Pt(85)
    p_list.paragraph_format.line_spacing = 1.4
    p_list.add_run("- Giám đốc Bệnh viện Nhi tỉnh Gia Lai\n")
    p_list.add_run("- Phòng Tổ chức - Hành chính quản trị\n")
    p_list.add_run(f"- {user.department}")

    # Thông tin người làm đơn
    p_name = document.add_paragraph()
    p_name.add_run("Tên tôi là: ").font.size = Pt(14)
    run_name = p_name.add_run(user.name.upper())
    run_name.bold = True
    run_name.font.size = Pt(14)
    p_name.add_run("    Sinh ngày: ").font.size = Pt(14)
    run_birth = p_name.add_run(leave.birth_date.strftime('%d/%m/%Y') if leave.birth_date else '')
    run_birth.bold = True
    run_birth.font.size = Pt(14)

    p_pos = document.add_paragraph()
    p_pos.add_run(f"Chức vụ: {user.position}    ").font.size = Pt(14)
    p_pos.add_run("Năm vào công tác: ......................").font.size = Pt(14)

    p_dep = document.add_paragraph(f"Đơn vị công tác: {user.department} - Bệnh viện Nhi tỉnh Gia Lai")
    p_dep.style.font.size = Pt(14)

    # Nội dung đơn
    document.add_paragraph(
        "Nay tôi làm đơn này trình Ban Giám đốc, Phòng Tổ chức - Hành chính quản trị xem xét và sắp xếp cho tôi được nghỉ phép.")
    document.add_paragraph(
        f"Thời gian nghỉ: từ ngày {leave.start_date.strftime('%d/%m/%Y')} đến ngày {leave.end_date.strftime('%d/%m/%Y')}.")
    document.add_paragraph(f"Lý do: {leave.reason}")
    document.add_paragraph(f"Nơi nghỉ phép: {leave.location}")
    document.add_paragraph("Tôi xin cam đoan sẽ bàn giao công việc đầy đủ và trở lại làm việc đúng thời gian quy định.")
    document.add_paragraph("Vậy kính mong các cấp giải quyết, tôi xin chân thành cảm ơn./.")

    # Footer ngày tháng căn phải, in nghiêng
    date_str = leave.start_date.strftime("Gia Lai, ngày %d tháng %m năm %Y") if leave.start_date else ""
    p_footer = document.add_paragraph(date_str)
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_footer.runs[0].italic = True
    p_footer.runs[0].font.size = Pt(14)

    # Bảng ký tên 3 cột căn giữa
    sign_table = document.add_table(rows=2, cols=3)
    sign_table.autofit = False
    widths = [Pt(180), Pt(180), Pt(180)]
    for idx, width in enumerate(widths):
        sign_table.columns[idx].width = width

    # Dòng 1
    cells = sign_table.rows[0].cells
    cells[0].text = user.department.upper()
    cells[1].text = "Người làm đơn"
    cells[2].text = "Giám đốc\nPhòng Tổ chức – HC QT"
    for cell in cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(14)
                run.bold = True

    # Dòng 2
    cells = sign_table.rows[1].cells
    cells[0].text = ""
    cells[1].text = "(Ký và ghi rõ họ tên)"
    cells[2].text = ""
    for cell in cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(14)
                run.italic = True

    # Tên người làm đơn căn giữa phía dưới bảng ký
    p_name_sign = document.add_paragraph(user.name.upper())
    p_name_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name_sign.runs[0].font.size = Pt(14)

    # Xuất file Word
    stream = BytesIO()
    document.save(stream)
    stream.seek(0)

    return send_file(
        stream,
        as_attachment=True,
        download_name=f"don_nghi_phep_{leave.id}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/test-export')
def test_export():
    return "Route hoạt động"

@app.route('/assign', methods=['GET', 'POST']) 
def assign_schedule():
    user_role = session.get('role')
    user_dept = session.get('department')

    # Lấy danh sách khoa theo role
    if user_role != 'admin':
        departments = [user_dept]
    else:
        departments = [d[0] for d in db.session.query(User.department)
                       .filter(User.department != None).distinct().all()]

    # Lấy khoa được chọn
    selected_department = (
        request.args.get('department') if request.method == 'GET'
        else request.form.get('department')
    )

    # Lấy danh sách nhân sự theo khoa
    users = User.query.filter_by(department=selected_department).all() if selected_department else []

    # --- Sắp xếp theo thứ tự chức danh ---
    priority_order = ['GĐ', 'PGĐ', 'TK', 'TP', 'PTK', 'PP', 'PK', 'BS', 'ĐDT', 'ĐD', 'KTV', 'NV', 'HL', 'BV', 'LX']

    def get_priority(pos):
        pos = pos.upper() if pos else ''
        for i, p in enumerate(priority_order):
            if p in pos:
                return i
        return len(priority_order)  # Nếu không khớp, cho xuống cuối

    users.sort(key=lambda x: get_priority(x.position))

    # **Sửa chỗ này: Lấy shifts theo thứ tự `order`**
    shifts = Shift.query.order_by(Shift.order).all()

    leaves = []

    # Xử lý POST (lưu lịch)
    if request.method == 'POST':
        start_date = datetime.strptime(request.form['start_date'], '%Y-%m-%d').date()
        end_date = datetime.strptime(request.form['end_date'], '%Y-%m-%d').date()
        duplicated_entries = []

        user_name = session.get('name')
        app.logger.info(f"[ASSIGN] User '{user_name}' phân lịch '{selected_department}' từ {start_date} đến {end_date}")

        # Lấy đơn nghỉ trong khoảng thời gian
        leaves = LeaveRequest.query.filter(
            LeaveRequest.start_date <= end_date,
            LeaveRequest.end_date >= start_date
        ).all()

        # Lưu từng checkbox lịch trực
        for checkbox in request.form.getlist('schedule'):
            user_id, shift_id = checkbox.split('-')
            user_id = int(user_id)
            shift_id = int(shift_id)

            current = start_date
            while current <= end_date:
                existing = Schedule.query.filter_by(
                    user_id=user_id,
                    shift_id=shift_id,
                    work_date=current
                ).first()

                if existing:
                    duplicated_entries.append(
                        f"{existing.user.name} đã có lịch ca {existing.shift.name} ngày {current.strftime('%d/%m/%Y')}"
                    )
                else:
                    new_schedule = Schedule(
                        user_id=user_id,
                        shift_id=shift_id,
                        work_date=current
                    )
                    db.session.add(new_schedule)

                current += timedelta(days=1)

        try:
            db.session.commit()
            if duplicated_entries:
                for message in duplicated_entries:
                    flash(f"⚠️ {message}", "danger")
            else:
                flash("✅ Đã lưu lịch thành công.", "success")
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"❌ Lỗi khi lưu lịch trực: {e}", exc_info=True)
            flash("❌ Lỗi máy chủ: không thể lưu lịch trực.", "danger")
            return redirect('/assign?department=' + selected_department)

        return redirect('/assign?department=' + selected_department)

    return render_template(
        'assign.html',
        departments=departments,
        selected_department=selected_department,
        users=users,
        shifts=shifts,
        leaves=leaves
    )


@app.route('/auto-assign')
def auto_assign_page():
    from models.user import User
    from models.shift import Shift
    from flask import request, session, render_template

    user_name = session.get('name')
    user_role = session.get('role')
    user_department = session.get('department')

    # Nếu là admin thì thấy tất cả khoa
    if user_role == 'admin':
        departments = db.session.query(User.department).distinct().all()
        departments = [d[0] for d in departments if d[0]]
    else:
        departments = [user_department] if user_department else []

    selected_department = request.args.get('department') or (departments[0] if departments else None)
    users = User.query.filter_by(department=selected_department).all() if selected_department else []
    shifts = Shift.query.all()

    app.logger.info(f"[AUTO_ASSIGN_VIEW] User '{user_name}' mở trang phân lịch nhanh cho khoa '{selected_department}'")

    return render_template('auto_assign.html',
                           departments=departments,
                           selected_department=selected_department,
                           users=users,
                           shifts=shifts)

def get_departments():
    return [d[0] for d in db.session.query(User.department).distinct().all() if d[0]]

@app.route('/auto-assign-submit', methods=['POST'])
def auto_assign_submit():
    from models.schedule import Schedule
    from models.shift import Shift
    from models.user import User

    # Lấy danh sách user theo thứ tự click
    ordered_user_ids = request.form.getlist('ordered_user_ids')

    # Lấy ngày bắt đầu, kết thúc, ca trực từ form
    start_date = request.form.get('start_date')  # dd/mm/yyyy
    end_date = request.form.get('end_date')      # dd/mm/yyyy
    selected_shifts = request.form.getlist('shift_ids')

    # TODO: Parse ngày & logic tạo lịch trực
    # Ví dụ: vòng qua các ngày trong khoảng start->end, gán lần lượt user theo ordered_user_ids

    # Sau khi tạo lịch, redirect lại trang lịch trực
    flash("Đã tạo lịch trực tự động theo thứ tự chọn.", "success")
    return redirect(url_for('view_schedule'))

@app.route('/auto-attendance', methods=['GET', 'POST'])
def auto_attendance_page():
    from models.user import User
    from models.shift import Shift
    from models.schedule import Schedule  # model lịch trực
    from models.attendance import Attendance
    from datetime import datetime, timedelta
    from flask import request, redirect, url_for, flash, render_template, session

    # Giới hạn danh sách khoa theo vai trò
    if session.get('role') == 'admin':
        departments = get_departments()
    else:
        user_department = session.get('department')
        departments = [user_department] if user_department else []

    if request.method == 'POST':
        selected_department = request.form.get('department')
    else:
        selected_department = request.args.get('department') or (departments[0] if departments else None)

    day_shifts = Shift.query.filter(Shift.name.ilike('%làm ngày%')).all()

    if selected_department:
        users = User.query.filter_by(department=selected_department).order_by(User.name).all()
    else:
        users = []

    if request.method == 'POST':
        selected_department = request.form.get('department')
        user_name = session.get('name')
        start_date_str = request.form.get('start_date')
        end_date_str = request.form.get('end_date')
        shift_code = request.form.get('shift_code')
        staff_ids = request.form.getlist('staff_ids')

        app.logger.info(f"[AUTO_ATTEND_START] User '{user_name}' bắt đầu tạo lịch trực {shift_code} cho khoa '{selected_department}' từ {start_date_str} đến {end_date_str} cho {len(staff_ids)} nhân viên.")

        if not (selected_department and start_date_str and end_date_str and shift_code and staff_ids):
            flash('Vui lòng chọn đầy đủ thông tin.', 'danger')
            return redirect(url_for('auto_attendance_page'))

        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()

        shift = Shift.query.filter_by(code=shift_code).first()
        if not shift:
            flash('Ca làm không hợp lệ.', 'danger')
            return redirect(url_for('auto_attendance_page'))

        staff_members = User.query.filter(User.id.in_(staff_ids)).all()

        current_date = start_date
        try:
            created_count = 0
            while current_date <= end_date:
                if current_date.weekday() in [5, 6] or current_date.strftime('%m-%d') in {'01-01', '04-30', '05-01', '09-02'}:
                    current_date += timedelta(days=1)
                    continue

                for staff in staff_members:
                    if not Schedule.query.filter_by(user_id=staff.id, work_date=current_date).first():
                        db.session.add(Schedule(user_id=staff.id, work_date=current_date, shift_id=shift.id))
                        created_count += 1
                current_date += timedelta(days=1)

            db.session.commit()
            app.logger.info(f"[AUTO_ATTEND_DONE] Đã tạo {created_count} dòng chấm công tự động thành công.")
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"[AUTO_ATTEND_ERROR] Lỗi khi tạo lịch trực: {e}")
            flash(f'Lỗi khi lưu lịch trực: {e}', 'danger')
            return redirect(url_for('auto_attendance_page', department=selected_department))

        flash(f'Đã tạo lịch trực cho {len(staff_members)} nhân viên từ {start_date} đến {end_date}.', 'success')
        return redirect(url_for('auto_attendance_page', department=selected_department))

    return render_template('auto_attendance.html',
                           departments=departments,
                           selected_department=selected_department,
                           users=users,
                           day_shifts=day_shifts)

@app.route('/sync-attendance', methods=['POST'])
def sync_attendance():
    from models.schedule import Schedule
    from models.attendance import Attendance
    from models.user import User
    from datetime import datetime
    from flask import request, flash, redirect, url_for
    from extensions import db

    start_date_str = request.form.get('start_date')
    end_date_str = request.form.get('end_date')
    department = request.form.get('department')

    if not (start_date_str and end_date_str and department):
        flash('Vui lòng cung cấp đủ thông tin: khoa, từ ngày, đến ngày.', 'danger')
        return redirect(url_for('auto_attendance_page'))

    start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
    end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()

    # Lấy tất cả lịch trực trong khoảng ngày và khoa được chọn
    schedules = Schedule.query.join(Schedule.user).filter(
        User.department == department,
        Schedule.work_date >= start_date,
        Schedule.work_date <= end_date
    ).all()

    # Xóa dữ liệu Attendance cũ trong khoảng thời gian và khoa để tránh trùng lặp
    Attendance.query.filter(
        Attendance.department == department,
        Attendance.date >= start_date,
        Attendance.date <= end_date
    ).delete()

    # Thêm dữ liệu Attendance mới dựa trên Schedule
    for schedule in schedules:
        attendance = Attendance(
            user_id=schedule.user_id,
            date=schedule.work_date,
            shift_id=schedule.shift_id,
            department=department,
            status='Công'
        )
        db.session.add(attendance)

    db.session.commit()
    flash(f'Đã đồng bộ {len(schedules)} bản ghi lịch trực sang bảng chấm công.', 'success')
    return redirect(url_for('auto_attendance_page', department=department))

from flask import Flask, request, render_template, redirect, url_for, flash, session
from datetime import datetime, timedelta
from extensions import db
from models.user import User
from models.shift import Shift
from models.schedule import Schedule

from utils.unit_config import get_unit_config

@app.route('/schedule', methods=['GET'])  # GET là đủ vì không xử lý POST
def view_schedule():
    user_role = session.get('role')
    user_dept = session.get('department')

    # Quyền: admin & admin1 được chọn khoa tùy ý, còn lại cố định theo khoa của user
    is_super = user_role in ('admin', 'admin1')
    selected_department = request.args.get('department') if is_super else user_dept

    # Danh sách khoa hiển thị trong combobox
    if is_super:
        departments = [d[0] for d in db.session.query(User.department)
                       .filter(User.department.isnot(None)).distinct().all()]
    else:
        departments = [user_dept] if user_dept else []

    # Ngày bắt đầu/kết thúc (mặc định 7 ngày)
    start_date_str = request.args.get('start_date')
    end_date_str = request.args.get('end_date')
    if start_date_str and end_date_str:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
    else:
        start_date = datetime.today().date()
        end_date = start_date + timedelta(days=6)

    # Lấy lịch trực theo khoảng ngày & khoa
    query = (Schedule.query
             .join(User)
             .join(Shift)
             .filter(Schedule.work_date.between(start_date, end_date)))
    if selected_department:
        query = query.filter(User.department == selected_department)

    schedules = query.order_by(Schedule.work_date).all()
    date_range = [start_date + timedelta(days=i)
                  for i in range((end_date - start_date).days + 1)]

    # Chuẩn bị dữ liệu hiển thị
    schedule_data = {}
    for s in schedules:
        u = s.user
        if not u:
            continue
        contract_type = getattr(u, 'contract_type', None)
        if u.id not in schedule_data:
            schedule_data[u.id] = {
                'id': u.id,
                'name': u.name,
                'position': u.position,
                'department': u.department,
                'contract_type': contract_type,
                'shifts': {},
                'shifts_full': {}
            }
        schedule_data[u.id]['shifts_full'].setdefault(s.work_date, []).append({
            'shift_id': s.shift.id if s.shift else None,
            'shift_name': s.shift.name if s.shift else '',
            'machine_type': getattr(s, 'machine_type', None),
            'work_hours': getattr(s, 'work_hours', None)
        })

    # Dữ liệu cho phần in (lọc chỉ ca trực, loại nghỉ)
    print_data = {}
    for uid, data in schedule_data.items():
        filtered = {}
        for d, items in data['shifts_full'].items():
            ca_truc = [ca for ca in items
                       if 'trực' in (ca.get('shift_name') or '').lower()
                       and 'nghỉ' not in (ca.get('shift_name') or '').lower()]
            if ca_truc:
                filtered[d] = ca_truc
        if filtered:
            print_data[uid] = {**{k: data[k] for k in
                                  ('id','name','position','department','contract_type')},
                               'shifts_full': filtered}

    # Thứ tự sắp xếp
    priority_order = ['GĐ','PGĐ','TK','TP','PTK','PTP','PK','PP',
                      'BS','ĐDT','KTVT','KTV','ĐD','NV','HL','BV','LX']
    def get_priority(pos):
        pos = (pos or '').upper()
        return priority_order.index(pos) if pos in priority_order else 999
    def get_contract_priority(ct):
        return 0 if ct and 'biên' in ct.lower() else 1

    schedule_data = dict(sorted(schedule_data.items(),
                         key=lambda kv: (get_priority(kv[1]['position']),
                                         get_contract_priority(kv[1]['contract_type']),
                                         kv[1]['name'])))
    print_data = dict(sorted(print_data.items(),
                      key=lambda kv: (get_priority(kv[1]['position']),
                                      get_contract_priority(kv[1]['contract_type']),
                                      kv[1]['name'])))

    # Chữ ký & khóa (nếu có model; nếu không có thì bỏ qua nhẹ nhàng)
    signature = None
    is_signed = False
    signed_at = None
    try:
        signature = ScheduleSignature.query.filter_by(
            department=selected_department, from_date=start_date, to_date=end_date
        ).first()
        is_signed = bool(signature)
        signed_at = signature.signed_at if signature else None
    except Exception:
        pass

    locked = False
    try:
        lock = ScheduleLock.query.filter_by(
            department=selected_department, start_date=start_date, end_date=end_date
        ).first()
        locked = bool(lock)
    except Exception:
        pass

    unit_config = {
        'name': 'BỆNH VIỆN NHI TỈNH GIA LAI',
        'address': '123 Đường ABC, Gia Lai',
        'phone': '0269 123456'
    }

    # Ngày lễ & đánh dấu
    from datetime import date
    HOLIDAYS = [date(2025,1,1), date(2025,4,30), date(2025,5,1), date(2025,9,2)]
    highlight_days = {d: ('holiday' if d in HOLIDAYS else 'weekend')
                      for d in date_range if d.weekday() in (5,6) or d in HOLIDAYS}

    return render_template(
        'schedule.html',
        departments=departments,
        selected_department=selected_department,
        schedule_data=schedule_data,
        print_data=print_data,
        date_range=date_range,
        start_date=start_date,
        end_date=end_date,
        now=datetime.now(),
        is_signed=is_signed,
        signed_at=signed_at,
        locked=locked,
        user={'role': user_role, 'department': user_dept, 'name': session.get('name')},
        unit_config=unit_config,
        highlight_days=highlight_days
    )

@app.route('/schedule/edit/<int:user_id>', methods=['GET', 'POST'])
def edit_user_schedule(user_id):
    user = User.query.get_or_404(user_id)
    shifts = Shift.query.order_by(Shift.order).all()

    # Lấy khoảng ngày từ query hoặc form (request.values ăn cả GET & POST)
    start_str = request.values.get('start')
    end_str = request.values.get('end')
    if start_str and end_str:
        start = datetime.strptime(start_str, '%Y-%m-%d').date()
        end = datetime.strptime(end_str, '%Y-%m-%d').date()
    else:
        # fallback: tháng hiện tại
        today = date.today()
        start = date(today.year, today.month, 1)
        end = date(today.year, today.month, calendar.monthrange(today.year, today.month)[1])

    # Chỉ load lịch trong khoảng bạn đang xem (VD: T8)
    schedules = (Schedule.query
                 .filter_by(user_id=user_id)
                 .filter(Schedule.work_date >= start,
                         Schedule.work_date <= end)
                 .order_by(Schedule.work_date)
                 .all())

    # Không cho sửa nếu đã khoá
    for s in schedules:
        locked = ScheduleLock.query.filter_by(department=user.department)\
            .filter(ScheduleLock.start_date <= s.work_date,
                    ScheduleLock.end_date >= s.work_date).first()
        if locked:
            return "Không thể chỉnh sửa. Lịch trực đã được ký xác nhận và khóa.", 403

    if request.method == 'POST':
        changed = 0
        for key, value in request.form.items():
            if not key.startswith('shift_') or not value:
                continue
            try:
                sched_id = int(key.split('_', 1)[1])
                new_shift_id = int(value)
            except ValueError:
                continue

            s = Schedule.query.get(sched_id)
            if not s or s.user_id != user_id:
                continue

            if s.shift_id != new_shift_id:
                s.shift_id = new_shift_id
                changed += 1

        if changed:
            try:
                db.session.commit()
                flash(f"✅ Đã lưu {changed} thay đổi ca trực.", "success")
            except Exception as e:
                db.session.rollback()
                flash("❌ Lỗi khi lưu lịch trực.", "danger")

        # Quay lại đúng khoảng ngày đang xem
        return redirect(url_for('view_schedule',
                                department=user.department,
                                start_date=start.strftime('%Y-%m-%d'),
                                end_date=end.strftime('%Y-%m-%d')))

    # Truyền start/end xuống template để form giữ lại khi POST
    return render_template('edit_schedule.html',
                           user=user, shifts=shifts, schedules=schedules,
                           start=start.strftime('%Y-%m-%d'),
                           end=end.strftime('%Y-%m-%d'))

@app.route('/schedule/delete-one', methods=['POST'])
def schedule_delete_one():
    role = session.get('role')
    if role not in ('admin', 'admin1', 'manager'):
        return "Bạn không có quyền thực hiện thao tác này.", 403

    user_id = int(request.form['user_id'])
    shift_id = int(request.form['shift_id'])
    work_date = datetime.strptime(request.form['work_date'], '%Y-%m-%d').date()

    s = Schedule.query.filter_by(
        user_id=user_id, shift_id=shift_id, work_date=work_date
    ).first()
    if not s:
        return "Không tìm thấy ca trực để xóa.", 404

    db.session.delete(s)
    db.session.commit()
    return redirect(request.referrer or url_for('view_schedule'))


@app.route('/schedule/sign', methods=['POST'])
@login_required
def sign_schedule():
    department = request.form.get('department')
    from_date_str = request.form.get('from_date')
    to_date_str = request.form.get('to_date')

    user_name = session.get('name')
    app.logger.info(f"[SIGN] User '{user_name}' ký lịch trực cho khoa '{department}' từ {from_date_str} đến {to_date_str}")

    if not department or not from_date_str or not to_date_str:
        flash("Thiếu thông tin để ký xác nhận.", "danger")
        return redirect('/schedule')

    from_date = datetime.strptime(from_date_str, '%Y-%m-%d').date()
    to_date = datetime.strptime(to_date_str, '%Y-%m-%d').date()

    # Kiểm tra đã ký chưa
    existing = ScheduleSignature.query.filter_by(
        department=department,
        from_date=from_date,
        to_date=to_date
    ).first()

    if existing:
        flash("📌 Bảng lịch trực này đã được ký xác nhận trước đó.", "info")
    else:
        signed_by = session.get('user_id')
        new_sig = ScheduleSignature(
            department=department,
            from_date=from_date,
            to_date=to_date,
            signed_by=signed_by,
            signed_at=datetime.now()
        )
        db.session.add(new_sig)
        db.session.commit()
        flash("✅ Đã ký xác nhận và khóa bảng lịch trực.", "success")

    return redirect('/schedule?department={}&start_date={}&end_date={}'.format(
        department, from_date_str, to_date_str
    ))

@app.route('/schedule/unsign', methods=['POST'])
@admin_required
def unsign_schedule():
    department = request.form['department']
    from_date = request.form['from_date']
    to_date = request.form['to_date']

    user_name = session.get('name')
    app.logger.info(f"[UNSIGN] Admin '{user_name}' hủy ký lịch khoa '{department}' từ {from_date} đến {to_date}")

    schedules = ScheduleAssignment.query.filter(
        ScheduleAssignment.department == department,
        ScheduleAssignment.date >= from_date,
        ScheduleAssignment.date <= to_date
    ).all()

    for sched in schedules:
        sched.is_signed = False
        sched.signed_at = None
        sched.signed_by = None

    db.session.commit()
    flash("🗑️ Đã hủy ký xác nhận. Có thể chỉnh sửa lịch trực.", "warning")
    return redirect('/schedule')

@app.route('/schedule/unlock', methods=['POST'])
def unlock_signature():
    if session.get('role') != 'admin':
        return "Bạn không có quyền thực hiện hành động này.", 403

    department = request.form.get('department')
    from_date = datetime.strptime(request.form.get('from_date'), '%Y-%m-%d').date()
    to_date = datetime.strptime(request.form.get('to_date'), '%Y-%m-%d').date()

    sig = ScheduleSignature.query.filter_by(
        department=department,
        from_date=from_date,
        to_date=to_date
    ).first()

    if sig:
        db.session.delete(sig)
        db.session.commit()
        flash("🧹 Đã hủy xác nhận và mở khóa lịch trực.", "warning")
    else:
        flash("Không tìm thấy bản ký xác nhận để hủy.", "danger")

    return redirect(f'/schedule?department={department}&start_date={from_date}&end_date={to_date}')

@app.route('/calendar')
def fullcalendar():
    selected_department = request.args.get('department')
    if session.get('role') in ['manager', 'user']:
        selected_department = session.get('department')

    departments = [d[0] for d in db.session.query(User.department).filter(User.department != None).distinct().all()]
    query = Schedule.query.join(User)

    if selected_department:
        query = query.filter(User.department == selected_department)

    schedules = query.order_by(Schedule.work_date).all()
    return render_template('fullcalendar.html',
                           schedules=schedules,
                           departments=departments,
                           selected_department=selected_department)

@app.route('/stats')
def stats():
    from sqlalchemy import func

    user_role = session.get('role')
    user_dept = session.get('department')

    query_users = User.query
    query_schedules = Schedule.query.join(User)

    if user_role in ['manager', 'user']:
        query_users = query_users.filter(User.department == user_dept)
        query_schedules = query_schedules.filter(User.department == user_dept)


    total_users = query_users.count()
    total_shifts = Shift.query.count()
    total_schedules = query_schedules.count()

    schedules_per_user = db.session.query(User.name, func.count(Schedule.id))\
        .join(Schedule)\
        .filter(User.department == user_dept if user_role != 'admin' else True)\
        .group_by(User.id).all()

    return render_template('stats.html',
                           total_users=total_users,
                           total_shifts=total_shifts,
                           total_schedules=total_schedules,
                           schedules_per_user=schedules_per_user)


@app.route('/report-by-department')
def report_by_department():
    user_role = session.get('role')
    user_dept = session.get('department')

    query = Schedule.query.join(User).join(Shift)
    if user_role in ['manager', 'user']:
        query = query.filter(User.department == user_dept)

    schedules = query.all()
    report = {}

    for s in schedules:
        dept = s.user.department
        if dept not in report:
            report[dept] = []
        report[dept].append(s)

    return render_template('report_by_department.html', report=report)

from sqlalchemy import or_

from sqlalchemy import or_

@app.route('/users-by-department')
def users_by_department():
    user_role = session.get('role')
    user_dept = session.get('department')

    # Lấy tham số lọc
    selected_department = (request.args.get('department') or '').strip()
    q = (request.args.get('q') or '').strip()

    # Danh sách khoa cho dropdown
    if user_role in ('admin', 'admin1'):
        departments = [d[0] for d in db.session.query(User.department)
                       .filter(User.department != None).distinct().all()]
    else:
        departments = [user_dept]
        selected_department = user_dept  # nhân viên/manager chỉ xem khoa mình

    # Base query
    query = User.query
    if selected_department:
        query = query.filter(User.department == selected_department)
    elif user_role not in ('admin', 'admin1'):
        # selected_department rỗng nhưng không phải admin
        query = query.filter(User.department == user_dept)

    # Tìm kiếm theo tên / username / chức danh
    if q:
        query = query.filter(or_(
            User.name.ilike(f"%{q}%"),
            User.username.ilike(f"%{q}%"),
            User.position.ilike(f"%{q}%")
        ))

    # (giữ nguyên order_by DB để ổn định, sẽ sắp xếp ưu tiên ở Python)
    if user_role in ('admin', 'admin1') and not selected_department:
        query = query.order_by(User.department, User.name)
    else:
        query = query.order_by(User.name)

    users = query.all()

    # --- SẮP XẾP ƯU TIÊN CHỨC DANH ---
    priority_order = ['GĐ', 'PGĐ', 'TK', 'TP', 'PTP', 'PP', 'PTK', 'PK', 'BS', 'ĐDT', 'ĐD', 'KTV', 'NV','HL', 'BV', 'LX']

    def priority_index(position: str) -> int:
        pos = (position or '').upper().strip()
        for i, pre in enumerate(priority_order):
            if pos.startswith(pre):
                return i
        return len(priority_order)  # không khớp → cuối

    if user_role in ('admin', 'admin1') and not selected_department:
        # Xem tất cả khoa: Khoa → Ưu tiên → Tên
        users = sorted(users, key=lambda u: (
            (u.department or '').lower().strip(),
            priority_index(u.position),
            (u.name or '').lower().strip()
        ))
    else:
        # Một khoa (hoặc không phải admin): Ưu tiên → Tên
        users = sorted(users, key=lambda u: (
            priority_index(u.position),
            (u.name or '').lower().strip()
        ))

    return render_template(
        'users_by_department.html',
        users=users,
        departments=departments,
        selected_department=selected_department,
        q=q,
        current_user_role=user_role
    )


@app.route('/users/inactive')
def inactive_users():
    users = User.query.filter_by(active=False).all()
    return render_template('inactive_users.html', users=users)

@app.route('/users/reactivate/<int:user_id>')
def reactivate_user(user_id):
    user = User.query.get_or_404(user_id)
    user.active = True
    db.session.commit()
    flash("✅ Nhân viên đã được khôi phục!", "success")
    return redirect('/users/inactive')

@app.route('/users/deactivate/<int:user_id>')
def deactivate_user(user_id):
    user = User.query.get_or_404(user_id)
    user.active = False
    db.session.commit()
    flash("🚫 Nhân viên đã được ngừng hoạt động!", "warning")
    return redirect(request.referrer or url_for('users_by_department'))

@app.route('/users/delete-all', methods=['GET', 'POST'])
def delete_all_users():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash("Không có quyền thực hiện chức năng này.", "danger")
        return redirect('/login')

    try:
        from models.user import User

        user_name = session.get('username', 'Unknown')
        user_role = session.get('role', 'unknown')
        selected_department = session.get('department', 'unknown')

        users_to_delete = User.query.filter(User.username != 'admin').all()
        print(f"🧹 Xoá {len(users_to_delete)} nhân sự...")
        app.logger.info(f"[USER_VIEW] User '{user_name}' ({user_role}) xem danh sách nhân sự khoa '{selected_department}'")

        for u in users_to_delete:
            print(f"→ Xoá: {u.username}")
            db.session.delete(u)

        db.session.commit()
        print("✅ Xoá xong.")
        flash("Đã xoá toàn bộ nhân sự (trừ admin).", "success")
    except Exception as e:
        db.session.rollback()
        print(f"❌ Lỗi khi xoá: {str(e)}")
        flash(f"Lỗi khi xoá: {str(e)}", "danger")

    return redirect('/users-by-department')

@app.route('/export-by-department', methods=['GET', 'POST'])
def export_by_department():
    from sqlalchemy import distinct

    user_role = session.get('role')
    user_dept = session.get('department')

    # Lấy danh sách khoa
    departments = [d[0] for d in db.session.query(distinct(User.department)).filter(User.department != None).all()]
    selected_department = request.form.get('department') if request.method == 'POST' else user_dept

    if user_role != 'admin':
        selected_department = user_dept

    # Tạo file Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(['Họ tên', 'Ca trực', 'Ngày trực'])

    # Truy vấn lịch có chứa từ "trực"
    query = Schedule.query.join(User).join(Shift).filter(Shift.name.ilike('%trực%'))

    if selected_department:
        query = query.filter(User.department == selected_department)

    schedules = query.order_by(Schedule.work_date).all()

    for s in schedules:
        if s.user and s.shift:
            ws.append([s.user.name, s.shift.name, s.work_date.strftime('%Y-%m-%d')])

    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return send_file(stream, as_attachment=True, download_name="lichtruc_theo_khoa.xlsx")

@app.route('/generate_schedule', methods=['GET', 'POST'])
def generate_schedule_route():
    from models.leave_request import LeaveRequest

    # Lấy danh sách khoa
    departments = [d[0] for d in db.session.query(User.department)
                   .filter(User.department.isnot(None)).distinct().all()]

    if request.method == 'POST':
        department = request.form.get('department')
        start_date = datetime.strptime(request.form['start_date'], '%Y-%m-%d').date()
        end_date = datetime.strptime(request.form['end_date'], '%Y-%m-%d').date()

        # Lấy danh sách nhân viên theo thứ tự click
        ordered_user_ids = request.form.getlist('ordered_user_ids')  # từ JS
        if not ordered_user_ids:
            # fallback nếu frontend chưa cập nhật
            ordered_user_ids = request.form.getlist('user_ids')

        shift_ids = request.form.getlist('shift_ids')
        people_per_day = int(request.form.get('people_per_day', 1))

        if not ordered_user_ids or not shift_ids:
            flash("⚠️ Vui lòng chọn ít nhất 1 nhân viên và 1 ca trực.", "danger")
            return redirect(request.referrer)

        ordered_user_ids = [int(uid) for uid in ordered_user_ids]
        shift_ids = [int(sid) for sid in shift_ids]
        user_count = len(ordered_user_ids)

        # Tạo danh sách ngày
        date_range = [start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)]

        assignments = []
        conflicts = []

        # Lấy đơn nghỉ phép trong khoảng
        leaves = LeaveRequest.query.filter(
            LeaveRequest.start_date <= end_date,
            LeaveRequest.end_date >= start_date
        ).all()

        user_index = 0
        for work_date in date_range:
            assigned = set()
            for shift_id in shift_ids:
                attempts = 0
                while len(assigned) < people_per_day and attempts < user_count * 2:
                    uid = ordered_user_ids[user_index % user_count]
                    user_index += 1
                    attempts += 1

                    # Kiểm tra nghỉ phép
                    if any(l.user_id == uid and l.start_date <= work_date <= l.end_date for l in leaves):
                        continue

                    # Kiểm tra trùng lịch
                    exists = Schedule.query.filter_by(user_id=uid, shift_id=shift_id, work_date=work_date).first()
                    if exists or uid in assigned:
                        continue

                    # Thêm lịch
                    assignments.append(Schedule(user_id=uid, shift_id=shift_id, work_date=work_date))
                    assigned.add(uid)

        # Ghi DB
        if assignments:
            db.session.add_all(assignments)
            db.session.commit()
            flash("✅ Đã tạo lịch trực tự động theo thứ tự chọn.", "success")
        else:
            flash("⚠️ Không có lịch nào được tạo. Có thể tất cả bị trùng hoặc nghỉ phép.", "warning")

        return redirect(url_for('generate_schedule_route'))

    # GET hiển thị form
    selected_department = request.args.get('department')
    users = User.query.filter_by(department=selected_department).all() if selected_department else []
    shifts = Shift.query.all()

    return render_template('generate_form.html',
                           departments=departments,
                           selected_department=selected_department,
                           users=users,
                           shifts=shifts)


@app.route('/export')
def export_excel():
    user_role = session.get('role')
    user_dept = session.get('department')
    wb.active = wb.active  # Đảm bảo chọn đúng sheet
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(['Họ tên', 'Ca trực', 'Ngày trực'])
    ws.freeze_panes = "A2"  # ✅ Cố định hàng tiêu đề

    # Lấy lịch trực có chứa từ "trực"
    query = Schedule.query.join(User).join(Shift).filter(Shift.name.ilike('%trực%'))
    if user_role != 'admin':
        query = query.filter(User.department == user_dept)

    schedules = query.order_by(Schedule.work_date).all()
    for s in schedules:
        if s.user and s.shift:
            ws.append([s.user.name, s.shift.name, s.work_date.strftime('%Y-%m-%d')])

    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return send_file(stream, as_attachment=True, download_name="lichtruc.xlsx")

from sqlalchemy import asc, desc, func

def normalize_shift_order():
    ORDER_COL = Shift.__table__.c.order
    # Lấy tất cả theo id để đánh lại chỉ mục 1..n
    all_ids = [row[0] for row in db.session.query(Shift.id).order_by(Shift.id.asc()).all()]
    for i, sid in enumerate(all_ids, start=1):
        db.session.query(Shift).filter(Shift.id == sid).update({ORDER_COL: i})
    db.session.commit()

import os
from urllib.parse import urlparse
from sqlalchemy import asc, desc, func, and_, or_

def normalize_shift_order():
    """Đưa thứ tự order về 1..n, loại NULL/trùng; dùng (order,id) để ổn định."""
    ORDER_COL = Shift.__table__.c.order
    rows = (db.session.query(Shift)
            .order_by(asc(func.coalesce(ORDER_COL, 10**9)), Shift.id.asc())
            .with_for_update()  # khoá tạm thời để tránh race
            .all())
    for i, r in enumerate(rows, start=1):
        r.order = i
    db.session.commit()


# --- DB URI: hỗ trợ Render Postgres & local SQLite ---
uri = os.getenv('DATABASE_URL', 'sqlite:///database.db')

# Render thường trả "postgres://", cần đổi sang "postgresql+psycopg2://"
if uri.startswith('postgres://'):
    uri = uri.replace('postgres://', 'postgresql+psycopg2://', 1)

# Bắt buộc SSL khi là Postgres
if uri.startswith('postgresql') and 'sslmode=' not in uri:
    sep = '&' if '?' in uri else '?'
    uri = f'{uri}{sep}sslmode=require'

app.config['SQLALCHEMY_DATABASE_URI'] = uri
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Tránh kết nối stale / rớt SSL
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    "pool_pre_ping": True,
    "pool_recycle": 300,
    "pool_size": 5,
    "max_overflow": 5,
    "connect_args": {"sslmode": "require"} if uri.startswith('postgresql') else {}
}

@app.route('/shifts/reindex', methods=['POST'])
def reindex_shifts():
    normalize_shift_order()
    flash("Đã sắp xếp lại thứ tự ca trực.", "success")
    return redirect(url_for('list_shifts'))

@app.route('/shifts')
def list_shifts():
    ORDER_COL = Shift.__table__.c.order
    try:
        shifts = Shift.query.order_by(ORDER_COL.asc(), Shift.id.asc()).all()
    except Exception:
        shifts = Shift.query.order_by(Shift.id.asc()).all()
    return render_template('shifts.html', shifts=shifts)


from flask import render_template, request, redirect, flash
from datetime import datetime
from models import Shift  # đảm bảo đã import đúng

def parse_time_string(time_str):
    for fmt in ("%H:%M:%S", "%H:%M"):
        try:
            return datetime.strptime(time_str, fmt).time()
        except ValueError:
            continue
    raise ValueError("❌ Định dạng thời gian không hợp lệ. Vui lòng nhập HH:MM hoặc HH:MM:SS.")

@app.route('/shifts/add', methods=['GET', 'POST'])
def add_shift():
    if request.method == 'POST':
        name = request.form['name']
        code = request.form['code']
        duration = request.form['duration']

        try:
            start_time = parse_time_string(request.form['start_time'])
            end_time = parse_time_string(request.form['end_time'])
            duration = float(duration)
        except ValueError as e:
            flash(str(e), 'danger')
            return render_template('add_shift.html', old=request.form)

        # Lấy order lớn nhất + 1
        max_order = db.session.query(db.func.max(Shift.order)).scalar() or 0
        new_order = max_order + 1

        shift = Shift(name=name, code=code, start_time=start_time,
                      end_time=end_time, duration=duration, order=new_order)

        db.session.add(shift)
        db.session.commit()

        user_name = session.get('name')
        app.logger.info(f"[SHIFT_ADD] {user_name} thêm ca trực '{name}' (Mã: {code}, {start_time}-{end_time})")

        flash("✅ Đã thêm ca trực mới.", "success")
        return redirect('/shifts')

    return render_template('add_shift.html')

from sqlalchemy import asc, desc

@app.route('/shifts/move_up/<int:shift_id>', methods=['GET', 'POST'])
def move_shift_up(shift_id):
    ORDER_COL = Shift.__table__.c.order
    normalize_shift_order()

    shift = (Shift.query.filter_by(id=shift_id)
             .with_for_update()
             .first_or_404())

    # Hàng ở ngay phía trên theo (order, id)
    above_shift = (Shift.query
                   .filter(or_(
                       ORDER_COL < shift.order,
                       and_(ORDER_COL == shift.order, Shift.id < shift.id)
                   ))
                   .order_by(ORDER_COL.desc(), Shift.id.desc())
                   .with_for_update()
                   .first())

    if above_shift:
        shift.order, above_shift.order = above_shift.order, shift.order
        db.session.commit()
    else:
        flash("Đã ở vị trí đầu tiên, không thể di chuyển lên.", "info")
    return redirect('/shifts')



@app.route('/shifts/move_down/<int:shift_id>', methods=['GET', 'POST'])
def move_shift_down(shift_id):
    ORDER_COL = Shift.__table__.c.order
    normalize_shift_order()

    shift = (Shift.query.filter_by(id=shift_id)
             .with_for_update()
             .first_or_404())

    # Hàng ở ngay phía dưới theo (order, id)
    below_shift = (Shift.query
                   .filter(or_(
                       ORDER_COL > shift.order,
                       and_(ORDER_COL == shift.order, Shift.id > shift.id)
                   ))
                   .order_by(ORDER_COL.asc(), Shift.id.asc())
                   .with_for_update()
                   .first())

    if below_shift:
        shift.order, below_shift.order = below_shift.order, shift.order
        db.session.commit()
    else:
        flash("Đã ở vị trí cuối cùng, không thể di chuyển xuống.", "info")
    return redirect('/shifts')



from flask import request, redirect, flash
import openpyxl
from datetime import datetime
from models.shift import Shift
from models import db

def parse_time_string(time_str):
    return datetime.strptime(time_str.strip(), '%H:%M').time()

from flask import flash, redirect, render_template

@app.route('/shifts/edit/<int:id>', methods=['GET', 'POST'])
def edit_shift(id):
    shift = Shift.query.get_or_404(id)
    old = {}

    if request.method == 'POST':
        old = {
            'code': request.form['code'],
            'name': request.form['name'],
            'start_time': request.form['start_time'],
            'end_time': request.form['end_time'],
            'duration': request.form['duration']
        }

        try:
            shift.code = old['code']
            shift.name = old['name']

            # Parse giờ HH:MM hoặc HH:MM:SS
            try:
                shift.start_time = datetime.strptime(old['start_time'], '%H:%M').time()
            except ValueError:
                shift.start_time = datetime.strptime(old['start_time'], '%H:%M:%S').time()

            try:
                shift.end_time = datetime.strptime(old['end_time'], '%H:%M').time()
            except ValueError:
                shift.end_time = datetime.strptime(old['end_time'], '%H:%M:%S').time()

            shift.duration = float(old['duration'])

            db.session.commit()

            user_name = session.get('name')
            app.logger.info(f"[SHIFT_EDIT] {user_name} sửa ca trực ID={id} thành tên: {shift.name}, mã: {shift.code}, giờ: {shift.start_time}-{shift.end_time}")

            return redirect('/shifts')

        except ValueError as ve:
            flash("⚠️ Vui lòng nhập giờ theo định dạng HH:MM hoặc HH:MM:SS", "danger")

    return render_template('edit_shift.html', shift=shift, old=old)

@app.route('/shifts/delete/<int:shift_id>')
def delete_shift(shift_id):
    shift = Shift.query.get_or_404(shift_id)

    from models.schedule import Schedule
    if Schedule.query.filter_by(shift_id=shift_id).first():
        return "Không thể xoá ca này vì đang được sử dụng trong lịch trực.", 400

    db.session.delete(shift)
    db.session.commit()

    user_name = session.get('name', 'Không xác định')
    app.logger.info(f"[SHIFT_DELETE] {user_name} đã xoá ca trực: {shift.name}")

    return redirect('/shifts')

@app.route('/export-shifts')
def export_shifts():
    import openpyxl
    from io import BytesIO
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Shifts"

    ws.append(["Tên ca", "Mã ca", "Giờ bắt đầu", "Giờ kết thúc", "Thời lượng"])

    for shift in Shift.query.order_by(Shift.name).all():
        ws.append([shift.name, shift.code, str(shift.start_time), str(shift.end_time), shift.duration])

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="danh_sach_ca.xlsx")

from datetime import datetime, time
from flask import flash  # cần import để sử dụng thông báo

@app.route('/import-shifts', methods=['POST'], endpoint='import_shifts')
def import_shifts_excel():  # ✅ Đổi tên hàm
    import openpyxl
    from datetime import datetime, time

    file = request.files['file']
    if not file:
        flash("Không có file được chọn.", "error")
        return redirect('/shifts')

    try:
        wb = openpyxl.load_workbook(file)
        ws = wb.active

        for idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            name, code, start, end, duration = row
            if name and code:
                def to_time(val):
                    if isinstance(val, time):
                        return val
                    if isinstance(val, datetime):
                        return val.time()
                    if isinstance(val, str):
                        for fmt in ("%H:%M", "%H:%M:%S"):
                            try:
                                return datetime.strptime(val.strip(), fmt).time()
                            except ValueError:
                                continue
                    return None

                start_time = to_time(start)
                end_time = to_time(end)

                if not start_time or not end_time:
                    flash(f"Dòng {idx}: Giờ bắt đầu/kết thúc sai định dạng. Dùng 'HH:MM' hoặc 'HH:MM:SS'.", "error")
                    continue

                existing = Shift.query.filter_by(code=code).first()
                if not existing:
                    new_shift = Shift(name=name, code=code, start_time=start_time, end_time=end_time, duration=duration)
                    db.session.add(new_shift)

        db.session.commit()
        flash("✅ Đã nhập ca trực thành công.", "success")
    except Exception as e:
        flash(f"❌ Lỗi khi đọc file: {str(e)}", "error")

    return redirect('/shifts')

@app.route('/users/edit/<int:user_id>', methods=['GET', 'POST'])
def edit_user(user_id):
    user = User.query.get_or_404(user_id)
    if request.method == 'POST':
        new_username = request.form['username']
        if User.query.filter(User.username == new_username, User.id != user_id).first():
            error_message = "User này đã có người dùng, bạn không thể cập nhật."
            return render_template('edit_user.html', user=user, error=error_message)

        user.name = request.form['name']
        user.username = new_username
        user.password = request.form['password']
        user.role = request.form['role']
        user.department = request.form['department']
        user.position = request.form['position']
        user.contract_type = request.form.get('contract_type')  # ✅ nếu có thêm trường này
        user.phone = request.form['phone']
        db.session.commit()
        return redirect('/users-by-department')

    return render_template('edit_user.html', user=user)


@app.route('/users/add', methods=['GET', 'POST'])
def add_user():
    current_role = session.get('role')

    if request.method == 'POST':
        name = request.form['name']
        username = request.form['username']
        password = request.form['password']
        role = request.form['role']
        department = request.form['department']
        position = request.form['position']
        contract_type = request.form.get('contract_type')
        phone = request.form.get('phone')

        # Nếu người tạo là manager -> ép role thành user
        if current_role == 'manager':
            role = 'user'

        # Kiểm tra username trùng
        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            flash("❌ Tên đăng nhập đã tồn tại. Vui lòng chọn tên khác.", "danger")
            return render_template('add_user.html', old=request.form)

        # Tạo user mới
        new_user = User(
            name=name,
            username=username,
            password=password,
            role=role,
            department=department,
            position=position,
            contract_type=contract_type,
            phone=phone,
        )
        db.session.add(new_user)
        db.session.commit()
        flash("✅ Đã thêm người dùng mới.", "success")
        return redirect('/users-by-department')

    # === DỮ LIỆU ĐỘNG ===
    # Lấy danh sách khoa từ DB (distinct department)
    departments = [d[0] for d in db.session.query(User.department)
                   .filter(User.department != None)
                   .distinct()
                   .all()]

    # Danh sách chức danh cố định (có thể chuyển sang DB nếu cần)
    positions = ['Bác sĩ', 'Điều dưỡng', 'Kỹ thuật viên', 'Nhân viên', 'Hộ lý', 'Bảo vệ', 'Lái xe']

    return render_template(
        'add_user.html',
        departments=departments,
        positions=positions
    )

@app.route('/import-users', methods=['GET', 'POST'])
def import_users():
    import openpyxl
    from sqlalchemy.exc import IntegrityError

    if request.method == 'POST':
        file = request.files['file']
        if file.filename.endswith('.xlsx'):
            wb = openpyxl.load_workbook(file)
            sheet = wb.active
            skipped_users = []
            imported_count = 0

            for idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
                if not any(row):
                    continue
                # Gán đúng thứ tự theo file:
                name          = row[0]
                username      = row[1]
                password      = row[2]
                role          = row[3]
                department    = row[4]
                position      = row[5]
                contract_type = row[6] if len(row) > 6 else None 
                phone         = row[7] if len(row) > 7 else None

                # Bỏ qua nếu thiếu tên đăng nhập hoặc đã tồn tại
                if not username or User.query.filter_by(username=username).first():
                    skipped_users.append(username or f"Hàng {idx}")
                    continue

                user = User(
                    name=name,
                    username=username,
                    password=password,
                    role=role,
                    department=department,
                    position=position,
                    contract_type=contract_type,
                    phone=phone
                )
                db.session.add(user)
                imported_count += 1

            try:
                db.session.commit()
            except IntegrityError:
                db.session.rollback()
                flash("❌ Lỗi khi lưu dữ liệu. Vui lòng kiểm tra nội dung file Excel.", "danger")
                return redirect('/users-by-department')

            if skipped_users:
                flash(f"⚠️ Đã nhập {imported_count} người dùng. Bỏ qua: {', '.join(skipped_users)}", "warning")
            else:
                flash(f"✅ Đã nhập thành công {imported_count} người dùng.", "success")

            return redirect('/users-by-department')
        else:
            flash("❌ Vui lòng chọn file Excel (.xlsx)", "danger")
            return redirect('/import-users')

    return render_template('import_users.html')

import logging
from datetime import datetime

# Thiết lập file log
logging.basicConfig(filename='phanquyen.log', level=logging.INFO)

@app.route('/roles', methods=['GET', 'POST'])
def manage_roles():
    # Kiểm tra quyền: admin và admin1 được truy cập
    if session.get('role') not in ['admin', 'admin1']:
        return "Bạn không có quyền truy cập trang này."

    search = request.args.get('search', '').strip()
    role_filter = request.args.get('role', '')
    department_filter = request.args.get('department', '').strip()
    page = int(request.args.get('page', 1))
    per_page = 10

    users_query = User.query

    # Admin1 không được xem admin
    if session.get('role') == 'admin1':
        users_query = users_query.filter(User.role != 'admin')

    # Lọc theo search, role, department
    if search:
        users_query = users_query.filter(User.name.ilike(f"%{search}%"))
    if role_filter:
        users_query = users_query.filter_by(role=role_filter)
    if department_filter:
        users_query = users_query.filter_by(department=department_filter)

    # Phân trang
    pagination = users_query.order_by(User.department).paginate(page=page, per_page=per_page, error_out=False)
    users = pagination.items

    modules = [
        'trang_chu', 'xem_lich_truc', 'yeu_cau_cv_ngoai_gio', 'don_nghi_phep',
        'xep_lich_truc', 'tong_hop_khth', 'cham_cong', 'bang_cham_cong_2', 'bang_cong_gop', 'bang_tinh_tien_truc',
        'cau_hinh_ca_truc', 'cau_hinh_muc_doc_hai','thiet_lap_phong_kham', 'nhan_su_theo_khoa',
        'cau_hinh_tien_truc', 'thiet_lap_khoa_hscc', 'cau_hinh_don_vi', 'phan_quyen',
        'danh_sach_cong_viec', 'xem_log', 'doi_mat_khau', 'module_config'
    ]

    module_names = {
        'trang_chu': 'Trang chủ',
        'xem_lich_truc': 'Xem lịch trực',
        'yeu_cau_cv_ngoai_gio': 'Yêu cầu công việc ngoài giờ',
        'don_nghi_phep': 'Đơn nghỉ phép',
        'xep_lich_truc': 'Xếp lịch trực',
        'tong_hop_khth': 'Tổng hợp KHTH',
        'cham_cong': 'Chấm công',
        'bang_cham_cong_cong_2': 'Bảng chấm công BHYT',
        'bang_cong_gop': 'Bảng công gộp',
        'bang_tinh_tien_truc': 'Bảng thanh toán tiền trực',
        'cau_hinh_ca_truc': 'Cấu hình ca trực',
        'cau_hinh_muc_doc_hai': 'Cấu hình mức độc hại',
        'thiet_lap_phong_kham': 'Thiết lập Phòng khám',
        'nhan_su_theo_khoa': 'Nhân sự theo khoa',
        'cau_hinh_tien_truc': 'Cấu hình tiền trực',
        'thiet_lap_khoa_hscc': 'Thiết lập khoa HSCC',
        'cau_hinh_don_vi': 'Cấu hình đơn vị',
        'phan_quyen': 'Phân quyền',
        'danh_sach_cong_viec': 'Danh sách yêu cầu công việc',
        'xem_log': 'Xem log hệ thống',
        'doi_mat_khau': 'Đổi mật khẩu',
        'module_config': 'Cấu hình phân hệ'
    }

    if request.method == 'POST':
        all_users = User.query.all()
        for user in all_users:
            selected_modules = request.form.getlist(f'modules_{user.id}[]')
            if not selected_modules:
                continue

            role = request.form.get(f'role_{user.id}')
            dept = request.form.get(f'department_{user.id}')
            position = request.form.get(f'position_{user.id}')

            if role and dept and position:
                if (user.role != role) or (user.department != dept) or (user.position != position):
                    logging.info(f"{datetime.now()} | Admin ID {session['user_id']} cập nhật: {user.username} → "
                                 f"Role: {user.role} → {role}, Dept: {user.department} → {dept}, Position: {user.position} → {position}")
                user.role = role
                user.department = dept
                user.position = position

            Permission.query.filter_by(user_id=user.id).delete()
            for mod in modules:
                db.session.add(Permission(
                    user_id=user.id,
                    module_name=mod,
                    can_access=(mod in selected_modules)
                ))

        db.session.commit()
        flash("✅ Đã lưu thay đổi phân quyền người dùng.", "success")
        return redirect('/roles')

    departments = [d[0] for d in db.session.query(User.department).distinct().all() if d[0]]
    roles = ['admin', 'admin1', 'manager', 'user']  # Thêm admin1
    positions = [p[0] for p in db.session.query(User.position).filter(User.position != None).distinct().all()]

    for user in users:
        user.modules = [perm.module_name for perm in user.permissions if perm.can_access]

    return render_template('manage_roles.html',
                           users=users,
                           departments=departments,
                           roles=roles,
                           positions=positions,
                           modules=modules,
                           module_names=module_names,
                           pagination=pagination,
                           current_search=search,
                           current_role=role_filter,
                           current_department=department_filter)

@app.route('/unit-config', methods=['GET', 'POST'])
def unit_config():
    if session.get('role') != 'admin':
        return "Bạn không có quyền truy cập."

    config = UnitConfig.query.first()
    if not config:
        config = UnitConfig()

    if request.method == 'POST':
        config.name = request.form['name']
        config.address = request.form['address']
        config.phone = request.form['phone']
        db.session.add(config)
        db.session.commit()
        flash("✅ Đã cập nhật thông tin đơn vị", "success")
        return redirect('/unit-config')

    return render_template('unit_config.html', config=config)

from flask import send_file

import os
from flask import Flask, render_template, request, redirect, session, send_file

LOG_DIR = os.path.join(os.path.dirname(__file__), 'logs')
LOG_FILE = os.path.join(LOG_DIR, 'activity.log')

def ensure_log_file():
    os.makedirs(LOG_DIR, exist_ok=True)
    if not os.path.exists(LOG_FILE):
        # tạo file rỗng để tránh lỗi không tồn tại
        open(LOG_FILE, 'a', encoding='utf-8').close()     

@app.route('/view-log')
def view_log():
    if session.get('role') != 'admin':
        return "Không có quyền truy cập", 403

    log_path = os.path.join('logs', 'activity.log')
    if not os.path.exists(log_path):
        return render_template('view_log.html', log_lines=[])

    with open(log_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    return render_template('view_log.html', log_lines=lines)

@app.route('/download-log')
def download_log():
    if session.get('role') != 'admin':
        return "Không có quyền tải log", 403

    log_path = os.path.join('logs', 'activity.log')
    return send_file(log_path, as_attachment=True)

@app.route('/clear-log', methods=['POST'])
def clear_log():
    if session.get('role') != 'admin':
        return "Không có quyền xoá log", 403

    log_path = os.path.join('logs', 'activity.log')
    open(log_path, 'w').close()
    return redirect('/view-log')


@app.route('/delete_user/<int:user_id>', methods=['POST'])
def delete_user(user_id):
    user = User.query.get(user_id)
    if user:
        if user.username == 'admin':
            flash("❌ Không thể xoá tài khoản admin.", "danger")
        else:
            db.session.delete(user)
            db.session.commit()
            flash(f"✅ Đã xoá: {user.name}", "success")
    else:
        flash("❌ Không tìm thấy người dùng.", "danger")

    return redirect('/users-by-department')

import traceback

@app.errorhandler(Exception)
def handle_exception(e):
    # Nếu đang debug, show lỗi chi tiết ra trình duyệt
    return f"""
        <h2 style='color: red;'>❌ Internal Server Error</h2>
        <pre>{traceback.format_exc()}</pre>
        <hr>
        <p style='color: gray;'>Vui lòng báo lỗi này cho quản trị viên hệ thống.</p>
    """, 500

@app.route('/export-template', methods=['POST'])
def export_template():
    department = request.form.get('department')
    start_date = request.form.get('start_date')
    end_date = request.form.get('end_date')

    print(">>> [EXPORT] Khoa:", department)
    print(">>> [EXPORT] Từ ngày:", start_date)
    print(">>> [EXPORT] Đến ngày:", end_date)

    query = Schedule.query.join(User).join(Shift)

    if department:
        query = query.filter(User.department == department)

    if start_date and end_date:
        try:
            start = datetime.strptime(start_date, '%Y-%m-%d').date()
            end = datetime.strptime(end_date, '%Y-%m-%d').date()
            query = query.filter(Schedule.work_date.between(start, end))
        except ValueError:
            return "Ngày không hợp lệ.", 400
    else:
        return "Vui lòng chọn khoảng thời gian.", 400

    schedules = query.order_by(Schedule.work_date).all()
    if not schedules:
        return "Không có dữ liệu lịch trực.", 404

    # Tập hợp ngày
    date_range = sorted({s.work_date for s in schedules})

    # Pivot dữ liệu người dùng
    schedule_data = {}
    for s in schedules:
        u = s.user
        if u.id not in schedule_data:
            schedule_data[u.id] = {
                'name': u.name,
                'position': u.position,
                'department': u.department,
                'shifts': {}
            }
        schedule_data[u.id]['shifts'][s.work_date] = s.shift.name

    # --- Tạo Excel ---
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Lịch trực ngang"

    # --- Quốc hiệu, tiêu đề đầu trang ---
    ws.merge_cells('A1:G1')
    ws['A1'] = "BỆNH VIỆN NHI TỈNH GIA LAI"
    ws.merge_cells('H1:N1')
    ws['H1'] = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    ws.merge_cells('H2:N2')
    ws['H2'] = "Độc lập - Tự do - Hạnh phúc"
    ws.merge_cells('A4:N4')
    ws['A4'] = f"BẢNG LỊCH TRỰC KHOA {department.upper() if department else ''}"
    ws.merge_cells('A5:N5')
    ws['A5'] = f"Lịch trực tuần ngày {start.strftime('%d/%m/%Y')} đến ngày {end.strftime('%d/%m/%Y')}"

    # --- Dòng tiêu đề bảng bắt đầu từ dòng 7 ---
    start_row = 7
    header = ['Họ tên', 'Chức danh', 'Khoa'] + [d.strftime('%d/%m') for d in date_range]
    ws.append(header)

    # Dữ liệu từng người
    for u in schedule_data.values():
        row = [u['name'], u['position'], u['department']]
        for d in date_range:
            row.append(u['shifts'].get(d, ''))  # Nếu không có ca → để trống
        ws.append(row)

    # --- Chân trang ---
    last_row = ws.max_row + 2
    ws[f'A{last_row}'] = "Nơi nhận:"
    ws[f'A{last_row+1}'] = "- Ban Giám đốc"
    ws[f'A{last_row+2}'] = "- Các khoa/phòng"
    ws[f'A{last_row+3}'] = "- Đăng website"
    ws[f'A{last_row+4}'] = "- Lưu: VP, KH-CNTT"

    ws.merge_cells(start_row=last_row, start_column=5, end_row=last_row, end_column=7)
    ws.cell(row=last_row, column=5).value = "Người lập bảng"
    ws.merge_cells(start_row=last_row, start_column=10, end_row=last_row, end_column=12)
    ws.cell(row=last_row, column=10).value = "GIÁM ĐỐC"

    ws.cell(row=last_row+1, column=5).value = "(Ký, ghi rõ họ tên)"
    ws.cell(row=last_row+1, column=10).value = "(Ký, ghi rõ họ tên)"

    # --- Xuất file ---
    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)

    return send_file(stream, as_attachment=True, download_name="lichtruc_dangngang.xlsx")

@app.route('/bang-cham-cong')
def bang_cham_cong():
    from datetime import datetime, timedelta, date
    from sqlalchemy import or_

    user_role = session.get('role')
    user_dept = session.get('department')

    today = datetime.today()
    start_date = request.args.get('start', today.replace(day=1).strftime('%Y-%m-%d'))
    end_date = request.args.get('end', today.strftime('%Y-%m-%d'))

    raw_department = request.args.get('department', '').strip()

    # Cho phép admin và admin1 xem tất cả khoa
    if user_role in ['admin', 'admin1']:
        selected_department = raw_department if raw_department else 'Tất cả'
        if raw_department:
            query = User.query.filter(User.department == raw_department)
        else:
            query = User.query
    else:
        selected_department = user_dept
        query = User.query.filter(User.department == selected_department)

    selected_contract = request.args.get('contract_type', '')
    print_filter = request.args.get('print_filter') == 'yes'

    start = datetime.strptime(start_date, '%Y-%m-%d').date()
    end = datetime.strptime(end_date, '%Y-%m-%d').date()
    days_in_range = [start + timedelta(days=i) for i in range((end - start).days + 1)]

    if selected_contract:
        if selected_contract.lower() == "hợp đồng":
            query = query.filter(or_(
                User.contract_type.ilike("hợp đồng%"),
                User.contract_type.ilike("%hợp đồng"),
                User.contract_type.ilike("%hợp đồng%")
            ))
        else:
            query = query.filter(User.contract_type.ilike(selected_contract))

    priority_order = ['GĐ', 'PGĐ', 'TK', 'TP', 'PTK', 'PTP', 'BS', 'BSCK1', 'BSCK2', 'ĐDT', 'KTVT', 'KTV', 'ĐD', 'NV', 'HL', 'BV']

    def sort_by_position(user):
        position = (user.position or '').upper().strip()
        for i, p in enumerate(priority_order):
            if position.startswith(p):
                return i
        return len(priority_order)

    users = query.filter(User.role != 'admin').all()
    users = sorted(users, key=sort_by_position)

    schedules = Schedule.query.join(Shift).filter(
        Schedule.user_id.in_([u.id for u in users]),
        Schedule.work_date.between(start, end)
    ).all()

    schedule_map = {}
    summary = {user.id: {'kl': 0, 'tg': 0, '100': 0, 'bhxh': 0} for user in users}
    used_keys = set()

    for s in schedules:
        key = (s.user_id, s.work_date)
        if key in used_keys:
            continue
        used_keys.add(key)

        code = s.shift.code.upper() if s.shift and s.shift.code else 'X'
        schedule_map[key] = code

        # KL: Không lương
        if code == "KL":
            summary[s.user_id]['kl'] += 1

        # Công hưởng lương thời gian
        elif code.startswith("X") and not code.startswith("XĐ") and not code.startswith("XĐL") and code not in ["/X", "/NT"]:
            # Bao gồm X, X1, X2, X3...
            summary[s.user_id]['tg'] += 1

        # Nửa công hưởng lương TG + nửa công 100%
        elif code in ["/X", "/NT"]:
            summary[s.user_id]['tg'] += 0.5
            summary[s.user_id]['100'] += 0.5

        # Công 100% (nghỉ bù, phép, hội nghị...)
        elif code in ["NB", "P", "H", "CT", "L", "NT", "PC", "NBL", "PT", "NBS", "NBC"]:
            summary[s.user_id]['100'] += 1

        # BHXH
        elif code in ["Ô", "CÔ", "DS", "TS", "TN"]:
            summary[s.user_id]['bhxh'] += 1

        # Công hưởng lương TG cho XĐ và XĐL các loại
        elif code.startswith("XĐ") or code.startswith("XĐL"):
            summary[s.user_id]['tg'] += 1

    holidays = [
        date(2025, 1, 1),
        date(2025, 4, 30),
        date(2025, 5, 1),
        date(2025, 9, 2),
    ]

    if user_role in ['admin', 'admin1']:
        departments = [d[0] for d in db.session.query(User.department).filter(User.department != None).distinct().all()]
    else:
        departments = [user_dept]

    now = datetime.today()
    month = start.month
    year = start.year

    return render_template("bang_cham_cong.html",
                           users=users,
                           departments=departments,
                           days_in_month=days_in_range,
                           schedule_map=schedule_map,
                           summary=summary,
                           holidays=holidays,
                           selected_department=selected_department,
                           selected_contract=selected_contract,
                           start_date=start_date,
                           end_date=end_date,
                           month=month,
                           year=year,
                           now=now)

@app.template_filter('break_code')
def break_code(code):
    if code and len(code) > 2:
        return code[:2] + '\n' + code[2:]
    return code

from flask import request, send_file
from models import User, Schedule

@app.route('/export-cham-cong')
def export_cham_cong():
    from datetime import datetime, timedelta
    from io import BytesIO
    import openpyxl
    from flask import send_file, request
    from sqlalchemy import or_

    # Lấy và xử lý tham số
    start_str = request.args.get('start', '').strip()
    end_str = request.args.get('end', '').strip()
    raw_department = request.args.get('department', '').strip()
    selected_contract = request.args.get('contract_type', '').strip().lower()

    # Sửa lỗi department=None từ URL
    if raw_department.lower() == 'none':
        raw_department = ''

    # Sửa lỗi ValueError do khoảng trắng dư
    start_date = datetime.strptime(start_str, "%Y-%m-%d").date()
    end_date = datetime.strptime(end_str, "%Y-%m-%d").date()
    days = [start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)]

    # Lọc nhân viên theo điều kiện
    query = User.query
    if raw_department:
        query = query.filter(User.department == raw_department)
    if selected_contract:
        if selected_contract == "hợp đồng":
            query = query.filter(or_(
                User.contract_type.ilike("hợp đồng%"),
                User.contract_type.ilike("%hợp đồng"),
                User.contract_type.ilike("%hợp đồng%")
            ))
        elif selected_contract == "biên chế":
            query = query.filter(User.contract_type.ilike("biên chế"))
    users = query.order_by(User.name).all()

    user_ids = [u.id for u in users]
    schedules = Schedule.query.join(Shift).filter(
        Schedule.user_id.in_(user_ids),
        Schedule.work_date.between(start_date, end_date)
    ).all()
    schedule_map = {(s.user_id, s.work_date): s.shift.code for s in schedules if s.shift and s.shift.code}

    def count_summary(uid):
        vals = [schedule_map.get((uid, d), '') for d in days]
        return {
            'kl': vals.count('KL'),
            'tg': sum(1 for c in vals if c in ["X", "/X", "XĐ", "XĐ16", "XĐ24", "XĐ2", "XĐ3", "XĐL16", "XĐL24"] or c.startswith("XĐ")),
            '100': sum(1 for c in vals if c in ["NB", "P", "H", "CT", "L", "NT", "PC", "NBL", "PT","NBS","NBC"]),
            'bhxh': sum(1 for c in vals if c in ["Ô", "CÔ", "DS", "TS", "TN"])
        }

    # Tạo file Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Bảng chấm công"

    # Đầu trang
    ws.append(["BỆNH VIỆN NHI TỈNH GIA LAI"])
    ws.append(["PHÒNG TỔ CHỨC - HCQT"])
    ws.append(["CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"])
    ws.append(["Độc lập - Tự do - Hạnh phúc"])
    ws.append([""])
    ws.append([f"BẢNG CHẤM CÔNG THÁNG {start_date.month} NĂM {start_date.year}"])
    ws.append([f"Khoa/Phòng: {raw_department or 'TOÀN VIỆN'}"])
    ws.append([f"Loại hợp đồng: {selected_contract.upper() if selected_contract else 'TẤT CẢ'}"])
    ws.append([""])
    
    header1 = ['STT', 'Họ và tên', 'Chức danh'] + [d.strftime('%d') for d in days] + ['Số công không hưởng lương', 'Số công hưởng lương TG', 'Số công nghỉ việc 100%', 'Số công BHXH']
    header2 = ['', '', ''] + ['CN' if d.weekday() == 6 else f"T{d.weekday() + 1}" for d in days] + ['', '', '', '']
    ws.append(header1)
    ws.append(header2)

    for idx, u in enumerate(users, start=1):
        row = [idx, u.name, u.position]
        row += [schedule_map.get((u.id, d), '') for d in days]
        s = count_summary(u.id)
        row += [s['kl'], s['tg'], s['100'], s['bhxh']]
        ws.append(row)

    # Chân trang
    ws.append([])
    footer = ['Nơi nhận:', '', '', 'Người lập']
    if raw_department:
        footer.append("Trưởng khoa" if "khoa" in raw_department.lower() else "Trưởng phòng")
    else:
        footer.append("")
    footer += ["Phòng TC - HCQT", "Giám đốc"]
    ws.append(footer)

    # DEBUG LOG
    print("✔️ Tổng số nhân viên:", len(users))
    print("✔️ Tổng ca trực lấy được:", len(schedules))
    print("✔️ Số dòng schedule_map:", len(schedule_map))

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="bang_cham_cong.xlsx")

# ====================== TIMESHEET 2 - HELPERS & ROUTES ======================
from models import db
from models.user import User
from models.shift import Shift
from models.schedule import Schedule
from models.timesheet2 import Timesheet2, Timesheet2Entry
# ====================== TIMESHEET 2 - HELPERS & ROUTES ======================
import re
from io import BytesIO
from datetime import datetime, timedelta, date
import openpyxl
from flask import request, session, redirect, render_template, send_file

# ========= TIMESHEET2 – RECALC USING YOUR CODE MAP =========
from collections import defaultdict
# Ở đầu app.py (khu import)
try:
    from sqlalchemy.dialects.postgresql import insert as pg_insert
except Exception:
    pg_insert = None  # chạy được cả khi không dùng Postgres

# ----- (giữ NGUYÊN các hằng & hàm bạn đã cung cấp) -----
SHIFT_NAME_TO_CODES = {
    'làm ngày': ['X'],
    'trực 24h': ['XĐ24'],
    'trực 16h': ['XĐ16'],
    'trực đêm': ['XĐ'],
    'trực thường trú': ['XĐT'],
    'nghỉ trực': ['NT'],
    'nghỉ bù': ['NB'],
}
_CODE_RE = re.compile(r'^/?[A-ZĐƠƯ]{1,4}[0-9]{0,3}%?$')

def _normalize_code(s: str) -> str:
    s = (s or '').strip().upper()
    if not s:
        return ''
    if s.startswith('/'):
        body = s[1:]
        body = re.sub(r'^XD(\d*)$', r'XĐ\1', body)
        return '/' + body
    return re.sub(r'^XD(\d*)$', r'XĐ\1', s)

def _codes_from_shift_name(name: str):
    s = (name or '').strip().lower()
    s = re.sub(r'\s+', ' ', s)
    for k, v in SHIFT_NAME_TO_CODES.items():
        if k in s:
            return v[:]
    if 'trực' in s and '24' in s: return ['XĐ24']
    if 'trực' in s and '16' in s: return ['XĐ16']
    if 'đêm'  in s:               return ['XĐ']
    if 'thường trú' in s:         return ['XĐT']
    if 'làm ngày' in s or s == 'ngày': return ['X']
    if 'nghỉ trực' in s:          return ['NT']
    if 'nghỉ bù' in s:            return ['NB']
    return []

def ts2_codes_from_shift(shift):
    code_field = getattr(shift, 'code', None)
    if not code_field:
        return _codes_from_shift_name(getattr(shift, 'name', ''))

    raw = str(code_field).strip()
    raw = re.sub(r'(\\n|/n|\r\n|\r)', '\n', raw)
    parts = re.split(r'[\n,;]+', raw)

    codes = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        token = p.split()[0]
        c = _normalize_code(token)
        if _CODE_RE.match(c):
            codes.append(c)
        else:
            codes += _codes_from_shift_name(p)

    seen, uniq = set(), []
    for c in codes:
        if c not in seen:
            seen.add(c); uniq.append(c)
    return uniq

def _resolve_codes_in_cell(codes):
    seen, ordered = set(), []
    for c in codes:
        c = _normalize_code(c)
        if c and c not in seen:
            seen.add(c); ordered.append(c)
    if 'X' in ordered:
        ordered = [c for c in ordered if not c.startswith('XĐ')]
    return ordered

# ==== Sort nhân sự theo chức danh rồi theo tên (dùng chung) ====
PRIORITY_ORDER = [
    'GĐ', 'PGĐ', 'TK', 'TP', 'PTK', 'PTP',
    'BS', 'BSCK1', 'BSCK2', 'ĐDT', 'KTVT', 'KTV',
    'ĐD', 'NV', 'HL', 'BV'
]

def _sort_by_position(u):
    """
    Trả về tuple (độ ưu tiên chức danh, tên viết thường) để dùng cho key của sorted().
    """
    pos = (getattr(u, 'position', '') or '').upper().strip()
    for i, p in enumerate(PRIORITY_ORDER):
        if pos.startswith(p):
            return (i, (getattr(u, 'name', '') or '').lower())
    return (len(PRIORITY_ORDER), (getattr(u, 'name', '') or '').lower())


MAP_KL    = {'KL'}                                              # Không hưởng lương
MAP_100   = {'NB','P','H','CT','L','NT','PC','NBL','PT','NBS','NBC','P100'}  # 100%
MAP_BHXH  = {'Ô','CÔ','DS','TS','TN','ỐM','OM'}                # BHXH
HALF_PAIR = {'/X','/NT'}                                        # 0.5 TG + 0.5 100%

def _is_tg(c: str) -> bool:
    c = _normalize_code(c)
    return (c == 'X') or c.startswith('XĐ') or c.startswith('XĐL') or c in {'XD','XĐ'}

def _resolve_cell_to_counts(code_list):
    seen, codes = set(), []
    for raw in code_list:
        c = _normalize_code(raw)
        if c and c not in seen:
            seen.add(c); codes.append(c)

    if any(c in HALF_PAIR for c in codes):
        return {'kl': 0.0, 'tg': 0.5, '100': 0.5, 'bhxh': 0.0}
    if any(c in MAP_KL for c in codes):
        return {'kl': 1.0, 'tg': 0.0, '100': 0.0, 'bhxh': 0.0}
    if any(c in MAP_100 for c in codes):
        return {'kl': 0.0, 'tg': 0.0, '100': 1.0, 'bhxh': 0.0}
    if any(c in MAP_BHXH for c in codes):
        return {'kl': 0.0, 'tg': 0.0, '100': 0.0, 'bhxh': 1.0}
    if 'X' in codes or any(_is_tg(c) for c in codes):
        return {'kl': 0.0, 'tg': 1.0, '100': 0.0, 'bhxh': 0.0}
    return {'kl': 0.0, 'tg': 0.0, '100': 0.0, 'bhxh': 0.0}
# ----- /giữ NGUYÊN -----

# ----- Parse mã từ 1 ô (Timesheet2Entry) -----
def _codes_from_free_text(txt: str):
    if not txt:
        return []
    raw = re.sub(r'(\\n|/n|\r\n|\r)', '\n', str(txt))
    parts = re.split(r'[\n,;]+', raw)
    codes = []
    for p in parts:
        token = _normalize_code(p.split()[0] if p else '')
        if token and _CODE_RE.match(token):
            codes.append(token)
        else:
            # nếu user gõ "trực 24h", "làm ngày"… thì map qua
            codes += _codes_from_shift_name(p)
    return codes

def _codes_from_entry(e, shift_by_id=None):
    """
    Trả về list mã của 1 ô:
      - ưu tiên text trong ô (code/value/note…)
      - nếu có shift_id và có bảng Shift: bổ sung theo ts2_codes_from_shift
    """
    texts = []
    for fld in ('code', 'value', 'note', 'text'):
        if hasattr(e, fld):
            v = getattr(e, fld) or ''
            if v:
                texts.append(v)
    codes = []
    for t in texts:
        codes += _codes_from_free_text(t)

    if shift_by_id and getattr(e, 'shift_id', None):
        sh = shift_by_id.get(e.shift_id)
        if sh:
            codes += ts2_codes_from_shift(sh)

    return _resolve_codes_in_cell(codes)

# ----- TÍNH LẠI 4 CỘT CHO 1 BẢNG (sheet_id) -----
def t2_recalc_summary(sheet_id: int):
    """
    Trả: { user_id: {'kl':..., 'tg':..., '100':..., 'bhxh':...} }
    """
    # đổi import theo model thực tế của bạn
    from models.timesheet2 import Timesheet2Entry
    from models.shift import Shift

    # cache Shift (nếu có)
    try:
        shift_by_id = {s.id: s for s in Shift.query.all()}
    except Exception:
        shift_by_id = {}

    totals = defaultdict(lambda: {'kl': 0.0, 'tg': 0.0, '100': 0.0, 'bhxh': 0.0})

    entries = (Timesheet2Entry.query
               .filter(Timesheet2Entry.sheet_id == sheet_id)
               .all())

    for e in entries:
        codes = _codes_from_entry(e, shift_by_id=shift_by_id)
        c = _resolve_cell_to_counts(codes)
        u = getattr(e, 'user_id', None)
        if not u:
            continue
        totals[u]['kl']   += c['kl']
        totals[u]['tg']   += c['tg']
        totals[u]['100']  += c['100']
        totals[u]['bhxh'] += c['bhxh']

    return totals
# ========= /TIMESHEET2 – RECALC =========

# ========================= INDEX =========================
@app.route('/timesheet2', methods=['GET', 'POST'])
def ts2_index():
    # Phải đăng nhập
    if 'user_id' not in session:
        return redirect('/login')

    role    = (session.get('role') or '').strip()
    my_dept = (session.get('department') or '').strip()
    is_admin = role in ('admin', 'admin1')  # manager KHÔNG phải admin

    # --- Danh sách khoa cho dropdown ---
    if is_admin:
        departments = [d[0] for d in db.session.query(User.department)
                       .filter(User.department.isnot(None))
                       .distinct().order_by(User.department).all()]
    else:
        # Manager/role khác: chỉ khoa của chính mình
        departments = [my_dept] if my_dept else []

    if request.method == 'POST':
        name = (request.form.get('name') or 'Bảng chấm công BHYT').strip()

        # Admin được chọn khoa; Manager bị ép theo khoa của mình (bỏ qua giá trị gửi lên)
        if is_admin:
            department = (request.form.get('department') or my_dept).strip()
        else:
            department = my_dept

        if not department:
            return "Chưa cấu hình khoa/phòng cho tài khoản.", 400

        start_date = datetime.strptime(request.form['start_date'], '%Y-%m-%d').date()
        end_date   = datetime.strptime(request.form['end_date'],   '%Y-%m-%d').date()

        sheet = Timesheet2(
            name=name,
            department=department,
            start_date=start_date,
            end_date=end_date,
            created_by=session.get('user_id')  # nếu model có cột này
        )
        db.session.add(sheet)
        db.session.commit()
        return redirect(f'/timesheet2/{sheet.id}')

    # --- Lọc danh sách theo quyền ---
    q = Timesheet2.query.order_by(Timesheet2.created_at.desc())
    if not is_admin:
        q = q.filter(Timesheet2.department == my_dept)
    sheets = q.all()

    return render_template(
        'timesheet2/index.html',
        sheets=sheets,
        departments=departments,
        role=role,
        user_dept=my_dept
    )

# ========== SAO CHÉP TỪ SCHEDULE (gộp mã 1 ô = 1 dòng + UPSERT) ==========
@app.route('/timesheet2/create-from-schedule', methods=['POST'])
def ts2_create_from_schedule():
    role = session.get('role')
    my_dept = session.get('department')

    department = request.form['department']
    if role not in ('admin', 'admin1','manager') and department != my_dept:
        return "Bạn không có quyền tạo bảng cho khoa này.", 403

    start_date = datetime.strptime(request.form['start_date'], '%Y-%m-%d').date()
    end_date   = datetime.strptime(request.form['end_date'],   '%Y-%m-%d').date()
    name = (request.form.get('name') or 'Bảng chấm công BHYT (sao chép)').strip()

    sheet = Timesheet2(name=name, department=department,
                       start_date=start_date, end_date=end_date)
    db.session.add(sheet)
    db.session.flush()

    q = (db.session.query(Schedule, User, Shift)
         .join(User, Schedule.user_id == User.id)
         .join(Shift, Schedule.shift_id == Shift.id)
         .filter(User.department == department)
         .filter(Schedule.work_date >= start_date, Schedule.work_date <= end_date))

    from collections import defaultdict
    bucket = defaultdict(list)   # (uid, date) -> list codes

    for s, u, sh in q:
        if not (u and sh):
            continue
        codes = ts2_codes_from_shift(sh)
        if not codes:
            continue
        for c in codes:
            bucket[(u.id, s.work_date)].append(c)

    # --- GỘP MÃ THEO Ô & UPSERT 1 DÒNG/Ô ---
    rows = []  # list[dict]: 1 row cho mỗi (uid, work_date)
    for (uid, wdate), codes in bucket.items():
        ordered = _resolve_codes_in_cell(codes)  # Ưu tiên X, loại XĐ*
        code_text = '\n'.join(_normalize_code(c) for c in ordered)  # 1 ô = 1 dòng
        rows.append({
            "sheet_id": sheet.id,
            "user_id": uid,
            "work_date": wdate,
            "code": code_text,
            "deleted": False
        })

    if rows:
        backend = (db.engine.url.get_backend_name() or "").lower()
        if pg_insert and backend.startswith('postgres'):
            stmt = pg_insert(Timesheet2Entry).values(rows)
            stmt = stmt.on_conflict_do_update(
                index_elements=['sheet_id', 'user_id', 'work_date'],
                set_={'code': stmt.excluded.code, 'deleted': False}
            )
            db.session.execute(stmt)
        else:
            # Fallback cho non‑Postgres: xóa ô trong khoảng ngày rồi chèn lại 1 dòng/ô
            Timesheet2Entry.query.filter(
                Timesheet2Entry.sheet_id == sheet.id,
                Timesheet2Entry.work_date >= start_date,
                Timesheet2Entry.work_date <= end_date
            ).delete(synchronize_session=False)
            db.session.bulk_insert_mappings(Timesheet2Entry, rows)

    db.session.commit()
    return redirect(f'/timesheet2/{sheet.id}')

from flask import send_from_directory
import os
@app.route('/favicon.ico')
def favicon():
    return send_from_directory(os.path.join(app.root_path, 'static'),
                               'favicon.ico', mimetype='image/vnd.microsoft.icon')

# ========================= VIEW =========================
def _norm_contract_type(s: str) -> str:
    s = (s or '').strip().lower()
    if s in ('bc','biên chế','bien che','bienché','biênchế'):
        return 'bienche'
    if s.startswith('hđ') or 'hợp' in s or 'hop dong' in s or 'hđlđ' in s:
        return 'hopdong'
    return 'khac'

@app.route('/timesheet2/<int:sheet_id>')
def ts2_view(sheet_id):
    sheet = Timesheet2.query.get_or_404(sheet_id)
    if session.get('role') not in ('admin','admin1', 'manager') and sheet.department != session.get('department'):
        return "Bạn không có quyền truy cập bảng này.", 403

    # --- Lọc loại hợp đồng (giống bảng chính) ---
    contract_type = (request.args.get('contract_type') or 'all').strip().lower()
    label_map = {'all':'LOẠI HỢP ĐỒNG','bienche':'BIÊN CHẾ','hopdong':'HỢP ĐỒNG'}
    contract_header_text = label_map.get(contract_type, 'LOẠI HỢP ĐỒNG')

    all_users = User.query.filter_by(department=sheet.department)\
                          .filter(User.role != 'admin').all()
    if contract_type in ('bienche','hopdong'):
        users = [u for u in all_users if _norm_contract_type(getattr(u, 'contract_type', '')) == contract_type]
    else:
        users = list(all_users)
    users = sorted(users, key=_sort_by_position)

    # --- Dải ngày ---
    days = []
    cur = sheet.start_date
    while cur <= sheet.end_date:
        days.append(cur)
        cur += timedelta(days=1)

    # --- Ngày Lễ để bôi màu (01/01, 30/04, 01/05, 02/09) ---
    def _year_holidays(y: int):
        return {date(y,1,1), date(y,4,30), date(y,5,1), date(y,9,2)}
    holiday_set = set()
    y = sheet.start_date.year
    while y <= sheet.end_date.year:
        holiday_set |= _year_holidays(y)
        y += 1
    holiday_keys = {d.strftime('%Y-%m-%d') for d in days if d in holiday_set}

    # --- Gom entries & tính tổng 4 cột ---
    from collections import defaultdict
    cell = defaultdict(list)
    for e in Timesheet2Entry.query.filter_by(sheet_id=sheet.id, deleted=False).all():
        cell[(e.user_id, e.work_date)].append(e)

    summary_by_user = {}
    for u in users:
        totals = {'kl':0, 'tg':0, '100':0, 'bhxh':0}
        for d in days:
            items = cell.get((u.id, d), [])
            raw = []
            for it in items:
                txt = re.sub(r'(\\n|/n|\r\n|\r)', '\n', (it.code or '').strip())
                raw += [p.strip() for p in txt.split('\n') if p.strip()]
            cnt = _resolve_cell_to_counts(raw)
            totals['kl']   += cnt['kl']
            totals['tg']   += cnt['tg']
            totals['100']  += cnt['100']
            totals['bhxh'] += cnt['bhxh']
        summary_by_user[u.id] = (float(totals['kl']), float(totals['tg']),
                                 float(totals['100']), float(totals['bhxh']))

    contract_options = [('all','Tất cả'), ('bienche','Biên chế'), ('hopdong','Hợp đồng')]

    return render_template('timesheet2/view.html',
                           sheet=sheet, users=users, days=days,
                           cell=cell, summary_by_user=summary_by_user,
                           holiday_keys=holiday_keys,
                           contract_type=contract_type,
                           contract_options=contract_options,
                           contract_header_text=contract_header_text)

# --- IN BẢNG CHẤM CÔNG (A4 ngang) ---
@app.route('/timesheet2/<int:sheet_id>/print')
def ts2_print(sheet_id):
    sheet = Timesheet2.query.get_or_404(sheet_id)
    if session.get('role') not in ('admin','admin1', 'manager') and sheet.department != session.get('department'):
        return "Không có quyền.", 403

    # Dải ngày
    days = []
    cur = sheet.start_date
    while cur <= sheet.end_date:
        days.append(cur)
        cur += timedelta(days=1)

    # Nhận tham số lọc
    ct = (request.args.get('contract_type') or '').strip().lower()   # 'hopdong' | 'bienche' | ''
    ct_col = getattr(User, 'contract_type', None)  # tên cột của bạn, vd 'contract_type'

    # Lấy user theo khoa + lọc loại hợp đồng (nếu có)
    users_q = User.query.filter_by(department=sheet.department).filter(User.role != 'admin')
    if ct_col is not None and ct in ('hopdong', 'bienche'):
        # chuẩn hoá về lowercase để so khớp “hợp đồng/biên chế”
        if ct == 'hopdong':
            users_q = users_q.filter(db.func.lower(ct_col).contains('hợp'))
        else:  # bienche
            users_q = users_q.filter(db.func.lower(ct_col).contains('biên'))
    users = sorted(users_q.all(), key=_sort_by_position)

    # Gom entries theo (user_id, date)
    entries = Timesheet2Entry.query.filter_by(sheet_id=sheet.id, deleted=False).all()
    from collections import defaultdict
    cell = defaultdict(list)
    for e in entries:
        cell[(e.user_id, e.work_date)].append(e)

    # Tính tổng 4 cột
    summary_by_user = {}
    for u in users:
        totals = {'kl':0, 'tg':0, '100':0, 'bhxh':0}
        for d in days:
            items = cell.get((u.id, d), [])
            raw_codes = []
            for it in items:
                txt = re.sub(r'(\\n|/n|\r\n|\r)', '\n', (it.code or '').strip())
                raw_codes += [p.strip() for p in txt.split('\n') if p.strip()]
            cnt = _resolve_cell_to_counts(raw_codes)
            totals['kl']   += cnt['kl']
            totals['tg']   += cnt['tg']
            totals['100']  += cnt['100']
            totals['bhxh'] += cnt['bhxh']
        summary_by_user[u.id] = (float(totals['kl']), float(totals['tg']),
                                 float(totals['100']), float(totals['bhxh']))

    # Nhãn “BIÊN CHẾ/HỢP ĐỒNG” đặt trong ô phần tiêu đề bên trái của bảng
    contract_label = ('HỢP ĐỒNG' if ct == 'hopdong'
                      else 'BIÊN CHẾ' if ct == 'bienche'
                      else 'LOẠI HỢP ĐỒNG')

    # Chân ký: phòng/khoa
    manager_label = 'TRƯỞNG KHOA' if (sheet.department and 'khoa' in sheet.department.lower()) else 'TRƯỞNG PHÒNG'

    today = datetime.today().date()
    return render_template('timesheet2/print.html',
                           sheet=sheet, users=users, days=days,
                           cell=cell, summary_by_user=summary_by_user,
                           contract_label=contract_label, manager_label=manager_label,
                           today=today)

# ========== CẬP NHẬT 1 Ô (gộp nhiều mã; UPSERT 1 dòng/ô) ==========
@app.route('/timesheet2/<int:sheet_id>/cell', methods=['POST'])
def timesheet2_update_cell(sheet_id):
    sheet = Timesheet2.query.get_or_404(sheet_id)
    if session.get('role') not in ('admin', 'admin1', 'manager') and sheet.department != session.get('department'):
        return {"ok": False, "msg": "Không có quyền"}, 403

    # Parse body an toàn
    data = request.get_json(silent=True) or (request.form.to_dict() if request.form else {})
    try:
        user_id   = int(data.get('user_id'))
        work_date = data.get('work_date')
        code_raw  = (data.get('code') or '').strip()
    except Exception:
        return {"ok": False, "msg": "Thiếu hoặc sai tham số"}, 400
    if not user_id or not work_date:
        return {"ok": False, "msg": "Thiếu user_id hoặc work_date"}, 400

    work_date = datetime.strptime(work_date, '%Y-%m-%d').date()

    # Chuẩn hoá & gộp mã thành 1 chuỗi
    if code_raw:
        code_raw = re.sub(r'(\\n|/n|\r\n|\r)', '\n', code_raw)
        parts = [p.strip() for p in code_raw.split('\n') if p.strip()]
        resolved = _resolve_codes_in_cell(parts)
        code_text = '\n'.join(_normalize_code(p) for p in resolved)
    else:
        code_text = ''

    # UPSERT 1 dòng/ô (Postgres) hoặc replace (SQLite)
    backend = (db.engine.url.get_backend_name() or "").lower()
    if pg_insert and backend.startswith('postgres'):
        stmt = pg_insert(Timesheet2Entry).values({
            "sheet_id": sheet.id,
            "user_id": user_id,
            "work_date": work_date,
            "code": code_text,
            "deleted": False
        }).on_conflict_do_update(
            index_elements=['sheet_id', 'user_id', 'work_date'],
            set_={'code': code_text, 'deleted': False}
        )
        db.session.execute(stmt)
    else:
        Timesheet2Entry.query.filter_by(
            sheet_id=sheet.id, user_id=user_id, work_date=work_date
        ).delete(synchronize_session=False)
        if code_text:
            db.session.add(Timesheet2Entry(
                sheet_id=sheet.id, user_id=user_id, work_date=work_date,
                code=code_text, deleted=False
            ))

    db.session.commit()
    return {"ok": True}

# ========== Thêm / Sửa / Xoá 1 CHIP (gộp vào 1 dòng/ô) ==========
@app.route('/timesheet2/<int:sheet_id>/item/add', methods=['POST'])
def ts2_add_item(sheet_id):
    sheet = Timesheet2.query.get_or_404(sheet_id)
    if session.get('role') not in ('admin', 'admin1', 'manager') and sheet.department != session.get('department'):
        return {"ok": False, "msg": "Không có quyền"}, 403

    data = request.get_json(force=True)
    user_id   = int(data['user_id'])
    work_date = datetime.strptime(data['work_date'], '%Y-%m-%d').date()
    raw = (data.get('code') or '').strip()
    new_code = _normalize_code(raw)
    if not new_code:
        return {"ok": False, "msg": "Ký hiệu trống!"}, 400

    # Lấy dòng hiện tại của ô (1 dòng/ô)
    row = Timesheet2Entry.query.filter_by(
        sheet_id=sheet.id, user_id=user_id, work_date=work_date
    ).first()

    if not row:
        # tạo mới
        db.session.add(Timesheet2Entry(
            sheet_id=sheet.id, user_id=user_id, work_date=work_date,
            code=new_code, deleted=False
        ))
        db.session.commit()
        return {"ok": True, "code": new_code}

    # đã có -> ghép mã vào cuối, loại trùng / ưu tiên X
    existing = []
    if row.code and row.code.strip():
        txt = re.sub(r'(\\n|/n|\r\n|\r)', '\n', row.code.strip())
        existing = [p.strip() for p in txt.split('\n') if p.strip()]

    # hợp nhất
    parts = existing + [new_code]
    resolved = _resolve_codes_in_cell(parts)
    row.code = '\n'.join(_normalize_code(p) for p in resolved)
    row.deleted = False
    db.session.commit()
    return {"ok": True, "code": new_code}

@app.route('/timesheet2/<int:sheet_id>/item/<int:item_id>/update', methods=['POST'])
def ts2_update_item(sheet_id, item_id):
    sheet = Timesheet2.query.get_or_404(sheet_id)
    if session.get('role') not in ('admin', 'admin1', 'manager') and sheet.department != session.get('department'):
        return {"ok": False, "msg": "Không có quyền"}, 403

    e = Timesheet2Entry.query.get_or_404(item_id)
    if e.sheet_id != sheet.id:
        return {"ok": False, "msg": "Sai sheet"}, 400

    data = request.get_json(force=True)
    raw = (data.get('code') or '').strip()
    code = _normalize_code(raw)
    if not code:
        return {"ok": False, "msg": "Ký hiệu trống!"}, 400
    e.code = code
    db.session.commit()
    return {"ok": True, "id": e.id, "code": e.code}

@app.route('/timesheet2/<int:sheet_id>/item/<int:item_id>/delete', methods=['POST'])
def ts2_delete_item(sheet_id, item_id):
    sheet = Timesheet2.query.get_or_404(sheet_id)
    if session.get('role') not in ('admin', 'admin1', 'manager') and sheet.department != session.get('department'):
        return {"ok": False, "msg": "Không có quyền"}, 403

    e = Timesheet2Entry.query.get_or_404(item_id)
    if e.sheet_id != sheet.id:
        return {"ok": False, "msg": "Sai sheet"}, 400

    e.deleted = True
    db.session.commit()
    return {"ok": True}

# ========== XÓA TOÀN DÒNG CỦA 1 NHÂN SỰ ==========
@app.route('/timesheet2/<int:sheet_id>/row/delete', methods=['POST'])
def ts2_delete_row(sheet_id):
    sheet = Timesheet2.query.get_or_404(sheet_id)
    if session.get('role') not in ('admin', 'admin1', 'manager') and sheet.department != session.get('department'):
        return {"ok": False, "msg": "Không có quyền"}, 403

    data = request.get_json(force=True)
    user_id = int(data['user_id'])
    Timesheet2Entry.query.filter_by(sheet_id=sheet.id, user_id=user_id)\
                         .update({"deleted": True, "code": ''})
    db.session.commit()
    return {"ok": True}

# ========== XOÁ CẢ BẢNG (duy nhất) ==========
@app.route('/timesheet2/<int:sheet_id>/delete', methods=['POST'])
def ts2_delete_sheet(sheet_id):
    sheet = Timesheet2.query.get_or_404(sheet_id)
    if session.get('role') not in ('admin','admin1', 'manager') and sheet.department != session.get('department'):
        return {"ok": False, "msg": "Không có quyền"}, 403
    Timesheet2Entry.query.filter_by(sheet_id=sheet.id).delete()
    db.session.delete(sheet)
    db.session.commit()
    return {"ok": True}

# ========== DỌN/CHUẨN HOÁ SHEET (gộp nhiều dòng/ô thành 1) ==========
@app.route('/timesheet2/<int:sheet_id>/normalize', methods=['POST'])
def ts2_normalize_sheet(sheet_id):
    sheet = Timesheet2.query.get_or_404(sheet_id)
    if session.get('role') not in ('admin', 'admin1', 'manager') and sheet.department != session.get('department'):
        return {"ok": False, "msg": "Không có quyền"}, 403

    from collections import defaultdict
    cell = defaultdict(list)
    entries = Timesheet2Entry.query.filter_by(sheet_id=sheet.id, deleted=False).all()
    for e in entries:
        cell[(e.user_id, e.work_date)].append(e)

    changed = 0
    for key, items in cell.items():
        raw_codes = []
        for it in items:
            txt = re.sub(r'(\\n|/n|\r\n|\r)', '\n', (it.code or '').strip())
            raw_codes += [p.strip() for p in txt.split('\n') if p.strip()]
            it.deleted = True

        resolved = _resolve_codes_in_cell(raw_codes)
        db.session.add(Timesheet2Entry(
            sheet_id=sheet.id, user_id=key[0], work_date=key[1],
            code='\n'.join(_normalize_code(c) for c in resolved), deleted=False
        ))
        changed += 1

    db.session.commit()
    return {"ok": True, "changed": changed}

# ========== EXPORT EXCEL ==========
@app.route('/timesheet2/<int:sheet_id>/export')
def ts2_export(sheet_id):
    sheet = Timesheet2.query.get_or_404(sheet_id)
    if session.get('role') not in ('admin', 'admin1', 'manager') and sheet.department != session.get('department'):
        return "Không có quyền.", 403

    # Dải ngày
    days = []
    cur = sheet.start_date
    while cur <= sheet.end_date:
        days.append(cur)
        cur += timedelta(days=1)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Bang cham cong BHYT'
    header = ['STT', 'Họ và tên', 'Chức vụ'] + [d.strftime('%d') for d in days]
    ws.append(header)

    users_q = User.query.filter_by(department=sheet.department).filter(User.role != 'admin')
    users = sorted(users_q.all(), key=_sort_by_position)
    entries = Timesheet2Entry.query.filter_by(sheet_id=sheet.id, deleted=False).all()

    from collections import defaultdict
    m = defaultdict(list)
    for e in entries:
        if e.code and e.code.strip():
            txt = re.sub(r'(\\n|/n|\r\n|\r)', '\n', e.code.strip())
            for part in [p.strip() for p in txt.split('\n') if p.strip()]:
                m[(e.user_id, e.work_date)].append(_normalize_code(part))

    for idx, u in enumerate(users, start=1):
        row = [idx, u.name, (u.position or '')]
        for d in days:
            cell_codes = '\n'.join(m.get((u.id, d), []))
            row.append(cell_codes)
        ws.append(row)

    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return send_file(stream, as_attachment=True,
                     download_name=f"bang-cham-cong-bhyt-{sheet.id}.xlsx")
# ==================== HẾT PHẦN TIMESHEET 2 ====================

from flask import render_template, request, send_file
from datetime import datetime, timedelta
from models import User, Shift, Schedule
import openpyxl
from io import BytesIO

@app.route('/report-all')
def report_all():
    start_str = request.args.get('start')
    end_str = request.args.get('end')

    if start_str and end_str:
        start_date = datetime.strptime(start_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_str, '%Y-%m-%d').date()
    else:
        start_date = datetime.today().date()
        end_date = start_date + timedelta(days=6)

    schedules = Schedule.query.join(User).join(Shift).filter(
        Schedule.work_date.between(start_date, end_date),
        Schedule.shift_id != None,
        User.department != None
    ).order_by(User.department, Schedule.work_date).all()

    grouped = {}

    for s in schedules:
        shift_name = s.shift.name.strip().lower()

        # Chỉ lấy ca trực thật sự
        if not shift_name.startswith('trực'):
            continue

        dept = s.user.department.strip() if s.user.department else 'Khác'
        key = s.work_date.strftime('%a %d/%m')
        position = s.user.position.strip() if s.user.position else ''
        name = s.user.name.strip()

        # Gắn chức vụ vào tên
        display_name = name if name.startswith(position) else f"{position}. {name}" if position else name

        # Xác định loại ca
        if 'thường trú' in shift_name:
            ca_text = 'Trực thường trú'
        elif '24' in shift_name:
            ca_text = 'Trực 24h'
        elif '16' in shift_name:
            ca_text = 'Trực 16h'
        elif '8' in shift_name:
            ca_text = 'Trực 8h'
        else:
            ca_text = f"Trực {int(s.shift.duration)}h"

        line = f"{display_name} ({ca_text})"

        grouped.setdefault(dept, {})
        grouped[dept].setdefault(key, [])
        grouped[dept][key].append((position, line))

    # Thứ tự ưu tiên chức danh
    priority_order = ['GĐ', 'PGĐ', 'TK', 'TP', 'PP', 'PTK', 'PK', 'BS', 'ĐDT', 'ĐD', 'KTV', 'NV']

    def get_priority(pos):
        pos = pos.upper() if pos else ''
        return priority_order.index(pos) if pos in priority_order else 999

    # Sắp xếp nhân sự trong từng ngày theo ưu tiên
    grouped_by_dept = {
        dept: {
            day: sorted(entries, key=lambda x: get_priority(x[0]))
            for day, entries in dept_data.items()
        }
        for dept, dept_data in grouped.items() if any(dept_data.values())
    }

    # Sắp xếp khoa: Ban giám đốc trước
    def sort_priority(name):
        name = name.lower()
        if 'giám đốc' in name or 'ban giám' in name:
            return '1_' + name
        elif 'khoa' in name:
            return '2_' + name
        else:
            return '3_' + name

    dept_ordered = sorted(grouped_by_dept.keys(), key=sort_priority)
    date_range = [start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)]

    return render_template(
        'report_all.html',
        grouped_by_dept=grouped_by_dept,
        dept_ordered=dept_ordered,
        date_range=date_range,
        start_date=start_date,
        end_date=end_date
    )

@app.route('/export-report-all')
def export_report_all():
    start_str = request.args.get('start')
    end_str = request.args.get('end')
    if start_str and end_str:
        start_date = datetime.strptime(start_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_str, '%Y-%m-%d').date()
    else:
        start_date = datetime.today().date()
        end_date = start_date + timedelta(days=6)

    schedules = Schedule.query.join(User).join(Shift) \
        .filter(Schedule.work_date.between(start_date, end_date)) \
        .order_by(Schedule.work_date).all()

    grouped = {}
    for s in schedules:
        dept = s.user.department or 'Khác'
        pos = s.user.position or ''
        key = (dept, pos)
        grouped.setdefault(key, {})
        day = s.work_date.strftime('%a %d/%m')
        grouped[key][day] = grouped[key].get(day, '') + f"{s.user.name} ({s.shift.name}); "

    date_range = [(start_date + timedelta(days=i)).strftime('%a %d/%m')
                  for i in range((end_date - start_date).days + 1)]

    wb = openpyxl.Workbook()
    ws = wb.active
    header = ['Khoa/Phòng', 'Chức danh'] + date_range
    ws.append(header)

    for (dept, pos), days in grouped.items():
        row = [dept, pos] + [days.get(d, '') for d in date_range]
        ws.append(row)

    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return send_file(stream, as_attachment=True, download_name='lich_truc_toan_vien.xlsx')

from models.ca_config import CaConfiguration

from models.ca import Ca

@app.route('/cas', methods=['GET', 'POST'])
def manage_cas():
    if session.get('role') not in ['admin', 'manager']:
        return redirect('/')

    user_role = session.get('role')
    user_dept = session.get('department')

    if user_role == 'admin':
        departments = [d[0] for d in db.session.query(User.department).filter(User.department != None).distinct().all()]
        selected_department = request.args.get('department') or request.form.get('department') or (departments[0] if departments else None)
    else:
        departments = [user_dept]
        selected_department = user_dept

    if request.method == 'POST':
        new_ca = Ca(
            name=request.form.get('ca_name'),
            department=selected_department,
            doctor_id=request.form.get('doctor_id'),
            nurse1_id=request.form.get('nurse1_id'),
            nurse2_id=request.form.get('nurse2_id'),
        )
        db.session.add(new_ca)
        db.session.commit()
        return redirect(f"/cas?department={selected_department}")

    doctors = User.query.filter(User.position == 'BS', User.department == selected_department).all()
    nurses = User.query.filter(User.position == 'ĐD', User.department == selected_department).all()
    cas = Ca.query.filter_by(department=selected_department).all()

    return render_template('cas.html',
                           departments=departments,
                           selected_department=selected_department,
                           doctors=doctors,
                           nurses=nurses,
                           cas=cas)

@app.route('/cas/delete/<int:ca_id>', methods=['POST'])
def delete_ca(ca_id):
    if session.get('role') not in ['admin', 'manager']:
        return "Không có quyền."
    ca = Ca.query.get_or_404(ca_id)
    db.session.delete(ca)
    db.session.commit()
    return redirect(request.referrer or '/cas')


@app.route('/cas/edit/<int:ca_id>', methods=['GET', 'POST'])
def edit_ca(ca_id):
    if session.get('role') not in ['admin', 'manager']:
        return redirect('/')

    ca = Ca.query.get_or_404(ca_id)
    departments = [d[0] for d in db.session.query(User.department).filter(User.department != None).distinct().all()]
    doctors = User.query.filter_by(department=ca.department, position='BS').all()
    nurses = User.query.filter_by(department=ca.department, position='ĐD').all()

    if request.method == 'POST':
        ca.name = request.form.get('ca_name')
        ca.doctor_id = request.form.get('doctor_id')
        ca.nurse1_id = request.form.get('nurse1_id')
        ca.nurse2_id = request.form.get('nurse2_id')
        db.session.commit()
        return redirect(f'/cas?department={ca.department}')

    return render_template('edit_ca.html', ca=ca, doctors=doctors, nurses=nurses, departments=departments)

from models.ca_config import CaConfiguration
from models.ca import Ca

@app.route('/generate-ca', methods=['GET', 'POST'])
def generate_ca_schedule_route():
    from models import Ca, Schedule, Shift, CaConfiguration, User

    selected_department = request.args.get('department') or request.form.get('department')
    model_type = request.form.get('model_type')
    start_date_str = request.form.get('start_date')
    end_date_str = request.form.get('end_date')

    # Lấy danh sách khoa có trong hệ thống
    departments = [d[0] for d in db.session.query(User.department).filter(User.department != None).distinct().all()]
    selected_config = None
    if selected_department:
        selected_config = CaConfiguration.query.filter_by(department=selected_department).first()

    # Nếu là POST và đủ dữ liệu thì tiến hành tạo lịch
    if request.method == 'POST' and selected_department and start_date_str and end_date_str:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
        days = (end_date - start_date).days + 1

        cas = Ca.query.filter_by(department=selected_department).all()
        shifts = Shift.query.all()

        if len(cas) < 3:
            flash("Phải có ít nhất 3 ca để chạy lịch 2 ca 3 kíp", "danger")
            return redirect(request.url)

        # Tạo mô hình 2 ca 3 kíp
        def pattern_2ca3kip(i):
            cycle = [
                [("Ca 1", "Làm ngày"), ("Trực Ca 2", "Trực đêm"), ("Ca 3", "Nghỉ")],
                [("Ca 3", "Làm ngày"), ("Trực Ca 1", "Trực đêm"), ("Ca 2", "Nghỉ")],
                [("Ca 2", "Làm ngày"), ("Trực Ca 3", "Trực đêm"), ("Ca 1", "Nghỉ")]
            ]
            return cycle[i % 3]

        assignments = []
        for i in range(days):
            current_date = start_date + timedelta(days=i)
            day_shift_name, night_shift_name = pattern_2ca3kip(i)

            day_shift = next((s for s in shifts if s.name.lower() == day_shift_name[0].lower()), None)
            night_shift = next((s for s in shifts if s.name.lower() == night_shift_name[0].lower()), None)

            if not day_shift or not night_shift:
                continue

            ca_day = cas[i % len(cas)]
            ca_night = cas[(i + 1) % len(cas)]

            # Phân công ca ngày
            assignments.extend([
                Schedule(user_id=ca_day.doctor_id, shift_id=day_shift.id, work_date=current_date),
                Schedule(user_id=ca_day.nurse1_id, shift_id=day_shift.id, work_date=current_date),
                Schedule(user_id=ca_day.nurse2_id, shift_id=day_shift.id, work_date=current_date),
            ])

            # Phân công ca đêm
            assignments.extend([
                Schedule(user_id=ca_night.doctor_id, shift_id=night_shift.id, work_date=current_date),
                Schedule(user_id=ca_night.nurse1_id, shift_id=night_shift.id, work_date=current_date),
                Schedule(user_id=ca_night.nurse2_id, shift_id=night_shift.id, work_date=current_date),
            ])

        db.session.add_all(assignments)
        db.session.commit()

        flash("Tạo lịch trực 2 ca 3 kíp thành công", "success")
        return redirect(url_for('view_schedule', department=selected_department,
                                start_date=start_date_str, end_date=end_date_str))

    # GET hoặc POST thiếu thông tin -> hiển thị lại form
    return render_template(
        "generate_ca.html",
        department=selected_department,
        departments=departments,
        selected_config=selected_config,
        model_type=model_type,
        start_date=start_date_str,
        end_date=end_date_str
    )

@app.route('/configure-ca', methods=['GET', 'POST'])
def configure_ca():
    if session.get('role') not in ['admin', 'manager']:
        return "Bạn không có quyền truy cập."

    user_role = session.get('role')
    user_dept = session.get('department')

    if user_role == 'admin':
        departments = [d[0] for d in db.session.query(User.department).filter(User.department != None).distinct().all()]
    else:
        departments = [user_dept]

    selected_config = None
    message = None

    if request.method == 'POST':
        config_id = request.form.get('config_id')
        department = request.form['department']
        num_shifts = int(request.form['num_shifts'])
        cas_per_shift = int(request.form['cas_per_shift'])
        doctors_per_ca = int(request.form['doctors_per_ca'])
        nurses_per_ca = int(request.form['nurses_per_ca'])

        if config_id:
            config = CaConfiguration.query.get(int(config_id))
            if config:
                config.num_shifts = num_shifts
                config.cas_per_shift = cas_per_shift
                config.doctors_per_ca = doctors_per_ca
                config.nurses_per_ca = nurses_per_ca
                message = "Cập nhật cấu hình thành công."
        else:
            config = CaConfiguration(
                department=department,
                num_shifts=num_shifts,
                cas_per_shift=cas_per_shift,
                doctors_per_ca=doctors_per_ca,
                nurses_per_ca=nurses_per_ca
            )
            db.session.add(config)
            message = "Đã lưu cấu hình mới."

        db.session.commit()

    edit_id = request.args.get('edit_id')
    if edit_id:
        selected_config = CaConfiguration.query.get(int(edit_id))

    configs = CaConfiguration.query.all()
    return render_template(
        'configure_ca.html',
        departments=departments,
        configs=configs,
        selected_config=selected_config,
        message=message
    )

@app.route('/configure-ca/edit/<int:id>', methods=['GET', 'POST'])
def edit_ca_config(id):
    if session.get('role') not in ['admin', 'manager']:
        return redirect('/')
        
    config = CaConfiguration.query.get_or_404(id)
    departments = [d[0] for d in db.session.query(User.department).filter(User.department != None).distinct().all()]

    if request.method == 'POST':
        config.department = request.form['department']
        config.num_shifts = int(request.form['num_shifts'])
        config.cas_per_shift = int(request.form['cas_per_shift'])
        config.doctors_per_ca = int(request.form['doctors_per_ca'])
        config.nurses_per_ca = int(request.form['nurses_per_ca'])
        db.session.commit()
        return redirect('/configure-ca')

    return render_template('edit_ca_config.html', config=config, departments=departments)

@app.route('/delete-ca-config/<int:config_id>', methods=['POST'])
def delete_ca_config(config_id):
    if session.get('role') not in ['admin', 'manager']:
        return "Không có quyền."

    config = CaConfiguration.query.get(config_id)
    if config:
        db.session.delete(config)
        db.session.commit()
        flash("Đã xoá cấu hình.")

    return redirect('/configure-ca')

from models.ca_config import CaConfiguration

@app.route('/clinic-rooms')
def clinic_rooms():
    rooms = ClinicRoom.query.all()
    return render_template('clinic_rooms.html', rooms=rooms)

@app.route('/clinic-rooms/add', methods=['POST'])
def add_clinic_room():
    name = request.form['name']
    description = request.form.get('description')
    db.session.add(ClinicRoom(name=name, description=description))
    db.session.commit()
    return redirect('/clinic-rooms')

@app.route('/clinic-rooms/delete/<int:room_id>')
def delete_clinic_room(room_id):
    room = ClinicRoom.query.get(room_id)
    db.session.delete(room)
    db.session.commit()
    return redirect('/clinic-rooms')

@app.route('/clinic-rooms/update/<int:room_id>', methods=['POST'])
def update_clinic_room(room_id):
    room = ClinicRoom.query.get(room_id)
    room.name = request.form['name']
    room.description = request.form.get('description')
    db.session.commit()
    return redirect('/clinic-rooms')

from collections import defaultdict

from flask import send_file, request
from openpyxl.styles import Alignment, Font
from io import BytesIO
from collections import defaultdict
from datetime import datetime, timedelta
from models import User, Shift, Schedule, ClinicRoom  # Đảm bảo bạn có các model này
import re  # <== DÒNG CẦN THÊM

@app.route('/export-clinic-schedule')
def export_clinic_schedule():
    # Bước 1: Nhận tham số ngày
    start_str = request.args.get('start')
    end_str = request.args.get('end')
    if not start_str or not end_str:
        return "Thiếu thông tin ngày", 400

    start_date = datetime.strptime(start_str, '%Y-%m-%d').date()
    end_date = datetime.strptime(end_str, '%Y-%m-%d').date()
    date_range = [start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)]

    # Bước 2: Lấy dữ liệu từ database
    rooms = ClinicRoom.query.all()
    schedules = Schedule.query.join(User).join(Shift).filter(Schedule.work_date.between(start_date, end_date)).all()

    # Bước 3: Tổ chức dữ liệu lịch trực
    clinic_schedule = {room.name: defaultdict(str) for room in rooms}
    user_positions = {}

    for s in schedules:
        user = s.user
        user_positions[user.name] = user.position
        shift_key = s.shift.name.lower().replace(" ", "")
        for room in rooms:
            room_key = room.name.lower().replace(" ", "")
            if room_key in shift_key:
                clinic_schedule[room.name][s.work_date] += f"{user.name}\n"
                break

    # Bước 4: Tạo file Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Lịch phòng khám"

    # Dòng tiêu đề
    ws.cell(row=1, column=1, value="Phòng Khám")
    for col, d in enumerate(date_range, start=2):
        ws.cell(row=1, column=col, value=d.strftime('%a %d/%m'))

    # Nội dung từng phòng
    for row_idx, (room, shifts) in enumerate(clinic_schedule.items(), start=2):
        ws.cell(row=row_idx, column=1, value=room)
        for col_idx, d in enumerate(date_range, start=2):
            raw = shifts[d].strip()
            formatted = []
            for name in raw.split("\n"):
                pos = user_positions.get(name, "").lower()
                prefix = "BS." if "bs" in pos or "bác" in pos else "ĐD." if "đd" in pos or "điều" in pos else ""
                if name.strip():
                    formatted.append(f"{prefix} {name.strip()}")
            value = "\n".join(formatted)
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    # Căn giữa dòng tiêu đề
    for col in ws[1]:
        col.alignment = Alignment(horizontal="center", vertical="center")
        col.font = Font(bold=True)

    # Xuất file
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="lich_phong_kham.xlsx")

@app.route('/change-password', methods=['GET', 'POST'])
def change_password():
    if 'user_id' not in session:
        return redirect('/login')

    user = db.session.get(User, session['user_id'])  # ✅ Dòng này đã cập 

    if request.method == 'POST':
        current_pw = request.form['current_password']
        new_pw = request.form['new_password']
        confirm_pw = request.form['confirm_password']

        if current_pw != user.password:
            flash("❌ Mật khẩu hiện tại không đúng.", "danger")
        elif new_pw != confirm_pw:
            flash("❌ Mật khẩu mới không khớp.", "danger")
        else:
            user.password = new_pw
            db.session.commit()
            flash("✅ Đổi mật khẩu thành công.", "success")
            return redirect('/')

    return render_template('change_password.html', user=user)

import re
from datetime import datetime, timedelta
from collections import defaultdict
from flask import render_template, request


@app.route('/print-clinic-schedule')
def print_clinic_schedule():
    from collections import defaultdict

    start_str = request.args.get('start')
    end_str = request.args.get('end')
    if not start_str or not end_str:
        return "Thiếu thông tin ngày bắt đầu hoặc kết thúc.", 400

    start_date = datetime.strptime(start_str, '%Y-%m-%d').date()
    end_date = datetime.strptime(end_str, '%Y-%m-%d').date()
    date_range = [start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)]

    # Lấy danh sách phòng khám
    all_rooms = ClinicRoom.query.all()
    rooms_dict = {room.name.lower(): room.name for room in all_rooms if "tiếp đón" not in room.name.lower()}

    # Khởi tạo dữ liệu lịch (dùng key 'phong_kham' viết thường)
    clinic_schedule = {
        "tiep_don": defaultdict(list),
        "phong_kham": {name: defaultdict(list) for name in rooms_dict.values()}
    }

    # Lấy dữ liệu phân công
    schedules = Schedule.query.join(User).join(Shift).filter(
        Schedule.work_date.between(start_date, end_date),
        Shift.name.ilike('%phòng khám%') | Shift.name.ilike('%tiếp đón%')
    ).all()

    # Tạo bảng chức vụ người dùng
    user_positions = {}
    for s in schedules:
        name = s.user.name
        user_positions[name] = s.user.position or ""
        date = s.work_date
        shift_name = s.shift.name.lower()

        if "tiếp đón" in shift_name:
            clinic_schedule["tiep_don"][date].append(name)
        else:
            for room_key in rooms_dict:
                if room_key in shift_name:
                    room_name = rooms_dict[room_key]
                    clinic_schedule["phong_kham"][room_name][date].append(name)
                    break

    # 1. Loại bỏ phòng trống
    clinic_schedule["phong_kham"] = {
        name: day_dict for name, day_dict in clinic_schedule["phong_kham"].items()
        if any(day_dict[d] for d in date_range)
    }

    # 2. Sắp xếp phòng theo thứ tự chuẩn
    desired_order = [
        "phòng khám 1", "phòng khám 2", "phòng khám 3",
        "phòng khám ngoại", "phòng khám tmh", "phòng khám rhm",
        "phòng khám mắt", "phòng khám 8 (tc)", "phòng khám 9 (tc)"
    ]
    ordered_schedule = {}
    for name in desired_order:
        original_name = rooms_dict.get(name)
        if original_name in clinic_schedule["phong_kham"]:
            ordered_schedule[original_name] = clinic_schedule["phong_kham"][original_name]
    clinic_schedule["phong_kham"] = ordered_schedule

    # Tạo danh sách rooms từ lịch đã sắp xếp
    rooms = list(clinic_schedule["phong_kham"].keys())

    return render_template(
        'print-clinic-schedule.html',
        start_date=start_date,
        end_date=end_date,
        date_range=date_range,
        clinic_schedule=clinic_schedule,
        user_positions=user_positions,
        rooms=rooms,
        now=datetime.now(),
        get_titled_names=get_titled_names
    )

# Các route như /print-clinic-schedule ở trên...

# === Cuối file hoặc phần helper ===
def get_titled_names(raw_names, user_positions):
    result = []
    for name in raw_names.split("\n"):
        name = name.strip()
        if name:
            role = user_positions.get(name)
            prefix = ""
            if role == "BS":
                prefix = "BS. "
            elif role == "ĐD":
                prefix = "ĐD. "
            result.append(f"{prefix}{name}")
    return "<br>".join(result)

@app.route('/print-clinic-dept-schedule')
def print_clinic_dept_schedule():
    from collections import defaultdict
    import re

    start_str = request.args.get('start')
    end_str = request.args.get('end')
    if not start_str or not end_str:
        return "Thiếu thông tin ngày bắt đầu hoặc kết thúc.", 400

    start_date = datetime.strptime(start_str, '%Y-%m-%d').date()
    end_date = datetime.strptime(end_str, '%Y-%m-%d').date()
    date_range = [start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)]

    # Lấy danh sách phòng khám
    all_rooms = ClinicRoom.query.all()
    rooms_dict = {room.name.lower(): room.name for room in all_rooms if "tiếp đón" not in room.name.lower()}

    # Khởi tạo dữ liệu lịch (dùng key 'phong_kham' viết thường)
    clinic_schedule = {
        "tiep_don": defaultdict(list),
        "phong_kham": {name: defaultdict(list) for name in rooms_dict.values()}
    }

    # Lấy dữ liệu phân công
    schedules = Schedule.query.join(User).join(Shift).filter(
        Schedule.work_date.between(start_date, end_date),
        Shift.name.ilike('%phòng khám%') | Shift.name.ilike('%tiếp đón%')
    ).all()

    # Tạo bảng chức vụ người dùng
    user_positions = {}
    for s in schedules:
        name = s.user.name
        user_positions[name] = s.user.position or ""
        date = s.work_date
        shift_name = s.shift.name.lower()

        if "tiếp đón" in shift_name:
            clinic_schedule["tiep_don"][date].append(name)
        else:
            for room_key in rooms_dict:
                if room_key in shift_name:
                    room_name = rooms_dict[room_key]
                    clinic_schedule["phong_kham"][room_name][date].append(name)
                    break

    # 1. Loại bỏ phòng trống
    clinic_schedule["phong_kham"] = {
        name: day_dict for name, day_dict in clinic_schedule["phong_kham"].items()
        if any(day_dict[d] for d in date_range)
    }

    # 2. Sắp xếp phòng theo thứ tự chuẩn
    desired_order = [
        "phòng khám 1", "phòng khám 2", "phòng khám 3",
        "phòng khám ngoại", "phòng khám tmh", "phòng khám rhm",
        "phòng khám mắt", "phòng khám 8 (tc)", "phòng khám 9 (tc)"
    ]
    ordered_schedule = {}
    for name in desired_order:
        original_name = rooms_dict.get(name)
        if original_name in clinic_schedule["phong_kham"]:
            ordered_schedule[original_name] = clinic_schedule["phong_kham"][original_name]
    clinic_schedule["phong_kham"] = ordered_schedule

    # Tạo danh sách rooms từ lịch đã sắp xếp
    rooms = list(clinic_schedule["phong_kham"].keys())

    return render_template(
        'print-clinic-dept-schedule.html',
        start_date=start_date,
        end_date=end_date,
        date_range=date_range,
        clinic_schedule=clinic_schedule,
        user_positions=user_positions,
        rooms=rooms,
        now=datetime.now(),
        get_titled_names=get_titled_names
    )

def get_titled_names(name_input, user_positions):
    # Nếu đầu vào là chuỗi, chuyển sang danh sách
    if isinstance(name_input, str):
        names = [n.strip() for n in name_input.split(',') if n.strip()]
    elif isinstance(name_input, list):
        names = [n.strip() for n in name_input if n.strip()]
    else:
        return ''

    titled = []
    for name in names:
        title = user_positions.get(name, '').strip().upper()
        if title:
            titled.append(f"{title}. {name}")
        else:
            titled.append(name)

    return '<br>'.join(titled)

# Đăng ký vào template
@app.context_processor
def inject_helpers():
    return dict(get_titled_names=get_titled_names)

@app.context_processor
def inject_user():
    user = None
    if 'user_id' in session:
        user = User.query.get(session['user_id'])
    return dict(user=user)

from flask import render_template, request, redirect
from models.shift_rate_config import ShiftRateConfig

@app.route('/shift-rate-config', methods=['GET', 'POST'])
def shift_rate_config():
    # Cho phép admin và admin1
    if session.get('role') not in ('admin', 'admin1'):
        return "Chỉ admin mới được phép truy cập."

    if request.method == 'POST':
        ca_loai = request.form['ca_loai']
        truc_loai = request.form['truc_loai']
        ngay_loai = request.form['ngay_loai']
        try:
            don_gia = int(request.form['don_gia'])
        except (TypeError, ValueError):
            don_gia = 0

        new_rate = ShiftRateConfig(
            ca_loai=ca_loai,
            truc_loai=truc_loai,
            ngay_loai=ngay_loai,
            don_gia=don_gia
        )
        db.session.add(new_rate)
        db.session.commit()
        return redirect('/shift-rate-config')

    rates = ShiftRateConfig.query.all()
    return render_template('shift_rate_config.html', rates=rates)

@app.route('/shift-rate-config/delete/<int:rate_id>')
def delete_shift_rate(rate_id):
    if session.get('role') != 'admin':
        return "Không có quyền"
    rate = ShiftRateConfig.query.get_or_404(rate_id)
    db.session.delete(rate)
    db.session.commit()
    return redirect('/shift-rate-config')

from models.hscc_department import HSCCDepartment  # Import đầu file

@app.route('/configure-hscc', methods=['GET', 'POST'])
def configure_hscc():
    if session.get('role') not in ('admin', 'admin1'):
        return "Chỉ admin hoặc admin1 được phép truy cập."

    if request.method == 'POST':
        new_dept = request.form.get('department').strip()
        if new_dept and not HSCCDepartment.query.filter_by(department_name=new_dept).first():
            hscc = HSCCDepartment(department_name=new_dept)
            db.session.add(hscc)
            db.session.commit()

    departments = HSCCDepartment.query.all()
    return render_template('configure_hscc.html', departments=departments)


@app.route('/configure-hscc/update/<int:id>', methods=['POST'])
def update_hscc(id):
    if session.get('role') not in ('admin', 'admin1'):
        return "Chỉ admin hoặc admin1 được phép truy cập."

    hscc = HSCCDepartment.query.get_or_404(id)
    new_name = request.form.get('new_name') or request.form.get('department')

    if new_name:
        hscc.department_name = new_name.strip()
        db.session.commit()

    return redirect('/configure-hscc')


@app.route('/shift-payment-view')
def shift_payment_view():
    user_role = session.get('role')
    user_dept = session.get('department')

    from collections import defaultdict

    def classify_day(date):
        ngay_le = {'01-01', '04-30', '05-01', '09-02'}
        mmdd = date.strftime('%m-%d')
        weekday = date.weekday()
        if mmdd in ngay_le:
            return 'ngày_lễ'
        elif weekday >= 5:
            return 'ngày_nghỉ'
        else:
            return 'ngày_thường'

    ca_chon = request.args.get('mode', '16h')
    selected_department = request.args.get('department')

    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')

    # Nếu không có ngày → mặc định tháng hiện tại
    if not start_date or not end_date:
        today = datetime.today()
        start_date_dt = today.replace(day=1)
        next_month = today.replace(day=28) + timedelta(days=4)
        last_day = (next_month - timedelta(days=next_month.day)).day
        end_date_dt = today.replace(day=last_day)
        start_date = start_date_dt.strftime('%Y-%m-%d')
        end_date = end_date_dt.strftime('%Y-%m-%d')
    else:
        start_date_dt = datetime.strptime(start_date, '%Y-%m-%d').date()
        end_date_dt = datetime.strptime(end_date, '%Y-%m-%d').date()

    # Danh sách khoa
    if user_role in ['admin', 'admin1']:
        departments = sorted([d[0] for d in db.session.query(User.department)
                              .filter(User.department.isnot(None))
                              .distinct().all()])
        departments.insert(0, 'Tất cả')
        if not selected_department:
            selected_department = 'Tất cả'
    else:
        departments = [user_dept] if user_dept else []
        selected_department = user_dept

    hscc_depts = [d.department_name for d in HSCCDepartment.query.all()]
    rates = {(r.ca_loai.lower().strip(), r.truc_loai.lower().strip(), r.ngay_loai.lower().strip()): r.don_gia
             for r in ShiftRateConfig.query.all()}

    # Query dữ liệu
    query = (
        Schedule.query
        .join(User).join(Shift)
        .filter(Schedule.work_date >= start_date_dt, Schedule.work_date <= end_date_dt)
        .filter(Shift.duration == (16 if ca_chon == '16h' else 24))
    )

    # Chỉ lọc khi không phải "Tất cả"
    if selected_department and selected_department != 'Tất cả':
        query = query.filter(User.department == selected_department)

    schedules = query.all()

    # Danh sách ca hợp lệ
    valid_shifts = [
        "trực 16h", "trực 16h t7cn",
        "trực 24h", "trực 24h t7cn",
        "trực lễ16h", "trực lễ 24h"
    ]

    # Gom dữ liệu
    data = defaultdict(lambda: defaultdict(int))
    for s in schedules:
        user = s.user
        shift = s.shift

        shift_name = shift.name.strip().lower()

        # Chỉ tính nếu ca hợp lệ
        if shift_name not in valid_shifts:
            continue

        # Bỏ ca thường trú
        if "thường trú" in shift_name:
            continue

        ngay_loai = classify_day(s.work_date)
        truc_loai = "HSCC" if user.department in hscc_depts else "thường"
        key = (truc_loai, ngay_loai)
        data[user][key] += 1

    # Sắp xếp chức danh
    priority_order = ['GĐ', 'PGĐ', 'TK', 'PTK', 'PK', 'BS', 'ĐDT', 'ĐD', 'KTV', 'NV', 'HL', 'BV']
    def get_priority(pos):
        pos = pos.upper() if pos else ''
        for i, p in enumerate(priority_order):
            if p in pos:
                return i
        return len(priority_order)

    # Chuẩn bị dữ liệu hiển thị
    rows = []
    co_ngay_le = False
    for user, info in data.items():
        if user.role == 'admin':
            continue

        row = {
            'user': user,
            'tong_ngay': sum(info.values()),
            'tien_ca': 0,
            'tien_an': sum(info.values()) * 15000,
            'tong_tien': 0,
            'is_contract': user.contract_type == "Hợp đồng",
            'detail': {},
            'priority': get_priority(user.position)
        }

        for key in [
            ("thường", "ngày_thường"), ("HSCC", "ngày_thường"),
            ("thường", "ngày_nghỉ"), ("HSCC", "ngày_nghỉ"),
            ("thường", "ngày_lễ"), ("HSCC", "ngày_lễ")
        ]:
            so_ngay = info.get(key, 0)
            if key[1] == 'ngày_lễ' and so_ngay > 0:
                co_ngay_le = True

            # Chuẩn hóa toàn bộ key khi lookup rates
            ca_key = ca_chon.lower().replace("trực", "").strip()
            truc_key = key[0].lower().strip()
            ngay_key = key[1].lower().strip()

            don_gia = (
                rates.get((ca_key, truc_key, ngay_key), 0)
                or rates.get((ca_chon, truc_key, ngay_key), 0)
                or rates.get((ca_key, key[0], key[1]), 0)
                or rates.get((ca_chon, key[0], key[1]), 0)
            )

            row['detail'][key] = {'so_ngay': so_ngay, 'don_gia': don_gia}
            row['tien_ca'] += so_ngay * don_gia

        row['tong_tien'] = row['tien_ca'] + row['tien_an']
        rows.append(row)

    # Sort theo chức danh + tên
    rows = sorted(rows, key=lambda r: (r['priority'], r['user'].name))

    thang = start_date_dt.month
    nam = start_date_dt.year

    return render_template(
        "shift_payment_view.html",
        ca_chon=ca_chon,
        rows=rows,
        departments=departments,
        selected_department=selected_department,
        mode=ca_chon,
        start_date=start_date,
        end_date=end_date,
        thang=thang,
        nam=nam,
        co_ngay_le=co_ngay_le
    )


@app.route('/tong-hop-cong-truc-view')
@login_required
def tong_hop_cong_truc_view():
    from collections import defaultdict
    from models.user import User
    from models.schedule import Schedule
    from models.shift import Shift
    from models.hscc_department import HSCCDepartment

    user_role = session.get('role')
    user_dept = session.get('department')

    # --- Danh sách khoa ---
    if user_role in ['admin', 'admin1']:
        departments = sorted([d[0] for d in db.session.query(User.department)
                              .filter(User.department.isnot(None))
                              .distinct().all()])
        departments.insert(0, 'Tất cả')
    else:
        departments = [user_dept] if user_dept else []

    # --- Lấy tham số department ---
    if user_role in ['admin', 'admin1']:
        selected_department = request.args.get('department', 'Tất cả')
    else:
        selected_department = user_dept  # User thường chỉ thấy khoa mình

    # --- Lấy ngày bắt đầu/kết thúc và chuyển sang kiểu date ---
    start_date_str = request.args.get('start_date')
    end_date_str = request.args.get('end_date')
    mode = request.args.get('mode', '16h')

    today = datetime.now()

    try:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date() if start_date_str else None
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date() if end_date_str else None
    except Exception:
        start_date, end_date = None, None

    # Nếu chưa chọn ngày → báo lỗi
    if not start_date or not end_date:
        return render_template(
            'tong_hop_cong_truc_view.html',
            rows=[],
            sum_row={},
            departments=departments,
            selected_department=selected_department,
            start_date=start_date_str,
            end_date=end_date_str,
            default_start=start_date_str,
            default_end=end_date_str,
            thang=today.month,
            nam=today.year,
            mode=mode,
            error_message="Bạn chưa chọn ngày bắt đầu và ngày kết thúc để xem báo cáo!"
        )

    # Lấy tháng/năm để hiển thị
    thang = start_date.month
    nam = start_date.year

    # --- Query lịch trực ---
    query = Schedule.query.join(User).join(Shift).filter(
        Schedule.work_date.between(start_date, end_date)
    )
    if selected_department not in ['Tất cả', 'all', None]:
        query = query.filter(User.department.ilike(selected_department))

    schedules = query.all()

    # --- Lấy danh sách HSCC (bọc try/except nếu bảng trống) ---
    try:
        hscc_depts = [d.department_name for d in HSCCDepartment.query.all()]
    except Exception:
        hscc_depts = []

    result_by_user = defaultdict(lambda: defaultdict(lambda: {'so_ngay': 0}))
    summary = defaultdict(int)

    # Danh sách ca hợp lệ
    valid_shifts = [
        "trực 16h", "trực 16h t7cn",
        "trực 24h", "trực 24h t7cn",
        "trực lễ16h", "trực lễ 24h"
    ]

    for s in schedules:
        if not s.shift or not s.user:
            continue

        shift_name = s.shift.name.strip().lower()

        # Chỉ tính ca hợp lệ
        if shift_name not in valid_shifts:
            continue

        # Bỏ ca thường trú
        if 'thường trú' in shift_name:
            continue

        # Bỏ các từ khóa nghỉ
        skip_keywords = ['nghỉ trực', 'nghỉ phép', 'làm ngày', 'làm 1/2 ngày', 'làm 1/2 ngày c', 'phòng khám']
        if any(x in shift_name for x in skip_keywords):
            continue

        # Lọc theo mode
        if mode == '24h' and '24h' not in shift_name:
            continue
        if mode == '16h' and '24h' in shift_name:
            continue

        loai_ca = 'HSCC' if s.user.department in hscc_depts else 'thường'

        mmdd = s.work_date.strftime('%m-%d')
        weekday = s.work_date.weekday()

        # Xác định loại ngày
        if mmdd in ['01-01', '04-30', '05-01', '09-02']:
            loai_ngay = 'ngày_lễ'
        elif weekday >= 5:
            loai_ngay = 'ngày_nghỉ'
        else:
            loai_ngay = 'ngày_thường'

        result_by_user[s.user_id][(loai_ca, loai_ngay)]['so_ngay'] += 1
        summary[(loai_ca, loai_ngay)] += 1

    user_ids = list(result_by_user.keys())
    users = User.query.filter(User.id.in_(user_ids), User.role != 'admin').all() if user_ids else []

    # Thứ tự ưu tiên chức danh
    priority_order = ['GĐ', 'PGĐ', 'TK', 'PTK', 'PK', 'BS', 'ĐDT', 'ĐD', 'KTV', 'NV', 'HL', 'BV']
    def get_priority(pos):
        pos = pos.upper() if pos else ''
        for i, p in enumerate(priority_order):
            if p in pos:
                return i
        return len(priority_order)

    # Chuẩn bị dữ liệu hiển thị
    rows = []
    for user in users:
        detail = result_by_user[user.id]
        rows.append({
            'user': user,
            'detail': detail,
            'tong_ngay': sum([v['so_ngay'] for v in detail.values()]),
            'ghi_chu': ''
        })

    rows.sort(key=lambda x: get_priority(x['user'].position))

    sum_row = {
        'detail': summary,
        'tong_ngay': sum(summary.values())
    }

    # --- Xử lý tên khoa/phòng hiển thị ---
    if selected_department in ['Tất cả', 'all', None]:
        if user_role == 'admin1' and user_dept:
            dept_display = user_dept
        else:
            dept_display = ""
    else:
        dept_display = selected_department

    return render_template(
        'tong_hop_cong_truc_view.html',
        rows=rows,
        sum_row=sum_row,
        departments=departments,
        selected_department=dept_display,
        start_date=start_date_str,
        end_date=end_date_str,
        default_start=start_date_str,
        default_end=end_date_str,
        thang=thang,
        nam=nam,
        mode=mode
    )

from models.unit_config import UnitConfig  # import model cấu hình đơn vị

@app.route('/tong-hop-cong-truc-print')
@login_required
def tong_hop_cong_truc_print():
    from collections import defaultdict
    from models.user import User
    from models.schedule import Schedule
    from models.shift import Shift
    from models.hscc_department import HSCCDepartment
    from models.unit_config import UnitConfig

    user_role = session.get('role')
    user_dept = session.get('department')

    # --- Lấy tham số ---
    selected_department = request.args.get('department', '')
    start_date_str = request.args.get('start_date')
    end_date_str = request.args.get('end_date')
    mode = request.args.get('mode', '16h')

    today = datetime.now()

    # --- Chuẩn hoá ngày ---
    try:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date() if start_date_str else None
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date() if end_date_str else None
    except Exception:
        start_date, end_date = None, None

    # Nếu chưa chọn ngày → báo lỗi (giữ nguyên template in)
    if not start_date or not end_date:
        return render_template(
            'tong_hop_cong_truc.html',
            rows=[],
            sum_row={},
            selected_department=selected_department,
            start_date=start_date_str,
            current_day=today.day,
            current_month=today.month,
            current_year=today.year,
            thang=today.month,
            nam=today.year,
            mode=mode,
            unit_config=UnitConfig.query.first(),
            error_message="Bạn chưa chọn ngày bắt đầu và ngày kết thúc để in báo cáo!"
        )

    thang = start_date.month
    nam = start_date.year

    # --- Query lịch trực (giống view) ---
    query = Schedule.query.join(User).join(Shift).filter(
        Schedule.work_date.between(start_date, end_date)
    )
    if selected_department not in ['Tất cả', 'all', None]:
        # dùng ilike như view
        query = query.filter(User.department.ilike(selected_department))

    schedules = query.all()

    # --- Danh sách HSCC (giống view, có try/except) ---
    try:
        hscc_depts = [d.department_name for d in HSCCDepartment.query.all()]
    except Exception:
        hscc_depts = []

    result_by_user = defaultdict(lambda: defaultdict(lambda: {'so_ngay': 0}))
    summary = defaultdict(int)

    # --- Danh sách ca hợp lệ (y như view) ---
    valid_shifts = [
        "trực 16h", "trực 16h t7cn",
        "trực 24h", "trực 24h t7cn",
        "trực lễ16h", "trực lễ 24h"
    ]
    skip_keywords = ['nghỉ trực', 'nghỉ phép', 'làm ngày', 'làm 1/2 ngày', 'làm 1/2 ngày c', 'phòng khám']

    for s in schedules:
        if not s.shift or not s.user:
            continue

        shift_name = s.shift.name.strip().lower()

        # chỉ tính ca hợp lệ
        if shift_name not in valid_shifts:
            continue

        # bỏ thường trú
        if 'thường trú' in shift_name:
            continue

        # bỏ các từ khoá nghỉ/làm ngày
        if any(x in shift_name for x in skip_keywords):
            continue

        # lọc theo mode
        if mode == '24h' and '24h' not in shift_name:
            continue
        if mode == '16h' and '24h' in shift_name:
            continue

        loai_ca = 'HSCC' if s.user.department in hscc_depts else 'thường'

        mmdd = s.work_date.strftime('%m-%d')
        weekday = s.work_date.weekday()
        if mmdd in ['01-01', '04-30', '05-01', '09-02']:
            loai_ngay = 'ngày_lễ'
        elif weekday >= 5:
            loai_ngay = 'ngày_nghỉ'
        else:
            loai_ngay = 'ngày_thường'

        result_by_user[s.user_id][(loai_ca, loai_ngay)]['so_ngay'] += 1
        summary[(loai_ca, loai_ngay)] += 1

    # --- Lấy user giống view: BỎ admin ---
    user_ids = list(result_by_user.keys())
    users = User.query.filter(User.id.in_(user_ids), User.role != 'admin').all() if user_ids else []

    # --- Sắp xếp theo chức danh (giữ nguyên) ---
    priority_order = ['GĐ', 'PGĐ', 'TK', 'PTK', 'PK', 'BS', 'ĐDT', 'ĐD', 'KTV', 'NV', 'HL', 'BV']
    def get_priority(pos):
        pos = pos.upper() if pos else ''
        for i, p in enumerate(priority_order):
            if p in pos:
                return i
        return len(priority_order)

    rows = []
    for user in users:
        detail = result_by_user[user.id]
        rows.append({
            'user': user,
            'detail': detail,
            'tong_ngay': sum([v['so_ngay'] for v in detail.values()]),
            'ghi_chu': ''
        })
    rows.sort(key=lambda x: get_priority(x['user'].position))

    sum_row = {
        'detail': summary,
        'tong_ngay': sum(summary.values())
    }

    # --- Hiển thị tên khoa/phòng (đồng bộ với view) ---
    if selected_department in ['Tất cả', 'all', None]:
        if user_role == 'admin1' and user_dept:
            dept_display = user_dept
        else:
            dept_display = ""
    else:
        dept_display = selected_department

    return render_template(
        'tong_hop_cong_truc.html',
        rows=rows,
        sum_row=sum_row,
        selected_department=dept_display,
        start_date=start_date_str,
        current_day=today.day,
        current_month=today.month,
        current_year=today.year,
        thang=thang,
        nam=nam,
        mode=mode,
        unit_config=UnitConfig.query.first()
    )

@app.route('/tong-hop-cong-truc-view/export-excel')
@login_required
def tong_hop_cong_truc_export_excel():
    from io import BytesIO
    from collections import defaultdict
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side

    # ==== Tham số giống VIEW ====
    user_role = session.get('role')
    user_dept = session.get('department')

    selected_department = request.args.get('department')
    start_date_str = request.args.get('start_date')
    end_date_str   = request.args.get('end_date')
    mode = request.args.get('mode', '16h')

    start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
    end_date   = datetime.strptime(end_date_str,   '%Y-%m-%d').date()

    # ==== Query lịch trực y như VIEW ====
    query = Schedule.query.join(User).join(Shift).filter(
        Schedule.work_date.between(start_date, end_date)
    )
    if user_role in ['admin','admin1']:
        if selected_department not in (None, '', 'Tất cả', 'all'):
            query = query.filter(User.department == selected_department)
    else:
        query = query.filter(User.department == user_dept)

    schedules = query.all()

    # ==== HSCC depts ====
    try:
        hscc_depts = [d.department_name for d in HSCCDepartment.query.all()]
    except Exception:
        hscc_depts = []

    # ==== Logic LỌC & ĐẾM GIỐNG HỆT VIEW ====
    result_by_user = defaultdict(lambda: defaultdict(lambda: {'so_ngay': 0}))
    summary = defaultdict(int)

    valid_shifts = [
        "trực 16h", "trực 16h t7cn",
        "trực 24h", "trực 24h t7cn",
        "trực lễ16h", "trực lễ 24h"
    ]
    for s in schedules:
        if not s.shift or not s.user:
            continue

        shift_name = (s.shift.name or '').strip().lower()

        # chỉ tính ca hợp lệ
        if shift_name not in valid_shifts:
            continue

        # bỏ ca thường trú
        if 'thường trú' in shift_name:
            continue

        # bỏ các từ khóa nghỉ/làm ngày/phòng khám
        skip_keywords = ['nghỉ trực', 'nghỉ phép', 'làm ngày', 'làm 1/2 ngày', 'làm 1/2 ngày c', 'phòng khám']
        if any(x in shift_name for x in skip_keywords):
            continue

        # lọc theo mode
        if mode == '24h' and '24h' not in shift_name:
            continue
        if mode == '16h' and '24h' in shift_name:
            continue

        loai_ca = 'HSCC' if s.user.department in hscc_depts else 'thường'

        mmdd = s.work_date.strftime('%m-%d')
        weekday = s.work_date.weekday()

        # phân loại ngày (y như VIEW)
        if mmdd in ['01-01', '04-30', '05-01', '09-02']:
            loai_ngay = 'ngày_lễ'
        elif weekday >= 5:
            loai_ngay = 'ngày_nghỉ'
        else:
            loai_ngay = 'ngày_thường'

        result_by_user[s.user_id][(loai_ca, loai_ngay)]['so_ngay'] += 1
        summary[(loai_ca, loai_ngay)] += 1

    user_ids = list(result_by_user.keys())
    users = User.query.filter(User.id.in_(user_ids), User.role != 'admin').all() if user_ids else []

    # sắp xếp theo chức danh như VIEW
    priority_order = ['GĐ', 'PGĐ', 'TK', 'PTK', 'PK', 'BS', 'ĐDT', 'ĐD', 'KTV', 'NV', 'HL', 'BV']
    def get_priority(pos):
        pos = (pos or '').upper()
        for i, p in enumerate(priority_order):
            if p in pos:
                return i
        return len(priority_order)
    users.sort(key=lambda u: (get_priority(u.position), (u.name or '').lower()))

    # ==== Tạo Excel ====
    wb = Workbook()
    ws = wb.active
    ws.title = "Tổng hợp công trực"

    bold = Font(bold=True)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Side(style='thin')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # tiêu đề
    ws.append(["STT","HỌ TÊN",
               "Trực thường - Ngày thường","Trực HSCC - Ngày thường",
               "Trực thường - Ngày nghỉ","Trực HSCC - Ngày nghỉ",
               "Trực thường - Ngày lễ","Trực HSCC - Ngày lễ",
               "Tổng số ngày trực","Ghi chú"])
    for c in ws[1]:
        c.font = bold; c.alignment = center; c.border = border

    def get_cnt(detail, key):
        return detail.get(key, {}).get('so_ngay', 0)

    for i, u in enumerate(users, 1):
        d = result_by_user[u.id]
        row = [
            i, u.name,
            get_cnt(d, ('thường','ngày_thường')),
            get_cnt(d, ('HSCC','ngày_thường')),
            get_cnt(d, ('thường','ngày_nghỉ')),
            get_cnt(d, ('HSCC','ngày_nghỉ')),
            get_cnt(d, ('thường','ngày_lễ')),
            get_cnt(d, ('HSCC','ngày_lễ')),
        ]
        row.append(sum(row[2:8]))  # tổng số ngày trực
        row.append("")             # ghi chú
        ws.append(row)
        for c in ws[ws.max_row]:
            c.alignment = center; c.border = border

    # có thể thêm dòng tổng cuối (nếu muốn giống view phần tổng)
    # s = summary
    # ws.append([])
    # ws.append(["TỔNG","","",
    #            s.get(('HSCC','ngày_thường'),0),
    #            s.get(('thường','ngày_nghỉ'),0),
    #            ...])

    bio = BytesIO()
    wb.save(bio); bio.seek(0)
    fname = f"tong_hop_cong_truc_{start_date:%Y%m%d}_{end_date:%Y%m%d}_{mode}.xlsx"
    return send_file(bio, as_attachment=True, download_name=fname,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

@app.route('/export-shift-payment-all')
def export_shift_payment_all():
    from io import BytesIO
    from collections import defaultdict
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side

    # --- Phân loại ngày theo lịch (giống bản in) ---
    def classify_day(d):
        ngay_le = {'01-01', '04-30', '05-01', '09-02'}  # có thể bổ sung nếu cần
        mmdd = d.strftime('%m-%d')
        if mmdd in ngay_le:
            return 'ngày_lễ'
        return 'ngày_nghỉ' if d.weekday() >= 5 else 'ngày_thường'

    # --- Tham số ---
    ca_chon = request.args.get('mode', '16h')                 # '16h' hoặc '24h'
    selected_department = request.args.get('department', 'all')
    start_date = request.args.get('start_date', '2025-06-01')
    end_date   = request.args.get('end_date',   '2025-06-30')

    start_date_dt = datetime.strptime(start_date, '%Y-%m-%d').date()
    end_date_dt   = datetime.strptime(end_date,   '%Y-%m-%d').date()
    thang, nam = start_date_dt.month, start_date_dt.year

    # --- Dữ liệu cấu hình ---
    hscc_depts = {d.department_name for d in HSCCDepartment.query.all()}
    rates = {(r.ca_loai, r.truc_loai, r.ngay_loai): r.don_gia for r in ShiftRateConfig.query.all()}

    # --- Lấy lịch theo khoảng & theo duration (chỉ 16h/24h như yêu cầu) ---
    wanted_duration = 16 if ca_chon == '16h' else 24
    q = (
        Schedule.query
        .join(User).join(Shift)
        .filter(Schedule.work_date >= start_date_dt,
                Schedule.work_date <= end_date_dt)
        .filter(Shift.duration == wanted_duration)
    )
    if selected_department != 'all':
        q = q.filter(User.department == selected_department)

    schedules = q.all()

    # --- Khử trùng lặp theo ngày: mỗi (user, date) chỉ tính 1 lần ---
    picked_per_day = {}  # key: (user_id, date) -> user_obj
    for s in schedules:
        if not s.user:
            continue
        key = (s.user_id, s.work_date)
        if key not in picked_per_day:
            picked_per_day[key] = s.user

    # --- Gom dữ liệu: user -> {(trực_loại, ngày_loại): count} ---
    data = defaultdict(lambda: defaultdict(int))
    for (uid, d), user in picked_per_day.items():
        ngay_loai = classify_day(d)
        truc_loai = "HSCC" if (user.department in hscc_depts) else "thường"
        data[user][(truc_loai, ngay_loai)] += 1

    # --- Tạo Excel ---
    wb = Workbook()
    ws = wb.active
    ws.title = f"BẢNG TRỰC {ca_chon}"

    bold = Font(bold=True)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Side(style='thin')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # Tiêu đề đầu trang
    ws.merge_cells("A1:M1"); ws["A1"] = "SỞ Y TẾ TỈNH GIA LAI"; ws["A1"].font = bold
    ws.merge_cells("A2:M2"); ws["A2"] = "BỆNH VIỆN NHI"; ws["A2"].font = bold
    ws.merge_cells("A4:M4"); ws["A4"] = f"BẢNG THANH TOÁN TIỀN TRỰC THÁNG {thang:02d} NĂM {nam}"
    ws["A4"].alignment = center; ws["A4"].font = Font(bold=True, size=13)

    # Header bảng
    ws.append([
        "STT", "HỌ TÊN",
        "Trực thường\n(Ngày thường)", "Trực HSCC\n(Ngày thường)",
        "Trực thường\n(Ngày nghỉ)",   "Trực HSCC\n(Ngày nghỉ)",
        "Trực thường\n(Ngày lễ)",     "Trực HSCC\n(Ngày lễ)",
        "Tổng số\nngày trực", "Tiền ca\n(QĐ 73)",
        "Tiền ăn\n(15k/ngày)", "Tổng cộng", "Ghi chú"
    ])
    for c in ws[ws.max_row]:
        c.font = bold; c.alignment = center; c.border = border

    # Ghi dữ liệu: sắp theo khoa rồi theo tên để in đẹp
    def cnt(info, key): return info.get(key, 0)

    sorted_users = sorted(
        data.items(),
        key=lambda kv: ((kv[0].department or ''), (kv[0].name or ''))
    )

    for i, (user, info) in enumerate(sorted_users, start=1):
        total_day = sum(info.values())
        tien_ca = 0

        row = [i, user.name]
        for key in [
            ("thường", "ngày_thường"), ("HSCC", "ngày_thường"),
            ("thường", "ngày_nghỉ"),   ("HSCC", "ngày_nghỉ"),
            ("thường", "ngày_lễ"),     ("HSCC", "ngày_lễ")
        ]:
            so_ngay = cnt(info, key)
            row.append(so_ngay)
            don_gia = rates.get((ca_chon, *key), 0)
            tien_ca += so_ngay * don_gia

        tien_an = total_day * 15000
        tong = tien_ca + tien_an
        ghi_chu = "HD" if getattr(user, "contract_type", "") == "Hợp đồng" else ""

        row += [total_day, tien_ca, tien_an, tong, ghi_chu]
        ws.append(row)
        for c in ws[ws.max_row]:
            c.alignment = center; c.border = border

    # Xuất file
    stream = BytesIO(); wb.save(stream); stream.seek(0)
    filename = f"BANG_THANH_TOAN_{thang:02d}_{nam}_{ca_chon}.xlsx"
    return send_file(stream, as_attachment=True, download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

@app.route('/print-shift-payment')
def print_shift_payment():
    from calendar import month_name
    from collections import defaultdict
    from datetime import datetime

    def classify_day(date):
        # Danh sách ngày lễ cố định
        ngay_le = {'01-01', '04-30', '05-01', '09-02'}
        mmdd = date.strftime('%m-%d')
        weekday = date.weekday()
        if mmdd in ngay_le:
            return 'ngày_lễ'
        elif weekday >= 5:
            return 'ngày_nghỉ'
        else:
            return 'ngày_thường'

    # Lấy params
    ca_chon = request.args.get('mode', '16h')
    selected_department = request.args.get('department', 'all')
    start_date = request.args.get('start_date', '2025-06-01')
    end_date = request.args.get('end_date', '2025-06-30')

    start_date_dt = datetime.strptime(start_date, '%Y-%m-%d').date()
    end_date_dt = datetime.strptime(end_date, '%Y-%m-%d').date()
    thang = start_date_dt.month
    nam = start_date_dt.year

    # Ngày in hiện tại
    today = datetime.now()
    current_day = today.day
    current_month = today.month
    current_year = today.year

    # Danh sách khoa HSCC
    hscc_depts = [d.department_name for d in HSCCDepartment.query.all()]
    rates = {(r.ca_loai, r.truc_loai, r.ngay_loai): r.don_gia for r in ShiftRateConfig.query.all()}

    # Query lịch trực
    query = (
        Schedule.query
        .join(User).join(Shift)
        .filter(Schedule.work_date >= start_date_dt, Schedule.work_date <= end_date_dt)
        .filter(Shift.duration == (16 if ca_chon == '16h' else 24))
    )
    if selected_department != 'all':
        query = query.filter(User.department == selected_department)

    schedules = query.all()

    # Gom dữ liệu
    data = defaultdict(lambda: defaultdict(int))
    for s in schedules:
        user = s.user
        shift = s.shift

        # Bỏ qua trực thường trú
        if "thường trú" in shift.name.strip().lower():
            continue

        ngay_loai = classify_day(s.work_date)
        truc_loai = "HSCC" if user.department in hscc_depts else "thường"
        key = (truc_loai, ngay_loai)
        data[user][key] += 1

    # Tổng cộng
    sum_row = {
        'tong_ngay': 0,
        'tien_ca': 0,
        'tien_an': 0,
        'tong_tien': 0,
        'detail': {
            ("thường", "ngày_thường"): 0,
            ("HSCC", "ngày_thường"): 0,
            ("thường", "ngày_nghỉ"): 0,
            ("HSCC", "ngày_nghỉ"): 0,
            ("thường", "ngày_lễ"): 0,
            ("HSCC", "ngày_lễ"): 0
        }
    }

    # Danh sách ưu tiên chức danh
    priority_order = ['GĐ', 'PGĐ', 'TK', 'PTK', 'PK', 'BS', 'ĐDT', 'ĐD', 'KTV', 'NV', 'HL', 'BV']

    def get_priority(pos):
        pos = pos.upper() if pos else ''
        for i, p in enumerate(priority_order):
            if p in pos:
                return i
        return len(priority_order)

    # Tạo rows
    rows = []
    for user, info in data.items():
        row = {
            'user': user,
            'tong_ngay': sum(info.values()),
            'tien_ca': 0,
            'tien_an': sum(info.values()) * 15000,
            'tong_tien': 0,
            'is_contract': user.contract_type == "Hợp đồng",
            'ghi_chu': 'HĐ' if user.contract_type == 'Hợp đồng' else '',
            'detail': {},
            'priority': get_priority(user.position)
        }

        for key in sum_row['detail'].keys():
            so_ngay = info.get(key, 0)
            don_gia = rates.get((ca_chon, *key), 0)
            row['detail'][key] = {'so_ngay': so_ngay, 'don_gia': don_gia}
            row['tien_ca'] += so_ngay * don_gia
            sum_row['detail'][key] += so_ngay

        row['tong_tien'] = row['tien_ca'] + row['tien_an']

        sum_row['tong_ngay'] += row['tong_ngay']
        sum_row['tien_ca'] += row['tien_ca']
        sum_row['tien_an'] += row['tien_an']
        sum_row['tong_tien'] += row['tong_tien']

        rows.append(row)

    # Sắp xếp theo thứ tự ưu tiên chức danh và tên
    rows = sorted(rows, key=lambda r: (r['priority'], r['user'].name))

    return render_template(
        "shift_payment_print.html",
        ca_chon=ca_chon,
        rows=rows,
        sum_row=sum_row,
        selected_department=selected_department,
        mode=ca_chon,
        start_date=start_date,
        end_date=end_date,
        thang=thang,
        nam=nam,
        current_day=current_day,
        current_month=current_month,
        current_year=current_year,
        tong_tien_bang_chu=num2text(int(sum_row['tong_tien']))
    )

from utils import num2text

from utils.num2text import num2text  # ✅ chỉ lấy hàm, không lấy module
text = num2text(1530000)
# Kết quả: "Một triệu năm trăm ba mươi nghìn đồng"

from datetime import datetime  # Đảm bảo đã import ở đầu file

@app.route('/print-shift-payment-summary')
def print_shift_payment_summary():
    from collections import defaultdict

    def classify_day(date):
        ngay_le = {'01-01', '04-30', '05-01', '09-02'}
        mmdd = date.strftime('%m-%d')
        weekday = date.weekday()
        if mmdd in ngay_le:
            return 'ngày_lễ'
        elif weekday >= 5:
            return 'ngày_nghỉ'
        else:
            return 'ngày_thường'

    def roman(num):
        roman_map = zip(
            (1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1),
            ("M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I")
        )
        result = ''
        for i, r in roman_map:
            while num >= i:
                result += r
                num -= i
        return result

    # Thứ tự chức danh ưu tiên
    priority_order = ['GĐ', 'PGĐ', 'TK', 'PTK', 'PK', 'BS', 'ĐDT', 'ĐD', 'KTV', 'NV', 'HL', 'BV']

    def get_priority(pos):
        pos = pos.upper() if pos else ''
        for i, p in enumerate(priority_order):
            if p in pos:
                return i
        return len(priority_order)

    ca_chon = request.args.get('mode', '16h')
    start_date = request.args.get('start_date', '2025-06-01')
    end_date = request.args.get('end_date', '2025-06-30')
    start_dt = datetime.strptime(start_date, '%Y-%m-%d').date()
    end_dt = datetime.strptime(end_date, '%Y-%m-%d').date()

    hscc_depts = [d.department_name for d in HSCCDepartment.query.all()]
    rates = {(r.ca_loai, r.truc_loai, r.ngay_loai): r.don_gia for r in ShiftRateConfig.query.all()}

    # Query lịch trực
    schedules = (
        Schedule.query.join(User).join(Shift)
        .filter(Schedule.work_date >= start_dt, Schedule.work_date <= end_dt)
        .filter(Shift.duration == (16 if ca_chon == '16h' else 24))
        .filter(~Shift.name.ilike('%thường trú%'))
        .all()
    )

    grouped = defaultdict(lambda: defaultdict(lambda: defaultdict(int)))

    for s in schedules:
        user = s.user
        dept = user.department or 'Không rõ'
        key = ("HSCC" if dept in hscc_depts else "thường", classify_day(s.work_date))
        grouped[dept][user][key] += 1

    summary_rows = []
    total = defaultdict(int)

    # Sắp xếp khoa theo tên
    for i, (dept, users) in enumerate(sorted(grouped.items(), key=lambda x: x[0]), start=1):
        # Thêm header khoa
        summary_rows.append({
            'is_dept': True,
            'index_label': f"{roman(i)}.",
            'department': dept
        })

        # Sắp xếp nhân sự trong khoa theo chức danh ưu tiên
        sorted_users = sorted(users.items(), key=lambda x: (get_priority(x[0].position), x[0].name))

        for j, (user, counts) in enumerate(sorted_users, start=1):
            row = {
                'is_dept': False,
                'index_label': str(j),
                'department': dept,
                'full_name': user.name,
                'is_contract': user.contract_type and 'hợp đồng' in user.contract_type.lower(),
                'thuong_thuong': 0,
                'hscc_thuong': 0,
                'thuong_nghi': 0,
                'hscc_nghi': 0,
                'thuong_le': 0,
                'hscc_le': 0,
                'total_shifts': 0,
                'tien_qd73': 0,
                'tien_an': 0,
                'tong_tien': 0
            }

            for key, so_ngay in counts.items():
                loai_truc, ngay_loai = key
                don_gia = rates.get((ca_chon, loai_truc, ngay_loai), 0)
                tien = so_ngay * don_gia

                if (loai_truc, ngay_loai) == ('thường', 'ngày_thường'):
                    row['thuong_thuong'] = so_ngay
                elif (loai_truc, ngay_loai) == ('HSCC', 'ngày_thường'):
                    row['hscc_thuong'] = so_ngay
                elif (loai_truc, ngay_loai) == ('thường', 'ngày_nghỉ'):
                    row['thuong_nghi'] = so_ngay
                elif (loai_truc, ngay_loai) == ('HSCC', 'ngày_nghỉ'):
                    row['hscc_nghi'] = so_ngay
                elif (loai_truc, ngay_loai) == ('thường', 'ngày_lễ'):
                    row['thuong_le'] = so_ngay
                elif (loai_truc, ngay_loai) == ('HSCC', 'ngày_lễ'):
                    row['hscc_le'] = so_ngay

                row['tien_qd73'] += tien
                row['total_shifts'] += so_ngay

            row['tien_an'] = row['total_shifts'] * 15000
            row['tong_tien'] = row['tien_qd73'] + row['tien_an']

            for k in ['thuong_thuong', 'hscc_thuong', 'thuong_nghi', 'hscc_nghi', 'thuong_le', 'hscc_le',
                      'total_shifts', 'tien_qd73', 'tien_an', 'tong_tien']:
                total[k] += row[k]

            summary_rows.append(row)

    now = datetime.now()

    return render_template(
        'shift_payment_summary_print.html',
        summary_rows=summary_rows,
        total_shifts=total['total_shifts'],
        total_qd73=total['tien_qd73'],
        total_an=total['tien_an'],
        total_sum=total['tong_tien'],
        sum_thuong_thuong=total['thuong_thuong'],
        sum_hscc_thuong=total['hscc_thuong'],
        sum_thuong_nghi=total['thuong_nghi'],
        sum_hscc_nghi=total['hscc_nghi'],
        sum_thuong_le=total['thuong_le'],
        sum_hscc_le=total['hscc_le'],
        tong_tien_bang_chu=num2text(int(total['tong_tien'])),
        thang=start_dt.month,
        nam=start_dt.year,
        now=now,
        mode=ca_chon
    )

@app.route('/configure-hscc/delete/<int:dept_id>', methods=['POST'])
def delete_hscc(dept_id):
    if session.get('role') != 'admin':
        return "Không có quyền."
    dept = HSCCDepartment.query.get_or_404(dept_id)
    db.session.delete(dept)
    db.session.commit()
    return redirect('/configure-hscc')

import unicodedata

def _normalize(s: str) -> str:
    # bỏ dấu + lower để so khớp ổn định
    s = unicodedata.normalize('NFD', s)
    s = ''.join(ch for ch in s if unicodedata.category(ch) != 'Mn')
    return s.lower().strip()

from flask import request, redirect, render_template
from datetime import datetime
import unicodedata
from sqlalchemy import func

def _normalize(s: str) -> str:
    s = unicodedata.normalize('NFD', s or '')
    s = ''.join(ch for ch in s if unicodedata.category(ch) != 'Mn')
    return s.lower().strip()

# =========================
# HAZARD CONFIG (NO MERGE)
# =========================

from datetime import datetime
from flask import request, redirect, render_template, session
# from models import db
# from models.user import User
# from models.hazard_config import HazardConfig

def _unit_normalize(unit_raw: str) -> str:
    """Chuẩn hoá đơn vị từ form → một trong: '%', 'đ', 'ngay', 'gio'."""
    u = (unit_raw or '').strip().lower()
    mapping = {
        '%': '%', 'percent': '%', 'pct': '%',
        'đ': 'đ', 'd': 'đ', 'vnđ': 'đ', 'vnd': 'đ',
        'ngày': 'ngay', 'ngay': 'ngay', 'day': 'ngay',
        'giờ': 'gio', 'gio': 'gio', 'hour': 'gio'
    }
    return mapping.get(u, u)


import unicodedata

# ---- ROUTE GIỮ NGUYÊN, CHỈ BỔ SUNG THEO YÊU CẦU ----
@app.route('/hazard-config', methods=['GET', 'POST'])
def hazard_config():
    if session.get('role') != 'admin':
        return "Bạn không có quyền truy cập.", 403

    if request.method == 'POST':
        try:
            departments = request.form.getlist('departments')
            if not departments:
                return "Chưa chọn khoa.", 400

            hazard_level = float(request.form['hazard_level'])

            unit = _unit_normalize(request.form.get('unit'))
            if unit not in ('%', 'đ', 'ngay', 'gio'):
                return "Đơn vị không hợp lệ. Chỉ chấp nhận %, đ, ngày, giờ.", 400

            duration_hours = float(request.form['duration_hours'])

            position = (request.form.get('position') or None)
            machine_type = (request.form.get('machine_type') or None)

            start_date = datetime.strptime(request.form['start_date'], '%Y-%m-%d').date()
            end_date   = datetime.strptime(request.form['end_date'],   '%Y-%m-%d').date()
            if start_date > end_date:
                return "Khoảng thời gian không hợp lệ.", 400

            # ★ NEW: cờ áp dụng riêng cho T7/CN và chọn nửa ngày/cả ngày
            weekend_only = (request.form.get('weekend_only') == 'on')          # checkbox
            weekend_part = request.form.get('weekend_part')  # 'full' | 'half' | None

            # ★ NEW: Nếu chọn T7/CN → chỉ sinh bản ghi cho các ngày T7/CN
            # ★ NEW: và CHỈ áp dụng cho khoa đã có cấu hình độc hại (HazardConfig) trước đó
            if weekend_only:
                # Ép đơn vị về giờ để thống nhất (dù user nhập gì)
                if unit != 'gio':
                    unit = 'gio'

                # Cả ngày = 17h, Nửa ngày = 7h
                weekend_hours = 17.0 if weekend_part == 'full' else 7.0

                for department in departments:
                    # ★ NEW: kiểm tra khoa đã từng có cấu hình độc hại nào chưa
                    has_hazard = db.session.query(HazardConfig.id).filter(
                        HazardConfig.department == department
                    ).first()

                    if not has_hazard:
                        # Chưa có cấu hình độc hại → bỏ qua khoa này
                        continue

                    # ★ NEW: thêm bản ghi cho từng ngày T7/CN trong khoảng
                    cur = start_date
                    while cur <= end_date:
                        # weekday(): Mon=0 ... Sun=6 → T7=5, CN=6
                        if cur.weekday() in (5, 6):
                            db.session.add(HazardConfig(
                                department=department,
                                position=position,
                                hazard_level=hazard_level,
                                unit=unit,
                                duration_hours=weekend_hours,
                                start_date=cur,   # chỉ áp dụng trong ngày
                                end_date=cur,
                                machine_type=machine_type
                            ))
                        cur += timedelta(days=1)

                db.session.commit()
                return redirect('/hazard-config')

            # --- KHÔNG GỘP: luôn tạo bản ghi mới cho từng khoa (luồng chuẩn cũ) ---
            for department in departments:
                db.session.add(HazardConfig(
                    department=department,
                    position=position,
                    hazard_level=hazard_level,
                    unit=unit,
                    duration_hours=duration_hours,
                    start_date=start_date,
                    end_date=end_date,
                    machine_type=machine_type
                ))

            db.session.commit()
            return redirect('/hazard-config')

        except Exception as e:
            db.session.rollback()
            return f"Lỗi xử lý: {e}", 400

    # --- GET: chuẩn bị dữ liệu cho form ---
    departments_raw = [
        d[0] for d in db.session.query(User.department)
        .filter(User.department.isnot(None))
        .distinct()
        .order_by(User.department)
        .all()
    ]

    # Hàm chuẩn hóa bỏ dấu tiếng Việt và lowercase
    def _normalize_local(s: str) -> str:
        s = (s or '').strip().lower()
        s = ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')
        return s

    departments = [
        {"name": d, "is_lab": ("xet nghiem" in _normalize_local(d))}
        for d in departments_raw
    ]

    machine_types = [
        ("", "Tất cả máy"),
        ("Máy huyết học", "Máy huyết học"),
        ("Máy truyền máu", "Máy truyền máu"),
        ("Máy vi sinh", "Máy vi sinh"),
    ]

    configs = HazardConfig.query.order_by(
        HazardConfig.department.asc(),
        HazardConfig.start_date.desc()
    ).all()

    return render_template(
        'hazard_config.html',
        configs=configs,
        departments=departments,
        machine_types=machine_types
    )


@app.route('/hazard-config/edit/<int:config_id>', methods=['GET', 'POST'])
def edit_hazard_config(config_id):
    if session.get('role') != 'admin':
        return "Bạn không có quyền truy cập.", 403

    config = HazardConfig.query.get_or_404(config_id)

    if request.method == 'POST':
        try:
            department = (request.form.get('department') or '').strip()
            if not department:
                return "Thiếu khoa.", 400

            hazard_level = float(request.form['hazard_level'])

            unit = _unit_normalize(request.form.get('unit'))
            if unit not in ('%', 'đ', 'ngay', 'gio'):
                return "Đơn vị không hợp lệ. Chỉ chấp nhận %, đ, ngày, giờ.", 400

            duration_hours = float(request.form['duration_hours'])
            position = (request.form.get('position') or None)
            machine_type = (request.form.get('machine_type') or None)

            start_date = datetime.strptime(request.form['start_date'], '%Y-%m-%d').date()
            end_date   = datetime.strptime(request.form['end_date'],   '%Y-%m-%d').date()
            if start_date > end_date:
                return "Khoảng thời gian không hợp lệ.", 400

            # ★ NEW: cho phép ép logic T7/CN khi chỉnh sửa
            weekend_only = (request.form.get('weekend_only') == 'on')
            weekend_part = request.form.get('weekend_part')  # 'full' | 'half' | None
            if weekend_only:
                unit = 'gio'  # ép đơn vị giờ
                duration_hours = 17.0 if weekend_part == 'full' else 7.0

            # KHÔNG GỘP: chỉ cập nhật bản hiện tại
            config.department = department
            config.position = position
            config.hazard_level = hazard_level
            config.unit = unit
            config.duration_hours = duration_hours
            config.start_date = start_date
            config.end_date = end_date
            config.machine_type = machine_type

            db.session.commit()
            return redirect('/hazard-config')

        except Exception as e:
            db.session.rollback()
            return f"Lỗi xử lý: {e}", 400

    # --- GET: dữ liệu cho form ---
    departments_raw = [
        d[0] for d in db.session.query(User.department)
        .filter(User.department.isnot(None))
        .distinct()
        .order_by(User.department)
        .all()
    ]

    def _normalize_local(s: str) -> str:
        return (s or '').lower().strip()

    departments = [
        {"name": d, "is_lab": ("xet nghiem" in _normalize_local(d))}
        for d in departments_raw
    ]

    machine_types = [
        ("", "Tất cả máy"),
        ("Máy huyết học", "Máy huyết học"),
        ("Máy truyền máu", "Máy truyền máu"),
        ("Máy vi sinh", "Máy vi sinh"),
    ]

    return render_template(
        'edit_hazard_config.html',
        config=config,
        departments=departments,
        machine_types=machine_types
    )

@app.route('/hazard-config/delete/<int:config_id>')
def delete_hazard_config(config_id):
    config = HazardConfig.query.get_or_404(config_id)
    db.session.delete(config)
    db.session.commit()
    return redirect('/hazard-config')

from models.user_machine_hazard import UserMachineHazard

from flask import jsonify
import unicodedata
from sqlalchemy import or_
from datetime import date, datetime, timedelta
import calendar, unicodedata
from collections import defaultdict

def _normalize(s: str) -> str:
    s = unicodedata.normalize('NFD', s or '')
    s = ''.join(ch for ch in s if unicodedata.category(ch) != 'Mn')
    return s.lower().strip()

# ========= /bang-doc-hai (VIEW) =========
@app.route('/bang-doc-hai', methods=['GET', 'POST'])
def bang_doc_hai():
    """
    VIEW bảng độc hại:
    - Bệnh nhiệt đới & Xét nghiệm: trực 24h -> 17h; HÔM SAU + nửa ngày -> 7h; còn lại 8h / 4h
    - Các khoa khác: 8h nếu làm (kể cả trực), 4h nếu nửa ngày
    - Chỉ hưởng khi đi làm; nghỉ / đi học / công tác... không hưởng
    - Mức %: lấy từ HazardConfig hiệu lực tại start_date, ưu tiên (chức vụ > máy > bản ghi mới hơn)
    - Cho phép lọc theo máy xét nghiệm (machine_type)
    """
    if session.get('role') not in ['admin', 'admin1', 'manager']:
        return "Bạn không có quyền truy cập."

    # ===== Inputs =====
    selected_department = request.values.get('department')
    selected_machine    = request.values.get('machine_type')  # có thể rỗng
    start_date_str      = request.values.get('start')
    end_date_str        = request.values.get('end')
    selected_user_ids   = request.values.getlist('hazard_user_ids')

    # ===== Dates =====
    from datetime import date, datetime, timedelta
    import calendar, unicodedata
    from collections import defaultdict
    from sqlalchemy import or_

    if not start_date_str or not end_date_str:
        today = date.today()
        start_date = date(today.year, today.month, 1)
        end_date   = date(today.year, today.month, calendar.monthrange(today.year, today.month)[1])
    else:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        end_date   = datetime.strptime(end_date_str,   '%Y-%m-%d').date()

    days = [start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)]

    # ===== Dropdown khoa =====
    role = session.get('role')
    user_dept = session.get('department')
    is_super = role in ('admin', 'admin1')  # admin1 nhìn như admin

    if is_super:
        departments = ['Tất cả'] + [d[0] for d in db.session.query(User.department)
                                    .filter(User.department.isnot(None))
                                    .distinct().order_by(User.department).all()]
    else:
        departments = [user_dept]

    # ===== Users =====
    users_q = User.query.filter(User.active == True)
    if is_super:
        if selected_department and selected_department != 'Tất cả':
            users_q = users_q.filter(User.department == selected_department)
    else:
        if user_dept:
            users_q = users_q.filter(User.department == user_dept)

    users = users_q.all()
    if selected_user_ids:
        ids_int = set(map(int, selected_user_ids))
        users = [u for u in users if u.id in ids_int]

    # Sắp xếp theo chức danh rồi tên
    PRIORITY_ORDER = ['GĐ','PGĐ','TK','TP','PTK','PTP','BS','BSCK1','BSCK2','ĐDT','KTV','KTVT','ĐD','NV','HL','BV']
    def _sort_key(u):
        pos = (u.position or '').upper().strip()
        for i, p in enumerate(PRIORITY_ORDER):
            if pos.startswith(p):
                return (i, (u.name or '').lower())
        return (len(PRIORITY_ORDER), (u.name or '').lower())
    users = sorted(users, key=_sort_key)

    # ===== Lịch & Ca =====
    query_start = start_date - timedelta(days=1)   # để bắt đêm trực hôm trước
    user_ids = [u.id for u in users]

    schedules_q = Schedule.query.filter(
        Schedule.work_date >= query_start,
        Schedule.work_date <= end_date
    )
    # Lọc theo danh sách user
    schedules_q = schedules_q.filter(
        Schedule.user_id.in_(user_ids)
    ) if user_ids else schedules_q.filter(db.text('1=0'))

    # Lọc theo máy xét nghiệm (nếu có chọn)
    if selected_machine:
        schedules_q = schedules_q.filter(
            or_(Schedule.machine_type == selected_machine,
                Schedule.machine_type.is_(None),
                Schedule.machine_type == '')
        )

    schedules = schedules_q.all()

    def _as_date(v):
        return v if isinstance(v, date) and not isinstance(v, datetime) else v.date()

    scheds_by_key = defaultdict(list)   # (user, date) -> list[Schedule]
    for s in schedules:
        scheds_by_key[(s.user_id, _as_date(s.work_date))].append(s)

    ca_map = {ca.id: ca for ca in Shift.query.all()}

    # ===== Hazard config =====
    hazard_cfgs = HazardConfig.query.filter(
        HazardConfig.start_date <= end_date,
        HazardConfig.end_date   >= start_date
    ).all()

    # --- CHUẨN HÓA CHUỖI ---
    def _norm(s: str) -> str:
        s = (s or '')
        s = unicodedata.normalize('NFD', s)
        s = ''.join(ch for ch in s if unicodedata.category(ch) != 'Mn')
        s = s.replace('Đ', 'D').replace('đ', 'd')
        return s.lower().strip()

    # Ưu tiên: chức vụ (2|1|0), máy (2|1|0), ngày bắt đầu (mới hơn tốt hơn)
    def pick_hazard_level(user) -> float:
        cand = [c for c in hazard_cfgs
                if c.department == user.department and c.start_date <= start_date <= c.end_date]
        if not cand:
            return 0.0

        up = (user.position or '').upper().strip()
        sel = (selected_machine or '').strip()

        def pos_score(c):
            cp = (c.position or '').upper().strip()
            if not cp: return 1
            return 2 if up.startswith(cp) else 0

        def mac_score(c):
            cm = (c.machine_type or '').strip()
            if not cm: return 1
            return 2 if cm == sel else 0

        best = max(cand, key=lambda c: (pos_score(c), mac_score(c), c.start_date or date.min))
        try:
            return float(best.hazard_level or 0.0)
        except:
            return 0.0

    # ===== Nhận diện ca =====
    def _is_oncall(sched_obj, shift_obj) -> bool:
        # theo giờ
        try: dur_shift = float(getattr(shift_obj, 'duration', 0) or 0.0)
        except: dur_shift = 0.0
        if dur_shift >= 16.5: return True
        try: dur_sched = float(getattr(sched_obj, 'work_hours', 0) or 0.0)
        except: dur_sched = 0.0
        if dur_sched >= 16.5: return True

        # theo text
        pool = []
        if shift_obj:
            pool += [getattr(shift_obj, 'name', ''), getattr(shift_obj, 'code', '')]
        if sched_obj:
            for fld in ('code','value','note','text'):
                pool.append(getattr(sched_obj, fld, '') or '')
        t = _norm(' '.join(pool))
        t_compact = t.replace(' ', '')
        if 'truc24' in t_compact or '24h' in t: return True
        return ('xd' in t) and ('24' in t)

    def _is_half(sched_obj, shift_obj) -> bool:
        h = None
        if sched_obj and sched_obj.work_hours not in (None, 0):
            try: h = float(sched_obj.work_hours)
            except: h = None
        if h is None and shift_obj:
            try: h = float(getattr(shift_obj, 'duration', 0) or 0)
            except: h = None
        if h is not None and 0 < h <= 4.5: return True
        pool = []
        if shift_obj:
            pool += [getattr(shift_obj, 'name', ''), getattr(shift_obj, 'code', '')]
        if sched_obj:
            for fld in ('code','value','note','text'):
                pool.append(getattr(sched_obj, fld, '') or '')
        t = _norm(' '.join(pool))
        return any(k in t for k in ['1/2','nua','half','nn/','nn','/x','/nt'])

    # ===== KHÔNG LÀM (off day) → không hưởng =====
    def _norm_text_for_flags(s: str) -> str:
        s = (s or '')
        s = unicodedata.normalize('NFD', s)
        s = ''.join(ch for ch in s if unicodedata.category(ch) != 'Mn')
        s = s.replace('Đ','D').replace('đ','d')
        return s.lower().strip()

    def _is_off_day(sched_obj, shift_obj) -> bool:
        """
        TRUE nếu là ngày KHÔNG LÀM → không hưởng độc hại:
        - nghỉ phép / nghỉ bù / nghỉ ốm / thai sản / nghỉ lễ...
        - đi học / tập huấn / đào tạo
        - đi công tác
        Nhận diện theo name/code/value/note/text của ca & lịch.
        """
        pool = []
        if shift_obj:
            pool += [getattr(shift_obj, 'name', ''), getattr(shift_obj, 'code', '')]
        if sched_obj:
            for fld in ('code', 'value', 'note', 'text'):
                pool.append(getattr(sched_obj, fld, '') or '')

        t = _norm_text_for_flags(' '.join(pool))
        t_compact = t.replace(' ', '')

        # Nghỉ
        if ('nghiphep' in t_compact or 'nghibu' in t_compact or 'nghiom' in t_compact or
            'omdau' in t_compact or 'thaisan' in t_compact or 'nghile' in t_compact):
            return True
        if ' nghi phep' in t or 'phep ' in t or t == 'p' or ' nghi bu' in t:
            return True

        # Học / đào tạo / tập huấn
        if any(k in t_compact for k in ['dihoc','hoc','taphuan','daotao','hocvien','huanluyen']):
            return True
        if '/h' in t or 'h/' in t:
            return True

        # Công tác
        if any(k in t_compact for k in ['congtac','ctac']):
            return True
        if '/ct' in t or 'ct/' in t:
            return True

        return False

    # Tóm tắt 1 ngày
    def _day_meta(user_id, d):
        """
        Trả: (max_hours, oncall_today, half_today, has_any, off_day)
        """
        lst = scheds_by_key.get((user_id, d), [])
        if not lst:
            return (0.0, False, False, False, False)

        max_hours = 0.0
        oncall = False
        half   = False
        off    = False

        for s in lst:
            ca = ca_map.get(getattr(s, 'shift_id', None))
            # giờ
            if s.work_hours not in (None, 0):
                try: h = float(s.work_hours)
                except: h = 0.0
            else:
                try: h = float(getattr(ca, 'duration', 0) or 0)
                except: h = 0.0
            if h > max_hours: max_hours = h

            # flags
            if _is_off_day(s, ca): off = True
            if _is_oncall(s, ca):  oncall = True
            if _is_half(s, ca):    half   = True

        return (max_hours, oncall, half, True, off)

    # ===== Build table =====
    nhom_chung, nhom_ho_ly = [], []

    for u in users:
        row = {
            'name': u.name,
            'position': u.position or '',
            'department': u.department,
            'daily_hours': [],
            'total_days': 0,
            'hazard_level': pick_hazard_level(u)
        }

        dept_key  = _norm(u.department or '')
        is_special = ('benh nhiet doi' in dept_key) or ('xet nghiem' in dept_key)

        for d in days:
            today_hours, today_oncall, today_half, today_has, today_off = _day_meta(u.id, d)
            if not today_has:
                row['daily_hours'].append('–')
                continue

            # KHÔNG LÀM → không hưởng
            if today_off:
                row['daily_hours'].append('–')
                continue

            prev_oncall = _day_meta(u.id, d - timedelta(days=1))[1]

            if is_special:
                # 17h / 7h / 8h / 4h
                if today_oncall:
                    desired = 17
                elif prev_oncall and today_half:
                    desired = 7
                elif today_half:
                    desired = 4
                else:
                    desired = 8 if (today_hours >= 7.5 or today_hours >= 6) else (4 if today_hours > 0 else 0)
            else:
                # Khoa khác: 8/4 (trực 24h vẫn 8h)
                desired = 4 if today_half else (8 if (today_hours > 0 or today_oncall) else 0)

            row['daily_hours'].append(f"{int(desired)}h" if desired else "–")
            if desired:
                row['total_days'] += 1

        (nhom_ho_ly if (row['position'] or '').upper().startswith('HL') else nhom_chung).append(row)

    return render_template(
        'bang_doc_hai.html',
        nhom_chung=nhom_chung,
        nhom_ho_ly=nhom_ho_ly,
        departments=departments,
        selected_department=selected_department,
        start=start_date.strftime('%Y-%m-%d'),
        end=end_date.strftime('%Y-%m-%d'),
        days_in_month=days,
        all_users=users,
        selected_user_ids=selected_user_ids,
        selected_machine=selected_machine
    )

@app.route('/machines-by-department', endpoint='machines_by_department', methods=['GET'])
def machines_by_department():
    from sqlalchemy import func
    import unicodedata  # ← cần import ở đây

    dept = (request.args.get('department') or '').strip()
    if not dept or dept == 'Tất cả':
        role = session.get('role')
        user_dept = session.get('department')
        # cho admin1 giống admin
        if role not in ('admin', 'admin1') and user_dept:
            dept = user_dept
        else:
            return jsonify([])

    q1 = (
        db.session.query(Schedule.machine_type)
        .join(User, User.id == Schedule.user_id)
        .filter(
            User.department == dept,
            Schedule.machine_type.isnot(None),
            func.trim(Schedule.machine_type) != ''
        )
        .distinct()
    )

    q2 = (
        db.session.query(HazardConfig.machine_type)
        .filter(
            HazardConfig.department == dept,
            HazardConfig.machine_type.isnot(None),
            func.trim(HazardConfig.machine_type) != ''
        )
        .distinct()
    )

    raw = [r[0] for r in q1.all()] + [r[0] for r in q2.all()]

    def _normalize_no_accent(s: str) -> str:
        s = (s or '').strip()
        s = unicodedata.normalize('NFD', s)
        return ''.join(ch for ch in s if unicodedata.category(ch) != 'Mn').casefold()

    best_by_key = {}
    for name in raw:
        key = _normalize_no_accent(name)
        if key not in best_by_key or len(name) > len(best_by_key[key]):
            best_by_key[key] = name.strip()

    return jsonify(sorted(best_by_key.values(), key=_normalize_no_accent))

@app.route('/user-machine-hazard/delete/<int:id>', methods=['POST'])
def delete_user_machine_hazard(id):
    if session.get('role') not in ['admin', 'admin1', 'manager']:
        return "Bạn không có quyền truy cập."
    mapping = UserMachineHazard.query.get_or_404(id)
    db.session.delete(mapping)
    db.session.commit()
    return redirect('/user-machine-hazard')

from flask import render_template, request, session
from datetime import datetime, timedelta, date
import calendar
from models import User, Schedule, Shift, HazardConfig

@app.route('/bang-doc-hai/print', methods=['POST'])
def bang_doc_hai_print():
    # Quyền
    if session.get('role') not in ['admin', 'admin1', 'manager']:
        return "Bạn không có quyền truy cập."

    selected_department = request.values.get('department')
    selected_machine    = request.values.get('machine_type')
    selected_user_ids   = request.values.getlist('hazard_user_ids')
    start               = request.values.get('start')
    end                 = request.values.get('end')

    from datetime import date, datetime, timedelta
    import calendar, unicodedata
    from collections import defaultdict
    from sqlalchemy import or_

    if not start or not end:
        today = date.today()
        start_date = date(today.year, today.month, 1)
        end_date   = date(today.year, today.month, calendar.monthrange(today.year, today.month)[1])
    else:
        start_date = datetime.strptime(start, '%Y-%m-%d').date()
        end_date   = datetime.strptime(end,   '%Y-%m-%d').date()

    days_range = [start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)]
    query_start = start_date - timedelta(days=1)

    # Users
    users_q = User.query.filter(User.active == True)
    if selected_department and selected_department != 'Tất cả':
        users_q = users_q.filter(User.department == selected_department)
    users = users_q.all()
    if selected_user_ids:
        ids_int = set(map(int, selected_user_ids))
        users = [u for u in users if u.id in ids_int]

    PRIORITY_ORDER = ['GĐ','PGĐ','TK','TP','PTK','PTP','BS','BSCK1','BSCK2','ĐDT','KTV','KTVT','ĐD','NV','HL','BV']
    def _sort_key(u):
        pos = (u.position or '').upper().strip()
        for i,p in enumerate(PRIORITY_ORDER):
            if pos.startswith(p):
                return (i, (u.name or '').lower())
        return (len(PRIORITY_ORDER), (u.name or '').lower())
    users = sorted(users, key=_sort_key)

    # Lịch & ca
    schedules_q = Schedule.query.filter(
        Schedule.work_date >= query_start,
        Schedule.work_date <= end_date
    )
    if selected_machine:
        schedules_q = schedules_q.filter(
            or_(Schedule.machine_type == selected_machine,
                Schedule.machine_type.is_(None),
                Schedule.machine_type == '')
        )
    schedules = schedules_q.all()
    ca_map = {ca.id: ca for ca in Shift.query.all()}

    def _as_date(v):
        return v if isinstance(v, date) and not isinstance(v, datetime) else v.date()
    scheds_by_key = defaultdict(list)
    for s in schedules:
        scheds_by_key[(s.user_id, _as_date(s.work_date))].append(s)

    # % cấu hình
    hazard_cfgs = HazardConfig.query.filter(
        HazardConfig.start_date <= end_date,
        HazardConfig.end_date   >= start_date
    ).all()

    def _norm(s: str) -> str:
        s = (s or '')
        s = unicodedata.normalize('NFD', s)
        s = ''.join(ch for ch in s if unicodedata.category(ch) != 'Mn')
        s = s.replace('Đ', 'D').replace('đ', 'd')
        return s.lower().strip()

    def pick_hazard_level(user) -> float:
        cand = [c for c in hazard_cfgs
                if c.department == user.department and c.start_date <= start_date <= c.end_date]
        if not cand:
            return 0.0
        up  = (user.position or '').upper().strip()
        sel = (selected_machine or '').strip()

        def pos_score(c):
            cp = (c.position or '').upper().strip()
            if not cp: return 1
            return 2 if up.startswith(cp) else 0

        def mac_score(c):
            cm = (c.machine_type or '').strip()
            if not cm: return 1
            return 2 if cm == sel else 0

        best = max(cand, key=lambda c: (pos_score(c), mac_score(c), c.start_date or date.min))
        try:
            return float(best.hazard_level or 0.0)
        except:
            return 0.0

    # Nhận diện ca/off tương tự VIEW
    OFF_KEYWORDS = [
        'nghi phep','phep','np','nb',
        'om','om dau','benh','dieu tri',
        'thai san','ts',
        'cong tac','ct ',
        'dao tao','hoc','tap huan','boi duong',
        'nghi bu','nghi le','le',
    ]
    def _is_off_day(sched_obj, shift_obj) -> bool:
        pool = []
        if shift_obj:
            pool += [getattr(shift_obj, 'name', ''), getattr(shift_obj, 'code', '')]
        if sched_obj:
            for fld in ('code','value','note','text'):
                pool.append(getattr(sched_obj, fld, '') or '')
        t = ' ' + _norm(' '.join(pool)) + ' '
        def has_kw(kw):
            kw = kw.strip()
            return (' ' + kw + ' ') in t or t.endswith(' ' + kw) or t.startswith(' ' + kw + ' ')
        return any(has_kw(kw) for kw in OFF_KEYWORDS)

    def _is_oncall(sched_obj, shift_obj) -> bool:
        try: dur_shift = float(getattr(shift_obj, 'duration', 0) or 0.0)
        except: dur_shift = 0.0
        if dur_shift >= 16.5: return True
        try: dur_sched = float(getattr(sched_obj, 'work_hours', 0) or 0.0)
        except: dur_sched = 0.0
        if dur_sched >= 16.5: return True
        pool = []
        if shift_obj:
            pool += [getattr(shift_obj, 'name', ''), getattr(shift_obj, 'code', '')]
        if sched_obj:
            for fld in ('code','value','note','text'):
                pool.append(getattr(sched_obj, fld, '') or '')
        t = _norm(' '.join(pool))
        t_compact = t.replace(' ', '')
        if 'truc24' in t_compact or '24h' in t: return True
        return ('xd' in t and '24' in t)

    def _is_half(sched_obj, shift_obj) -> bool:
        h = None
        if sched_obj and sched_obj.work_hours not in (None, 0):
            try: h = float(sched_obj.work_hours)
            except: h = None
        if h is None and shift_obj:
            try: h = float(getattr(shift_obj, 'duration', 0) or 0)
            except: h = None
        if h is not None and 0 < h <= 4.5: return True
        pool = []
        if shift_obj:
            pool += [getattr(shift_obj, 'name', ''), getattr(shift_obj, 'code', '')]
        if sched_obj:
            for fld in ('code','value','note','text'):
                pool.append(getattr(sched_obj, fld, '') or '')
        t = _norm(' '.join(pool))
        return any(k in t for k in ['1/2','nua','half','nn/','nn','/x','/nt'])

    def _is_full_day_hours(h: float) -> bool:
        try: return float(h or 0) >= 7.5
        except: return False

    def _day_meta(user_id, d):
        lst = scheds_by_key.get((user_id, d), [])
        if not lst:
            return (0.0, False, False, False, False)
        max_hours = 0.0
        oncall = False
        half   = False
        off    = False
        for s in lst:
            ca = ca_map.get(getattr(s, 'shift_id', None))
            if s.work_hours not in (None, 0):
                try: h = float(s.work_hours)
                except: h = 0.0
            else:
                try: h = float(getattr(ca, 'duration', 0) or 0)
                except: h = 0.0
            if h > max_hours: max_hours = h
            if _is_oncall(s, ca): oncall = True
            if _is_half(s, ca):   half   = True
            if _is_off_day(s, ca): off  = True
        return (max_hours, oncall, half, True, off)

    # Build table
    table_data = []
    for u in users:
        row = {
            'name': u.name,
            'position': u.position or '',
            'hazard_level': pick_hazard_level(u),
            'daily_hours': [],
            'total_days': 0
        }

        dept_key = _norm(u.department or '')
        is_special = ('benh nhiet doi' in dept_key) or ('xet nghiem' in dept_key)

        for d in days_range:
            today_hours, today_oncall, today_half, today_has, today_off = _day_meta(u.id, d)
            if not today_has or today_off:
                row['daily_hours'].append('–')
                continue

            prev_oncall = _day_meta(u.id, d - timedelta(days=1))[1]

            if is_special:
                if today_oncall:
                    desired = 17
                elif prev_oncall and today_half:
                    desired = 7
                elif today_half:
                    desired = 4
                else:
                    desired = 8 if (_is_full_day_hours(today_hours) or today_hours >= 6) else (4 if today_hours > 0 else 0)
            else:
                desired = 4 if today_half else (8 if (today_hours > 0 or today_oncall) else 0)

            row['daily_hours'].append(f"{int(desired)}h" if desired else "–")
            if desired:
                row['total_days'] += 1

        table_data.append(row)

    return render_template(
        'bang_doc_hai_print.html',
        table_data=table_data,
        department=selected_department or 'Tất cả',
        start=start_date,
        end=end_date,
        days_range=days_range,
        now=datetime.now()
    )

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from io import BytesIO
from datetime import timedelta

def export_bang_doc_hai_excel_file(users, schedules, shifts, hazard_configs, start_date, end_date, selected_user_ids):
    wb = Workbook()
    ws = wb.active

    bold_font = Font(bold=True)
    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    ws.merge_cells('A1:AF1')
    ws['A1'] = "BỆNH VIỆN NHI TỈNH GIA LAI"
    ws['A1'].font = bold_font

    ws.merge_cells('A2:AF2')
    ws['A2'] = f"KHOA {users[0].department.upper()}" if users else ""
    ws['A2'].font = bold_font

    ws.merge_cells('A3:AF3')
    ws['A3'] = "BẢNG CHẤM CÔNG HƯỞNG MỨC ĐỘC HẠI 0.2"
    ws['A3'].font = Font(bold=True, size=14)
    ws['A3'].alignment = center_align

    ws.merge_cells('A4:AF4')
    ws['A4'] = f"Từ ngày {start_date.strftime('%d/%m/%Y')} đến {end_date.strftime('%d/%m/%Y')}"

    headers = ['STT', 'Họ tên', 'Chức vụ']
    num_days = (end_date - start_date).days + 1
    headers.extend([(start_date + timedelta(days=i)).day for i in range(num_days)])
    headers.append('Tổng ngày')
    ws.append(headers)

    shift_dict = {s.id: s for s in shifts}
    schedule_map = {(s.user_id, s.work_date): s for s in schedules}

    for idx, user in enumerate(users, 1):
        if selected_user_ids and str(user.id) not in selected_user_ids:
            continue

        row = [idx, user.name, user.position]
        total = 0

        for i in range(num_days):
            date_i = start_date + timedelta(days=i)
            sched = schedule_map.get((user.id, date_i))
            if not sched or not sched.shift_id:
                row.append('–')
                continue

            ca = shift_dict.get(sched.shift_id)
            if not ca:
                row.append('–')
                continue

            configs_in_day = [cfg for cfg in hazard_configs if cfg.department == user.department and cfg.start_date <= date_i <= cfg.end_date]
            if not configs_in_day:
                row.append('–')
                continue

            best_match = sorted(configs_in_day, key=lambda cfg: abs(cfg.duration_hours - ca.duration))[0]
            row.append(f"{int(best_match.duration_hours)}h")
            total += 1

        row.append(f"{total} ngày")
        ws.append(row)

    for row in ws.iter_rows(min_row=5, max_row=ws.max_row, max_col=ws.max_column):
        for cell in row:
            cell.alignment = center_align
            cell.border = thin_border

    sign_row = ws.max_row + 3
    ws.merge_cells(start_row=sign_row, start_column=1, end_row=sign_row, end_column=5)
    ws.merge_cells(start_row=sign_row, start_column=6, end_row=sign_row, end_column=10)
    ws.merge_cells(start_row=sign_row, start_column=11, end_row=sign_row, end_column=15)
    ws.merge_cells(start_row=sign_row, start_column=16, end_row=sign_row, end_column=20)

    ws.cell(row=sign_row, column=1, value="NGƯỜI LẬP BẢNG\n(Ký, ghi rõ họ tên)").alignment = center_align
    ws.cell(row=sign_row, column=6, value="TRƯỞNG KHOA\n(Ký, ghi rõ họ tên)").alignment = center_align
    ws.cell(row=sign_row, column=11, value="PHÒNG TỔ CHỨC - HCQT\n(Ký, ghi rõ họ tên)").alignment = center_align
    ws.cell(row=sign_row, column=16, value="GIÁM ĐỐC\n(Ký, ghi rõ họ tên)").alignment = center_align

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output

# ✅ ROUTE EXPORT FILE EXCEL
@app.route('/bang-doc-hai/export-excel', methods=['POST'], endpoint='export_bang_doc_hai_excel')
def export_bang_doc_hai_excel():
    # Cho phép admin, admin1 và manager xuất Excel
    if session.get('role') not in ['admin', 'admin1', 'manager']:
        return "Bạn không có quyền truy cập."

    selected_department = request.values.get('department')
    start_date = datetime.strptime(request.values.get('start'), '%Y-%m-%d').date()
    end_date = datetime.strptime(request.values.get('end'), '%Y-%m-%d').date()
    selected_user_ids = request.values.getlist('hazard_user_ids')

    # Lấy danh sách user theo khoa
    users = User.query.filter(User.active == True)
    if selected_department and selected_department != 'Tất cả':
        users = users.filter(User.department == selected_department)
    users = users.all()

    # Lấy dữ liệu lịch trực và hazard config
    schedules = Schedule.query.filter(
        Schedule.work_date >= start_date,
        Schedule.work_date <= end_date
    ).all()

    hazard_configs = HazardConfig.query.filter(
        HazardConfig.start_date <= end_date,
        HazardConfig.end_date >= start_date
    ).all()

    shifts = Shift.query.all()

    # Xuất file Excel
    output = export_bang_doc_hai_excel_file(
        users, schedules, shifts, hazard_configs,
        start_date, end_date, selected_user_ids
    )

    filename = f"bang_doc_hai_{selected_department or 'tatca'}.xlsx"
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

from flask import request, render_template
from datetime import datetime, timedelta
from collections import defaultdict
from dateutil.relativedelta import relativedelta

@app.route('/theo-doi-nghi-bu', methods=['GET'], endpoint='theo_doi_nghi_bu')
def report_compensations():
    from collections import defaultdict
    from dateutil.relativedelta import relativedelta
    from flask_login import current_user

    # Kiểm tra quyền admin
    is_admin = getattr(current_user, 'role', '') == 'admin'

    # Lấy khoa phòng
    if is_admin:
        selected_department = request.args.get('department')
        if not selected_department:
            first_dept = db.session.query(User.department).filter(User.department.isnot(None)).first()
            if first_dept:
                selected_department = first_dept[0]
    else:
        selected_department = getattr(current_user, 'department', None)

    if not selected_department:
        return "Không xác định được khoa để hiển thị."

    # Ngày bắt đầu/kết thúc
    start = request.args.get('start')
    end = request.args.get('end')
    try:
        start_date = datetime.strptime(start, '%Y-%m-%d').date() if start else datetime.today().replace(day=1)
        end_date = datetime.strptime(end, '%Y-%m-%d').date() if end else datetime.today()
    except ValueError:
        return "Ngày không hợp lệ. Định dạng cần là YYYY-MM-DD."

    # Dải ngày
    days_range = []
    cur_day = start_date
    while cur_day <= end_date:
        days_range.append(cur_day)
        cur_day += timedelta(days=1)

    # Các mã ca trực
    CA_TRUC_CODES = ['XĐ', 'XĐ16', 'XĐ24', 'XĐ3', 'XĐL16', 'XĐL24', 'XĐT']

    def is_ca_truc(name):
        return name.startswith('TRỰC') or name in CA_TRUC_CODES

    # Hàm nhận diện loại nghỉ
    def detect_nghi_type(name):
        name = name.upper().strip()

        # NB nửa ngày (NBS, NBC hoặc có 1/2) → hiển thị /X
        if 'NBS' in name or 'NBC' in name or '1/2' in name:
            return '/X'

        # NB nguyên ngày
        if name == 'NB' or 'NGHỈ BÙ' in name:
            return 'NB'

        # Nghỉ trực nguyên ngày
        if 'NT' in name or 'NGHỈ TRỰC' in name:
            return 'NT'

        # Nghỉ 1/2 ngày khác
        if '/X' in name:
            return '/X'

        return None

    # Tính số dư tháng trước
    first_day_this_month = start_date.replace(day=1)
    first_day_last_month = first_day_this_month - relativedelta(months=1)
    last_day_last_month = first_day_this_month - timedelta(days=1)

    query_last = Schedule.query.join(User).join(Shift).filter(
        Schedule.work_date.between(first_day_last_month, last_day_last_month),
        User.department == selected_department
    )

    prev_comp = defaultdict(float)
    for s in query_last.all():
        shift_name = s.shift.name.strip().upper()
        nghi_type = detect_nghi_type(shift_name)
        if is_ca_truc(shift_name):
            prev_comp[s.user_id] += 1
        if nghi_type == 'NB':
            prev_comp[s.user_id] -= 1
        elif nghi_type == '/X':
            prev_comp[s.user_id] -= 0.5
        elif nghi_type == 'NT':
            prev_comp[s.user_id] -= 1

    # Dữ liệu tháng hiện tại
    query = Schedule.query.join(User).join(Shift).filter(
        Schedule.work_date.between(start_date, end_date),
        User.department == selected_department
    )

    users_data = defaultdict(lambda: {
        'name': '',
        'position': '',
        'department': '',
        'days': {},
        'prev_total': 0,
        'remain': 0
    })

    schedules_by_user = defaultdict(list)
    for s in query.all():
        schedules_by_user[s.user_id].append(s)

    for uid in schedules_by_user:
        schedules_by_user[uid] = sorted(schedules_by_user[uid], key=lambda x: x.work_date)

    # Xử lý từng nhân viên
    for uid, schedules in schedules_by_user.items():
        user = schedules[0].user
        users_data[uid]['name'] = user.name
        users_data[uid]['position'] = user.position
        users_data[uid]['department'] = user.department
        users_data[uid]['prev_total'] = prev_comp.get(uid, 0)

        total = 0
        for s in schedules:
            day_str = s.work_date.strftime('%d')
            shift_name = s.shift.name.strip().upper()
            nghi_type = detect_nghi_type(shift_name)

            # Ẩn chữ "LÀM NGÀY" (bỏ qua không hiển thị)
            if 'LÀM NGÀY' in shift_name or shift_name == 'X':
                continue

            if is_ca_truc(shift_name):
                users_data[uid]['days'][day_str] = 'XĐ'
                total += 1
            elif nghi_type:
                users_data[uid]['days'][day_str] = nghi_type
                if nghi_type == 'NB':
                    total -= 1
                elif nghi_type == '/X':
                    total -= 0.5
                elif nghi_type == 'NT':
                    total -= 1

        users_data[uid]['remain'] = total + users_data[uid]['prev_total']

    # Sắp xếp
    position_order = ['GĐ', 'PGĐ', 'TK', 'PTK', 'BS', 'ĐDT', 'ĐD', 'NV', 'HL', 'BV']
    sorted_users = sorted(
        users_data.values(),
        key=lambda x: (
            position_order.index(x['position']) if x['position'] in position_order else 99,
            x['name']
        )
    )

    # Highlight ngày nghỉ
    fixed_holidays = [(4, 30), (5, 1), (9, 2), (1, 1)]
    highlight_days = []
    for i, d in enumerate(days_range):
        if d.weekday() in [5, 6] or (d.month, d.day) in fixed_holidays:
            highlight_days.append(i)

    departments = []
    if is_admin:
        departments = [row[0] for row in db.session.query(User.department).distinct().all() if row[0]]

    return render_template(
        'report_compensations.html',
        days_range=days_range,
        users_data=sorted_users,
        highlight_days=highlight_days,
        departments=departments,
        selected_department=selected_department,
        is_admin=is_admin,
        start=start_date.strftime('%Y-%m-%d'),
        end=end_date.strftime('%Y-%m-%d')
    )


@app.before_first_request
def create_missing_tables():
    from models.hazard_config import HazardConfig
    db.create_all()
    print("✅ Đã kiểm tra và tạo các bảng thiếu (nếu có)")

from models.user import User

from sqlalchemy import inspect

@app.route('/init-db')
def init_db():
    inspector = inspect(db.engine)
    existing_tables = inspector.get_table_names()

    # Chỉ tạo bảng nếu chưa tồn tại
    if 'user' not in existing_tables:
        db.create_all()
        return "✅ Đã tạo bảng vào PostgreSQL"
    return "⚠️ Bảng đã tồn tại, không tạo lại."


@app.route('/run-seed')
def run_seed():
    try:
        import seed
        return "✅ Đã chạy seed.py thành công!"
    except Exception as e:
        return f"❌ Lỗi khi chạy seed.py: {str(e)}"

import time

@app.before_request
def start_timer():
    request.start_time = time.time()

@app.after_request
def log_request_time(response):
    duration = time.time() - request.start_time
    app.logger.info(f"[REQUEST TIME] {request.method} {request.path} - {duration:.3f}s")
    return response

@app.cli.command("db-upgrade")
def db_upgrade():
    from flask_migrate import upgrade
    upgrade()

@app.errorhandler(500)
def internal_error(error):
    return render_template('500.html'), 500

import traceback

@app.errorhandler(Exception)
def handle_exception(e):
    return f"<h2>Internal Server Error</h2><pre>{traceback.format_exc()}</pre>", 500


@app.before_first_request
def setup_initial_data():
    import os
    from sqlalchemy import inspect

    inspector = inspect(db.engine)
    existing_tables = inspector.get_table_names()
    required_tables = {'user', 'permission'}

    # ✅ Tạo bảng nếu thiếu
    if not required_tables.issubset(set(existing_tables)):
        from models.permission import Permission
        from models.user import User
        db.create_all()
        print("✅ Đã tạo tất cả bảng cần thiết.")
    else:
        print("✅ Các bảng chính đã tồn tại.")

    # ✅ Thêm cột 'active' nếu chưa có
    try:
        with db.engine.connect() as connection:
            result = connection.execute(
                "SELECT column_name FROM information_schema.columns WHERE table_name='user' AND column_name='active';"
            )
            if not result.fetchone():
                connection.execute('ALTER TABLE "user" ADD COLUMN active BOOLEAN DEFAULT TRUE;')
                print("✅ Đã thêm cột 'active' vào bảng user.")
            else:
                print("✅ Cột 'active' đã tồn tại.")
    except Exception as e:
        print(f"❌ Lỗi khi kiểm tra/thêm cột 'active': {e}")

    # ✅ Thêm tài khoản admin nếu chưa có
    from models.user import User
    if not User.query.filter_by(username='admin').first():
        admin = User(
            name="Quản trị viên",
            username="admin",
            password="admin",
            role="admin",
            department="Phòng CNTT",
            position="Bác sĩ"
        )
        db.session.add(admin)
        db.session.commit()
        print("✅ Đã tạo tài khoản admin.")
    else:
        print("⚠️ Tài khoản admin đã tồn tại.")

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)



